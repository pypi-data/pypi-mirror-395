"""
DOCX Exporter - Converts TipTap JSON back to Word document format.

This module handles the reverse conversion from TipTap editor JSON
back to a Word document (.docx) using python-docx.

Supports:
- Basic text formatting (bold, italic)
- Headings (H1-H6)
- Tables
- Track changes (insertions/deletions) via OOXML
- Comments via python-docx native API
"""

import copy
import json
from io import BytesIO
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .utils import base64_to_element


class DocxExporter:
    """
    Exports TipTap JSON documents to Word (.docx) format.

    This class encapsulates the conversion logic and maintains state
    (like revision counter, comment tracking) as instance variables
    for thread safety.
    """

    def __init__(
        self,
        comments: Optional[list[dict]] = None,
        template_bytes: Optional[bytes] = None,
    ):
        """
        Initialize the exporter.

        Args:
            comments: Optional list of comment dictionaries with id, author, text, date
            template_bytes: Optional bytes of a .docx file to use as template.
                If provided, the template's styles will be preserved but content
                will be replaced with the TipTap content.
        """
        self._comments_dict = {c["id"]: c for c in (comments or [])}
        self._template_bytes = template_bytes
        self._revision_id_counter = 0
        self._comment_runs_map: dict[str, list] = {}
        self._doc: Optional[Document] = None

    def export(self, tiptap_json: dict) -> BytesIO:
        """
        Convert TipTap JSON document to a Word document.

        Args:
            tiptap_json: TipTap document JSON with structure:
                {
                    "type": "doc",
                    "content": [...]
                }

        Returns:
            BytesIO buffer containing the .docx file
        """
        # Reset state for this export
        self._revision_id_counter = 0
        self._comment_runs_map = {}

        # Recalculate numbering based on numIlvl values (in case levels changed in TipTap)
        tiptap_json = self._recalculate_numbering(tiptap_json)

        # Restore raw styles from storage node
        tiptap_json = self._restore_raw_styles(tiptap_json)

        # Create document from template or blank
        if self._template_bytes:
            template_buffer = BytesIO(self._template_bytes)
            self._doc = Document(template_buffer)
            self._clear_document_content()
        else:
            self._doc = Document()

        # Process content
        content = tiptap_json.get("content", [])
        for node in content:
            self._process_node(node)

        # Add comments after processing all content
        self._add_comments_to_document()

        # Save to BytesIO buffer
        buffer = BytesIO()
        self._doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _next_revision_id(self) -> str:
        """Generate next revision ID."""
        self._revision_id_counter += 1
        return str(self._revision_id_counter)

    def _clear_document_content(self) -> None:
        """
        Clear all content from a document while preserving styles.

        This removes paragraphs and tables from the document body but keeps
        styles, headers, footers, and other document properties intact.
        """
        body = self._doc.element.body

        # Remove all paragraph and table elements from the body
        # We iterate in reverse to avoid issues with modifying during iteration
        for child in list(body):
            # Keep section properties (w:sectPr) as they contain page layout
            if child.tag.endswith("}sectPr"):
                continue
            body.remove(child)

    def _restore_raw_styles(self, tiptap_json: dict) -> dict:
        """
        Restore raw OOXML styles from rawStylesStorage node back into content.

        Finds the storage node, deserializes the data, and puts rawTblPr,
        rawTblGrid, rawXml, rawPPr back into the appropriate nodes.
        """
        doc = copy.deepcopy(tiptap_json)
        content = doc.get("content", [])

        # Find and extract rawStylesStorage node
        raw_styles = {}
        new_content = []
        for node in content:
            if node.get("type") == "rawStylesStorage":
                data = node.get("attrs", {}).get("data", "{}")
                raw_styles = json.loads(data)
            else:
                new_content.append(node)

        doc["content"] = new_content

        if not raw_styles:
            return doc

        # Restore styles to nodes
        para_counter = [0]  # Use list for mutable counter in nested function

        def process_node(node: dict):
            node_type = node.get("type")

            # Restore paragraph/heading rawPPr (numbering overrides, etc.)
            if node_type in ("paragraph", "heading"):
                para_key = f"para:{para_counter[0]}:pPr"
                if para_key in raw_styles:
                    if "attrs" not in node:
                        node["attrs"] = {}
                    node["attrs"]["rawPPr"] = raw_styles[para_key]
                para_counter[0] += 1

            if node_type == "table":
                attrs = node.get("attrs", {})
                table_id = attrs.get("id")
                if table_id:
                    # Restore table-level raw XML
                    if f"table:{table_id}:tblPr" in raw_styles:
                        attrs["rawTblPr"] = raw_styles[f"table:{table_id}:tblPr"]
                    if f"table:{table_id}:tblGrid" in raw_styles:
                        attrs["rawTblGrid"] = raw_styles[f"table:{table_id}:tblGrid"]

                    # Restore row and cell styles
                    for row_idx, row in enumerate(node.get("content", [])):
                        row_key = f"table:{table_id}:row:{row_idx}"
                        if row_key in raw_styles:
                            if "attrs" not in row:
                                row["attrs"] = {}
                            row["attrs"]["rawXml"] = raw_styles[row_key]

                        for cell_idx, cell in enumerate(row.get("content", [])):
                            cell_key = f"table:{table_id}:row:{row_idx}:cell:{cell_idx}"
                            if cell_key in raw_styles:
                                if "attrs" not in cell:
                                    cell["attrs"] = {}
                                cell["attrs"]["rawXml"] = raw_styles[cell_key]

            for child in node.get("content", []):
                if isinstance(child, dict):
                    process_node(child)

        for node in doc.get("content", []):
            process_node(node)

        return doc

    def _recalculate_numbering(self, tiptap_json: dict) -> dict:
        """
        Recalculate styleNumbering for all paragraphs based on numIlvl values.

        This ensures numbering is correct even after Tab/Shift-Tab changes in TipTap
        that may have changed numIlvl values without proper recalculation of the
        display numbers.

        IMPORTANT: This also updates styleName when numIlvl differs from what the
        style would normally provide. Word uses style-based numbering where different
        styles (SH1Legal, SH2Legal, SH3Legal, etc.) correspond to different levels.
        When the user changes the level in TipTap, we need to change the style.

        The algorithm:
        1. Extract the style-numbering map from rawStylesStorage
        2. Group paragraphs by numId
        3. For each paragraph, if numIlvl differs from the style's default ilvl,
           find and apply the correct style for that numId+ilvl combination
        4. Recalculate the display numbers
        """
        doc = copy.deepcopy(tiptap_json)

        # Extract the style-numbering map from rawStylesStorage
        style_numbering_map = {}
        for node in doc.get("content", []):
            if node.get("type") == "rawStylesStorage":
                data = node.get("attrs", {}).get("data", "{}")
                raw_styles = json.loads(data)
                style_numbering_map = raw_styles.get("__style_numbering_map__", {})
                break

        style_to_num = style_numbering_map.get("style_to_num", {})
        num_to_style = style_numbering_map.get("num_to_style", {})

        # Collect all numbered paragraphs grouped by numId
        # We need to track their position in the document for ordering
        numbered_paras: dict[str, list[tuple[int, dict]]] = {}  # numId -> [(index, node), ...]
        para_index = [0]  # Mutable counter for document order

        def collect_numbered(node: dict, parent_content: list = None, node_index: int = 0):
            """Collect numbered paragraphs in document order."""
            node_type = node.get("type")

            if node_type in ("paragraph", "heading"):
                attrs = node.get("attrs", {})
                num_id = attrs.get("numId")
                if num_id:
                    if num_id not in numbered_paras:
                        numbered_paras[num_id] = []
                    numbered_paras[num_id].append((para_index[0], node))
                para_index[0] += 1

            # Recurse into content
            for i, child in enumerate(node.get("content", [])):
                if isinstance(child, dict):
                    collect_numbered(child, node.get("content", []), i)

        for node in doc.get("content", []):
            collect_numbered(node)

        # Now recalculate numbering for each numId group
        for num_id, paras in numbered_paras.items():
            # paras are already in document order (by para_index)
            paras.sort(key=lambda x: x[0])

            # Track counters at each level
            counters: list[int] = []
            parent_prefixes: list[str] = [""]
            last_level = -1

            for _, node in paras:
                attrs = node.get("attrs", {})
                level = attrs.get("numIlvl", 0)
                current_style = attrs.get("styleName")

                # Check if we need to update the style based on numIlvl
                # Word uses style-based numbering: SH1Legal->ilvl=0, SH2Legal->ilvl=1, etc.
                # When user changes level via Tab/Shift-Tab, we need to change the style
                if current_style and num_to_style and num_id in num_to_style:
                    # Get what level the current style would provide
                    style_info = style_to_num.get(current_style, {})
                    style_default_ilvl = style_info.get("ilvl")

                    # If the current numIlvl differs from the style's default, find correct style
                    if style_default_ilvl is not None and style_default_ilvl != level:
                        # Look up the correct style for this numId + ilvl combination
                        # Note: JSON keys are strings, so convert level to string
                        correct_style = num_to_style.get(num_id, {}).get(str(level))
                        if correct_style:
                            attrs["styleName"] = correct_style

                if level > last_level:
                    # Going deeper - initialize new level counters
                    for i in range(last_level + 1, level + 1):
                        while len(counters) <= i:
                            counters.append(0)
                        counters[i] = 1
                        # Build parent prefix for this level
                        while len(parent_prefixes) <= i:
                            parent_prefixes.append("")
                        if i > 0:
                            parent_num = counters[i - 1] if i - 1 < len(counters) else 1
                            parent_prefixes[i] = parent_prefixes[i - 1] + str(parent_num) + "."
                        else:
                            parent_prefixes[i] = ""
                elif level < last_level:
                    # Going shallower - increment counter at this level
                    while len(counters) <= level:
                        counters.append(0)
                    counters[level] = counters[level] + 1
                    # Reset deeper level counters
                    for i in range(level + 1, len(counters)):
                        counters[i] = 0
                else:
                    # Same level - increment counter
                    while len(counters) <= level:
                        counters.append(0)
                    counters[level] = counters[level] + 1

                # Build the new number
                prefix = parent_prefixes[level] if level < len(parent_prefixes) else ""
                new_number = prefix + str(counters[level]) + "."

                # Update the node's styleNumbering
                attrs["styleNumbering"] = new_number
                last_level = level

        return doc

    def _process_node(self, node: dict, table_cell=None) -> None:
        """
        Process a TipTap node and add it to the document.

        Args:
            node: A TipTap node dictionary
            table_cell: Optional table cell to add content to (for nested content)
        """
        node_type = node.get("type")

        if node_type == "paragraph":
            self._process_paragraph(node, table_cell)
        elif node_type == "heading":
            self._process_heading(node, table_cell)
        elif node_type == "table":
            self._process_table(node, table_cell)
        elif node_type == "section":
            self._process_section(node)

    def _process_paragraph(self, node: dict, table_cell=None) -> None:
        """Process a paragraph node."""
        attrs = node.get("attrs", {})
        style_name = attrs.get("styleName")
        raw_pPr = attrs.get("rawPPr")
        num_ilvl = attrs.get("numIlvl")  # Current indentation level from TipTap
        num_id = attrs.get("numId")  # Numbering definition ID
        format_change = attrs.get("formatChange")  # Tracked formatting change
        # styleNumbering is stored as an attr, not as text content
        # Word will regenerate numbering from the style definition

        if table_cell is not None:
            para = (
                table_cell.paragraphs[0]
                if table_cell.paragraphs
                else table_cell.add_paragraph()
            )
            if para.text == "" and len(table_cell.paragraphs) == 1:
                pass
            else:
                para = table_cell.add_paragraph()
        else:
            para = self._doc.add_paragraph()

        # Apply style if specified
        if style_name:
            try:
                para.style = style_name
            except KeyError:
                pass  # Style not found in document, skip

        # Apply raw pPr if available (preserves numId="0" overrides, etc.)
        if raw_pPr:
            self._restore_raw_paragraph_properties(para, raw_pPr, num_ilvl, format_change)
        elif num_id is not None and num_ilvl is not None:
            # No raw pPr but we have numbering - add numPr element
            self._add_numbering_to_paragraph(para, num_id, num_ilvl, format_change)

        content = node.get("content", [])

        for child_node in content:
            if child_node.get("type") == "text":
                self._add_text_with_marks(para, child_node)
            elif child_node.get("type") == "hardBreak":
                self._add_break(para, child_node)
            elif child_node.get("type") == "tab":
                self._add_tab(para, child_node)

    def _process_heading(self, node: dict, table_cell=None) -> None:
        """Process a heading node."""
        attrs = node.get("attrs", {})
        level = attrs.get("level", 1)
        style_name = attrs.get("styleName")
        raw_pPr = attrs.get("rawPPr")
        num_ilvl = attrs.get("numIlvl")  # Current indentation level from TipTap
        num_id = attrs.get("numId")  # Numbering definition ID
        format_change = attrs.get("formatChange")  # Tracked formatting change
        content = node.get("content", [])
        # styleNumbering is stored as an attr, not as text content
        # Word will regenerate numbering from the style definition

        if table_cell is not None:
            para = table_cell.add_paragraph()
            if style_name:
                try:
                    para.style = style_name
                except KeyError:
                    pass
            # Apply raw pPr if available
            if raw_pPr:
                self._restore_raw_paragraph_properties(para, raw_pPr, num_ilvl, format_change)
            elif num_id is not None and num_ilvl is not None:
                self._add_numbering_to_paragraph(para, num_id, num_ilvl, format_change)
            for child_node in content:
                if child_node.get("type") == "text":
                    run = para.add_run(child_node.get("text", ""))
                    run.bold = True
                    self._apply_basic_marks(run, child_node.get("marks", []))
                elif child_node.get("type") == "hardBreak":
                    self._add_break(para, child_node)
                elif child_node.get("type") == "tab":
                    self._add_tab(para, child_node)
        else:
            # If we have a custom style name, use it; otherwise use standard heading
            if style_name:
                para = self._doc.add_paragraph()
                try:
                    para.style = style_name
                except KeyError:
                    # Fall back to standard heading if style not found
                    para.style = f"Heading {level}"
            else:
                para = self._doc.add_heading(level=level)

            # Apply raw pPr if available (preserves numId="0" overrides, etc.)
            if raw_pPr:
                self._restore_raw_paragraph_properties(para, raw_pPr, num_ilvl, format_change)
            elif num_id is not None and num_ilvl is not None:
                self._add_numbering_to_paragraph(para, num_id, num_ilvl, format_change)

            for child_node in content:
                if child_node.get("type") == "text":
                    self._add_text_with_marks(para, child_node)
                elif child_node.get("type") == "hardBreak":
                    self._add_break(para, child_node)
                elif child_node.get("type") == "tab":
                    self._add_tab(para, child_node)

    def _process_table(self, node: dict, parent_cell=None) -> None:
        """Process a table node with support for merged cells (colspan/rowspan) and styling.

        Args:
            node: The TipTap table node
            parent_cell: Optional parent cell for nested tables
        """
        rows_data = node.get("content", [])
        if not rows_data:
            return

        # Calculate actual grid dimensions accounting for colspan
        num_rows = len(rows_data)
        num_cols = 0
        for row_node in rows_data:
            row_cols = 0
            for cell_node in row_node.get("content", []):
                colspan = cell_node.get("attrs", {}).get("colspan", 1)
                row_cols += colspan
            num_cols = max(num_cols, row_cols)

        if num_rows == 0 or num_cols == 0:
            return

        # Create table in the appropriate parent (document root or cell)
        if parent_cell is not None:
            table = parent_cell.add_table(rows=num_rows, cols=num_cols)
        else:
            table = self._doc.add_table(rows=num_rows, cols=num_cols)

        # Get table-level attributes
        table_attrs = node.get("attrs", {})

        # Check if we have raw XML to restore (lossless round-trip)
        raw_tblPr = table_attrs.get("rawTblPr")
        raw_tblGrid = table_attrs.get("rawTblGrid")

        if raw_tblPr or raw_tblGrid:
            # Restore raw table properties for exact style preservation
            self._restore_raw_table_properties(table, raw_tblPr, raw_tblGrid)
        else:
            # Fall back to individual style attributes
            style_name = table_attrs.get("styleName")
            if style_name:
                try:
                    table.style = style_name
                except KeyError:
                    pass  # Style not found, use default
            else:
                # Try to apply Table Grid style as default
                try:
                    table.style = "Table Grid"
                except KeyError:
                    try:
                        table.style = "TableGrid"
                    except KeyError:
                        pass  # Use default table style

            # Apply table alignment and column widths
            self._apply_table_style(table, table_attrs)

        # Track which cells are covered by merges (row_idx, col_idx) -> True
        covered_cells: set[tuple[int, int]] = set()

        # Track merges to apply after filling content
        merges_to_apply: list[tuple] = []

        for row_idx, row_node in enumerate(rows_data):
            cells_data = row_node.get("content", [])
            grid_col = 0  # Current position in the grid

            # Restore raw row properties if present
            row_attrs = row_node.get("attrs", {})
            raw_row_xml = row_attrs.get("rawXml")
            if raw_row_xml:
                self._restore_raw_row_properties(table.rows[row_idx], raw_row_xml)

            for cell_node in cells_data:
                # Skip grid positions that are covered by previous merges
                while (row_idx, grid_col) in covered_cells and grid_col < num_cols:
                    grid_col += 1

                if grid_col >= num_cols:
                    break

                attrs = cell_node.get("attrs", {})
                colspan = attrs.get("colspan", 1)
                rowspan = attrs.get("rowspan", 1)

                cell = table.rows[row_idx].cells[grid_col]
                if cell.paragraphs:
                    cell.paragraphs[0].clear()

                # Fill cell content
                cell_content = cell_node.get("content", [])
                self._fill_cell_content(cell, cell_content)

                # Apply cell styling - prefer raw XML for lossless round-trip
                raw_cell_xml = attrs.get("rawXml")
                if raw_cell_xml:
                    self._restore_raw_cell_properties(cell, raw_cell_xml)
                else:
                    self._apply_cell_style(cell, attrs)

                # Mark covered cells and prepare merge if needed
                if colspan > 1 or rowspan > 1:
                    # Mark all cells in the merge range as covered
                    for r in range(row_idx, row_idx + rowspan):
                        for c in range(grid_col, grid_col + colspan):
                            if r != row_idx or c != grid_col:  # Don't mark origin
                                covered_cells.add((r, c))

                    # Record merge to apply
                    end_row = min(row_idx + rowspan - 1, num_rows - 1)
                    end_col = min(grid_col + colspan - 1, num_cols - 1)
                    if end_row > row_idx or end_col > grid_col:
                        start_cell = table.rows[row_idx].cells[grid_col]
                        end_cell = table.rows[end_row].cells[end_col]
                        merges_to_apply.append((start_cell, end_cell))

                grid_col += colspan

        # Apply all merges after content is filled
        for start_cell, end_cell in merges_to_apply:
            try:
                start_cell.merge(end_cell)
            except Exception:
                pass  # Skip invalid merges gracefully

        # Clean up: remove extra tc elements that are artifacts of python-docx merge
        self._cleanup_merged_cells(table)

    def _fill_cell_content(self, cell, cell_content: list) -> None:
        """Fill a table cell with content."""
        for i, content_node in enumerate(cell_content):
            if i == 0 and cell.paragraphs:
                node_type = content_node.get("type")
                attrs = content_node.get("attrs", {})
                has_style_numbering = attrs.get("hasStyleNumbering", False)
                skip_first_text = has_style_numbering

                if node_type == "paragraph":
                    # Apply style if specified
                    style_name = attrs.get("styleName")
                    if style_name:
                        try:
                            cell.paragraphs[0].style = style_name
                        except KeyError:
                            pass
                    for child_node in content_node.get("content", []):
                        if child_node.get("type") == "text":
                            if skip_first_text:
                                skip_first_text = False
                                continue
                            self._add_text_with_marks(cell.paragraphs[0], child_node)
                        elif child_node.get("type") == "hardBreak":
                            self._add_break(cell.paragraphs[0], child_node)
                        elif child_node.get("type") == "tab":
                            self._add_tab(cell.paragraphs[0], child_node)
                elif node_type == "heading":
                    # Apply style if specified
                    style_name = attrs.get("styleName")
                    if style_name:
                        try:
                            cell.paragraphs[0].style = style_name
                        except KeyError:
                            pass
                    for child_node in content_node.get("content", []):
                        if child_node.get("type") == "text":
                            if skip_first_text:
                                skip_first_text = False
                                continue
                            run = cell.paragraphs[0].add_run(child_node.get("text", ""))
                            run.bold = True
                            self._apply_basic_marks(run, child_node.get("marks", []))
                        elif child_node.get("type") == "hardBreak":
                            self._add_break(cell.paragraphs[0], child_node)
                        elif child_node.get("type") == "tab":
                            self._add_tab(cell.paragraphs[0], child_node)
            else:
                self._process_node(content_node, table_cell=cell)

    def _process_section(self, node: dict) -> None:
        """Process a section node (just process its content)."""
        content = node.get("content", [])
        for child_node in content:
            self._process_node(child_node)

    def _add_text_with_marks(self, para, text_node: dict) -> None:
        """
        Add text to a paragraph, handling track changes and comments.

        Args:
            para: The python-docx Paragraph object
            text_node: The TipTap text node
        """
        text = text_node.get("text", "")
        if not text:
            return

        marks = text_node.get("marks", [])

        # Check for track change marks and raw style
        insertion_mark = None
        deletion_mark = None
        raw_style_mark = None
        comment_ids = []
        basic_marks = []

        for mark in marks:
            mark_type = mark.get("type")
            if mark_type == "insertion":
                insertion_mark = mark.get("attrs", {})
            elif mark_type == "deletion":
                deletion_mark = mark.get("attrs", {})
            elif mark_type == "rawStyle":
                raw_style_mark = mark.get("attrs", {})
            elif mark_type == "comment":
                comment_id = mark.get("attrs", {}).get("commentId")
                if comment_id:
                    comment_ids.append(comment_id)
            elif mark_type in ("bold", "italic"):
                basic_marks.append(mark)

        # Handle track changes
        if insertion_mark:
            self._add_insertion(para, text, insertion_mark, basic_marks, raw_style_mark)
        elif deletion_mark:
            self._add_deletion(para, text, deletion_mark, basic_marks, raw_style_mark)
        else:
            # Regular text - use raw rPr if available for full style preservation
            if raw_style_mark and raw_style_mark.get("rPr"):
                self._add_run_with_raw_rPr(para, text, raw_style_mark["rPr"])
            else:
                run = para.add_run(text)
                self._apply_basic_marks(run, basic_marks)

            # Track runs for comments
            for comment_id in comment_ids:
                if comment_id not in self._comment_runs_map:
                    self._comment_runs_map[comment_id] = []
                self._comment_runs_map[comment_id].append(run if not raw_style_mark else None)

    def _add_run_with_raw_rPr(self, para, text: str, raw_rPr: str) -> None:
        """
        Add a run to a paragraph using raw rPr XML for full style preservation.

        Args:
            para: The python-docx Paragraph object
            text: The text content
            raw_rPr: Base64-encoded w:rPr element
        """
        p_elem = para._p

        # Create run element
        r = OxmlElement("w:r")

        # Restore the full rPr from base64
        rPr = base64_to_element(raw_rPr)
        if rPr is not None:
            r.append(rPr)

        # Add text element
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)

        p_elem.append(r)

    def _add_insertion(
        self, para, text: str, attrs: dict, basic_marks: list, raw_style_mark: dict = None
    ) -> None:
        """
        Add text as an insertion (tracked change).

        Creates a w:ins element wrapping the run.
        """
        p_elem = para._p

        # Create w:ins element
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), self._next_revision_id())
        ins.set(qn("w:author"), attrs.get("author", "Unknown"))
        if attrs.get("date"):
            ins.set(qn("w:date"), attrs["date"])

        # Create run inside insertion
        r = OxmlElement("w:r")

        # Add run properties - prefer raw rPr for full style preservation
        if raw_style_mark and raw_style_mark.get("rPr"):
            rPr = base64_to_element(raw_style_mark["rPr"])
            if rPr is not None:
                r.append(rPr)
        elif basic_marks:
            rPr = OxmlElement("w:rPr")
            for mark in basic_marks:
                if mark.get("type") == "bold":
                    rPr.append(OxmlElement("w:b"))
                elif mark.get("type") == "italic":
                    rPr.append(OxmlElement("w:i"))
            r.append(rPr)

        # Add text element
        t = OxmlElement("w:t")
        t.text = text
        # Preserve spaces
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)

        ins.append(r)
        p_elem.append(ins)

    def _add_deletion(
        self, para, text: str, attrs: dict, basic_marks: list, raw_style_mark: dict = None
    ) -> None:
        """
        Add text as a deletion (tracked change).

        Creates a w:del element with w:delText inside.
        """
        p_elem = para._p

        # Create w:del element
        del_elem = OxmlElement("w:del")
        del_elem.set(qn("w:id"), self._next_revision_id())
        del_elem.set(qn("w:author"), attrs.get("author", "Unknown"))
        if attrs.get("date"):
            del_elem.set(qn("w:date"), attrs["date"])

        # Create run inside deletion
        r = OxmlElement("w:r")

        # Add run properties - prefer raw rPr for full style preservation
        if raw_style_mark and raw_style_mark.get("rPr"):
            rPr = base64_to_element(raw_style_mark["rPr"])
            if rPr is not None:
                r.append(rPr)
        elif basic_marks:
            rPr = OxmlElement("w:rPr")
            for mark in basic_marks:
                if mark.get("type") == "bold":
                    rPr.append(OxmlElement("w:b"))
                elif mark.get("type") == "italic":
                    rPr.append(OxmlElement("w:i"))
            r.append(rPr)

        # Add deleted text element (w:delText instead of w:t)
        del_text = OxmlElement("w:delText")
        del_text.text = text
        del_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(del_text)

        del_elem.append(r)
        p_elem.append(del_elem)

    def _apply_basic_marks(self, run, marks: list) -> None:
        """Apply basic formatting marks (bold, italic) to a run."""
        for mark in marks:
            mark_type = mark.get("type")
            if mark_type == "bold":
                run.bold = True
            elif mark_type == "italic":
                run.italic = True

    def _add_break(self, para, node: dict) -> None:
        """
        Add a break element to a paragraph.

        Args:
            para: The python-docx Paragraph object
            node: The TipTap hardBreak node
        """
        p_elem = para._p
        attrs = node.get("attrs", {})
        break_type = attrs.get("breakType")

        # Create run to contain the break
        r = OxmlElement("w:r")

        # Create break element
        br = OxmlElement("w:br")
        if break_type == "page":
            br.set(qn("w:type"), "page")
        elif break_type == "column":
            br.set(qn("w:type"), "column")
        # else: line break (no type attribute needed)

        r.append(br)
        p_elem.append(r)

    def _add_tab(self, para, node: dict) -> None:
        """
        Add a tab character to a paragraph, preserving styling (like dotted underlines).

        Args:
            para: The python-docx Paragraph object
            node: The TipTap tab node with optional marks
        """
        p_elem = para._p
        marks = node.get("marks", [])

        # Create run element
        r = OxmlElement("w:r")

        # Check for raw style mark (contains w:rPr with underline, etc.)
        raw_style_mark = None
        for mark in marks:
            if mark.get("type") == "rawStyle":
                raw_style_mark = mark.get("attrs", {})
                break

        # Restore rPr if available (preserves dotted underlines, etc.)
        if raw_style_mark and raw_style_mark.get("rPr"):
            rPr = base64_to_element(raw_style_mark["rPr"])
            if rPr is not None:
                r.append(rPr)

        # Create tab element
        tab = OxmlElement("w:tab")
        r.append(tab)

        p_elem.append(r)

    def _restore_raw_paragraph_properties(
        self, para, raw_xml: str, num_ilvl: int = None, format_change: dict = None
    ) -> None:
        """
        Restore raw pPr element from base64-encoded XML.

        This preserves ALL direct paragraph formatting for lossless round-tripping:
        - numPr (numbering, including numId="0" to disable numbering)
        - ind (indentation)
        - spacing (paragraph spacing)
        - pBdr (paragraph borders)
        - rPr (run properties default)
        - jc (justification)
        - pPrChange (tracked formatting changes)
        - And any other direct formatting elements

        Args:
            para: The python-docx Paragraph object
            raw_xml: Base64-encoded pPr element
            num_ilvl: If provided, update the ilvl in numPr to this value
                     (for when user changed indentation level in TipTap)
            format_change: If provided, add/update pPrChange element with tracked formatting info
        """
        if not raw_xml:
            return

        p = para._p
        new_pPr = base64_to_element(raw_xml)
        if new_pPr is None:
            return

        # If num_ilvl is provided, update the ilvl element in numPr
        if num_ilvl is not None:
            numPr = new_pPr.find(qn("w:numPr"))
            if numPr is not None:
                ilvl = numPr.find(qn("w:ilvl"))
                if ilvl is not None:
                    ilvl.set(qn("w:val"), str(num_ilvl))
                else:
                    # Create ilvl element if it doesn't exist
                    ilvl = OxmlElement("w:ilvl")
                    ilvl.set(qn("w:val"), str(num_ilvl))
                    # ilvl should be first child of numPr
                    numPr.insert(0, ilvl)

        # Handle format change tracking (pPrChange)
        if format_change:
            self._add_or_update_pPr_change(new_pPr, format_change)

        # Get existing pPr (python-docx creates one when we set para.style)
        existing_pPr = p.find(qn("w:pPr"))

        if existing_pPr is not None:
            # Merge: keep pStyle from existing, add all other elements from new_pPr
            # This is because python-docx sets pStyle when we assign para.style
            for child in list(new_pPr):
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag == "pStyle":
                    # Skip pStyle - we keep the one python-docx set
                    continue

                # Remove any existing element with the same tag
                existing_child = existing_pPr.find(child.tag)
                if existing_child is not None:
                    existing_pPr.remove(existing_child)

                # Add the new element - insert in proper position
                # OOXML has a specific element order requirement
                self._insert_pPr_element(existing_pPr, child)
        else:
            # No existing pPr, just add the new one
            p.insert(0, new_pPr)

    def _add_numbering_to_paragraph(
        self, para, num_id: str, num_ilvl: int, format_change: dict = None
    ) -> None:
        """
        Add numbering properties to a paragraph that doesn't have rawPPr.

        This is used for paragraphs created via Enter key in TipTap that
        inherit numbering but don't have preserved raw XML.

        Args:
            para: The python-docx Paragraph object
            num_id: The Word numbering definition ID
            num_ilvl: The indentation level (0-8)
            format_change: If provided, add pPrChange element with tracked formatting info
        """
        p = para._p
        pPr = p.get_or_add_pPr()

        # Create numPr element
        numPr = OxmlElement("w:numPr")

        # Add ilvl
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(num_ilvl))
        numPr.append(ilvl)

        # Add numId
        numId_elem = OxmlElement("w:numId")
        numId_elem.set(qn("w:val"), str(num_id))
        numPr.append(numId_elem)

        # Insert numPr in the correct position
        self._insert_pPr_element(pPr, numPr)

        # Handle format change tracking (pPrChange)
        if format_change:
            self._add_or_update_pPr_change(pPr, format_change)

    def _add_or_update_pPr_change(self, pPr, format_change: dict) -> None:
        """
        Add or update a pPrChange element to track formatting changes.

        The pPrChange element stores the OLD formatting state, allowing Word
        to show what the paragraph was changed FROM. The current formatting
        is stored in the parent pPr.

        Args:
            pPr: The pPr element to add pPrChange to
            format_change: Dict with id, author, date, oldStyle, oldNumIlvl
        """
        # Remove existing pPrChange if present (we'll recreate it)
        existing_pPr_change = pPr.find(qn("w:pPrChange"))
        if existing_pPr_change is not None:
            pPr.remove(existing_pPr_change)

        # Create new pPrChange element
        pPr_change = OxmlElement("w:pPrChange")

        # Set attributes
        if format_change.get("id"):
            pPr_change.set(qn("w:id"), str(format_change["id"]))
        else:
            pPr_change.set(qn("w:id"), self._next_revision_id())

        if format_change.get("author"):
            pPr_change.set(qn("w:author"), format_change["author"])
        else:
            pPr_change.set(qn("w:author"), "Unknown")

        if format_change.get("date"):
            pPr_change.set(qn("w:date"), format_change["date"])

        # Create the old pPr element inside pPrChange
        old_pPr = OxmlElement("w:pPr")

        # Add old style if present
        old_style = format_change.get("oldStyle")
        if old_style:
            old_pStyle = OxmlElement("w:pStyle")
            old_pStyle.set(qn("w:val"), old_style)
            old_pPr.append(old_pStyle)

        # Add old numbering level if present
        old_num_ilvl = format_change.get("oldNumIlvl")
        if old_num_ilvl is not None:
            # We need to create numPr with the old ilvl
            # But we need the numId from the current pPr
            current_numPr = pPr.find(qn("w:numPr"))
            if current_numPr is not None:
                current_numId = current_numPr.find(qn("w:numId"))
                if current_numId is not None:
                    old_numPr = OxmlElement("w:numPr")
                    old_ilvl = OxmlElement("w:ilvl")
                    old_ilvl.set(qn("w:val"), str(old_num_ilvl))
                    old_numPr.append(old_ilvl)
                    # Copy the numId
                    old_numId = OxmlElement("w:numId")
                    old_numId.set(qn("w:val"), current_numId.get(qn("w:val")))
                    old_numPr.append(old_numId)
                    old_pPr.append(old_numPr)

        pPr_change.append(old_pPr)

        # pPrChange must be the last element in pPr per OOXML schema
        pPr.append(pPr_change)

    def _insert_pPr_element(self, pPr, element) -> None:
        """
        Insert an element into pPr in the correct OOXML order.

        OOXML requires elements in a specific sequence. This ensures
        elements are inserted at the right position.
        """
        # OOXML pPr element order (simplified - covers most common elements)
        # See: http://www.datypic.com/sc/ooxml/e-w_pPr-1.html
        element_order = [
            "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
            "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd",
            "tabs", "suppressAutoHyphens", "kinsoku", "wordWrap",
            "overflowPunct", "topLinePunct", "autoSpaceDE", "autoSpaceDN",
            "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind",
            "contextualSpacing", "mirrorIndents", "suppressOverlap", "jc",
            "textDirection", "textAlignment", "textboxTightWrap",
            "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr", "pPrChange"
        ]

        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        try:
            target_idx = element_order.index(tag)
        except ValueError:
            # Unknown element, append at end
            pPr.append(element)
            return

        # Find the right position to insert
        for i, existing in enumerate(pPr):
            existing_tag = existing.tag.split("}")[-1] if "}" in existing.tag else existing.tag
            try:
                existing_idx = element_order.index(existing_tag)
                if existing_idx > target_idx:
                    pPr.insert(i, element)
                    return
            except ValueError:
                continue

        # If we get here, append at end
        pPr.append(element)

    def _restore_raw_cell_properties(self, cell, raw_xml: str) -> None:
        """Restore raw tcPr element from base64-encoded XML."""
        if not raw_xml:
            return

        tc = cell._tc
        new_tcPr = base64_to_element(raw_xml)
        if new_tcPr is None:
            return

        # Remove existing tcPr if present
        existing_tcPr = tc.find(qn("w:tcPr"))
        if existing_tcPr is not None:
            tc.remove(existing_tcPr)

        # Insert new tcPr at the beginning (it should be first child)
        tc.insert(0, new_tcPr)

    def _restore_raw_row_properties(self, row, raw_xml: str) -> None:
        """Restore raw trPr element from base64-encoded XML."""
        if not raw_xml:
            return

        tr = row._tr
        new_trPr = base64_to_element(raw_xml)
        if new_trPr is None:
            return

        # Remove existing trPr if present
        existing_trPr = tr.find(qn("w:trPr"))
        if existing_trPr is not None:
            tr.remove(existing_trPr)

        # Insert new trPr at the beginning (it should be first child)
        tr.insert(0, new_trPr)

    def _restore_raw_table_properties(
        self, table, raw_tblPr: str, raw_tblGrid: str
    ) -> None:
        """Restore raw tblPr and tblGrid elements from base64-encoded XML."""
        tbl = table._tbl

        # Restore tblPr
        if raw_tblPr:
            new_tblPr = base64_to_element(raw_tblPr)
            if new_tblPr is not None:
                existing_tblPr = tbl.find(qn("w:tblPr"))
                if existing_tblPr is not None:
                    tbl.remove(existing_tblPr)
                # tblPr should be first child
                tbl.insert(0, new_tblPr)

        # Restore tblGrid
        if raw_tblGrid:
            new_tblGrid = base64_to_element(raw_tblGrid)
            if new_tblGrid is not None:
                existing_tblGrid = tbl.find(qn("w:tblGrid"))
                if existing_tblGrid is not None:
                    tbl.remove(existing_tblGrid)
                # tblGrid should be after tblPr
                # Find the insert position (after tblPr if present)
                tblPr = tbl.find(qn("w:tblPr"))
                if tblPr is not None:
                    tblPr_idx = list(tbl).index(tblPr)
                    tbl.insert(tblPr_idx + 1, new_tblGrid)
                else:
                    tbl.insert(0, new_tblGrid)

    def _cleanup_merged_cells(self, table) -> None:
        """
        Remove extra tc elements that python-docx leaves behind after merging.

        When python-docx merges cells horizontally, it sets gridSpan on the first cell
        but leaves the subsequent cells. We need to remove cells that push the total
        grid span beyond the table's column count.

        For vertical merges (vMerge), we keep the continuation cells as they're needed
        for the vertical merge structure.
        """
        tbl = table._tbl

        # Get the number of grid columns from tblGrid
        tblGrid = tbl.tblGrid
        if tblGrid is None:
            return
        num_grid_cols = len(tblGrid.findall(qn("w:gridCol")))
        if num_grid_cols == 0:
            return

        for row in table.rows:
            tr = row._tr
            tc_list = list(tr.tc_lst)

            # Track grid position and find cells to remove
            grid_pos = 0
            cells_to_remove = []

            for tc in tc_list:
                tcPr = tc.find(qn("w:tcPr"))

                # Get gridSpan (horizontal span)
                gridSpan = 1
                if tcPr is not None:
                    gs = tcPr.find(qn("w:gridSpan"))
                    if gs is not None:
                        try:
                            gridSpan = int(gs.get(qn("w:val")))
                        except (ValueError, TypeError):
                            gridSpan = 1

                # If this cell would start beyond the grid, it's an artifact
                if grid_pos >= num_grid_cols:
                    cells_to_remove.append(tc)
                else:
                    grid_pos += gridSpan

            # Remove artifact cells from the row
            for tc in cells_to_remove:
                tr.remove(tc)

    def _apply_cell_style(self, cell, attrs: dict) -> None:
        """Apply styling to a table cell from TipTap attributes."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Background color
        bg_color = attrs.get("backgroundColor")
        if bg_color:
            # Remove # prefix if present
            color = bg_color.lstrip("#")
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), color)
            tcPr.append(shd)

        # Vertical alignment
        v_align = attrs.get("verticalAlign")
        if v_align:
            vAlign = OxmlElement("w:vAlign")
            vAlign.set(qn("w:val"), v_align)
            tcPr.append(vAlign)

        # Cell width
        colwidth = attrs.get("colwidth")
        if colwidth and len(colwidth) > 0:
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(colwidth[0]))
            tcW.set(qn("w:type"), "dxa")
            # Remove existing tcW if present
            existing_tcW = tcPr.find(qn("w:tcW"))
            if existing_tcW is not None:
                tcPr.remove(existing_tcW)
            tcPr.insert(0, tcW)  # tcW should be first

        # Borders
        borders = attrs.get("borders")
        if borders:
            tcBorders = OxmlElement("w:tcBorders")
            for side in ["top", "bottom", "left", "right"]:
                border_data = borders.get(side)
                if border_data:
                    border_elem = OxmlElement(f"w:{side}")
                    if border_data.get("style"):
                        border_elem.set(qn("w:val"), border_data["style"])
                    if border_data.get("width"):
                        border_elem.set(qn("w:sz"), str(border_data["width"]))
                    if border_data.get("color"):
                        border_elem.set(qn("w:color"), border_data["color"])
                    tcBorders.append(border_elem)
            tcPr.append(tcBorders)

        # Text alignment (applied to first paragraph)
        text_align = attrs.get("textAlign")
        if text_align and cell.paragraphs:
            para = cell.paragraphs[0]
            pPr = para._p.get_or_add_pPr()
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), text_align)
            pPr.append(jc)

    def _apply_table_style(self, table, table_attrs: dict) -> None:
        """Apply styling to a table from TipTap attributes."""
        tbl = table._tbl

        # Table alignment
        alignment = table_attrs.get("alignment")
        if alignment:
            from docx.enum.table import WD_TABLE_ALIGNMENT

            alignment_map = {
                "left": WD_TABLE_ALIGNMENT.LEFT,
                "center": WD_TABLE_ALIGNMENT.CENTER,
                "right": WD_TABLE_ALIGNMENT.RIGHT,
            }
            if alignment in alignment_map:
                table.alignment = alignment_map[alignment]

        # Column widths
        colwidths = table_attrs.get("colwidths")
        if colwidths:
            # Set column widths via tblGrid
            tblGrid = tbl.tblGrid
            if tblGrid is not None:
                gridCols = tblGrid.findall(qn("w:gridCol"))
                for i, width in enumerate(colwidths):
                    if i < len(gridCols):
                        gridCols[i].set(qn("w:w"), str(width))

    def _add_comments_to_document(self) -> None:
        """
        Add comments to the document using python-docx's add_comment API.
        """
        for comment_id, runs in self._comment_runs_map.items():
            if comment_id not in self._comments_dict:
                continue

            comment_data = self._comments_dict[comment_id]
            if not runs:
                continue

            # Get first and last run for the comment range
            first_run = runs[0]
            last_run = runs[-1] if len(runs) > 1 else runs[0]

            try:
                # Use python-docx native add_comment API
                self._doc.add_comment(
                    first_run,
                    text=comment_data.get("text", ""),
                    author=comment_data.get("author", "Unknown"),
                    initials=comment_data.get("initials", ""),
                )

                # If comment spans multiple runs, mark the range
                if len(runs) > 1:
                    first_run.mark_comment_range(last_run, self._doc.comments[-1].comment_id)
            except Exception as e:
                # If comment creation fails, continue without it
                print(f"Warning: Could not add comment {comment_id}: {e}")


def create_docx_from_tiptap(
    tiptap_json: dict,
    comments: Optional[list[dict]] = None,
    template_bytes: Optional[bytes] = None,
) -> BytesIO:
    """
    Convert TipTap JSON document to a Word document.

    This is the public API function that maintains backward compatibility.
    It delegates to DocxExporter internally.

    Args:
        tiptap_json: TipTap document JSON with structure:
            {
                "type": "doc",
                "content": [...]
            }
        comments: Optional list of comment dictionaries with id, author, text, date
        template_bytes: Optional bytes of a .docx file to use as template.
            If provided, the template's styles will be preserved but content
            will be replaced with the TipTap content.

    Returns:
        BytesIO buffer containing the .docx file
    """
    exporter = DocxExporter(comments=comments, template_bytes=template_bytes)
    return exporter.export(tiptap_json)
