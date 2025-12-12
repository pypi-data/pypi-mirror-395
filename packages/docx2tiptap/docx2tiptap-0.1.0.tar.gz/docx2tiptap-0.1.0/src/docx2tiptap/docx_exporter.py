"""
DOCX Exporter - Converts TipTap JSON back to Word document format.

This module handles the reverse conversion from TipTap editor JSON
back to a Word document (.docx) using python-docx.

Supports:
- Basic text formatting (bold, italic)
- Headings (H1-H6)
- Tables
- Track changes (insertions/deletions) via OOXML
- Comments via python-docx native API
"""

from io import BytesIO
from typing import Any, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Global counter for revision IDs
_revision_id_counter = 0


def _next_revision_id() -> str:
    """Generate next revision ID."""
    global _revision_id_counter
    _revision_id_counter += 1
    return str(_revision_id_counter)


def _reset_revision_counter():
    """Reset revision counter (call at start of each export)."""
    global _revision_id_counter
    _revision_id_counter = 0


def _clear_document_content(doc: Document) -> None:
    """
    Clear all content from a document while preserving styles.

    This removes paragraphs and tables from the document body but keeps
    styles, headers, footers, and other document properties intact.
    """
    # Get the body element
    body = doc.element.body

    # Remove all paragraph and table elements from the body
    # We iterate in reverse to avoid issues with modifying during iteration
    for child in list(body):
        # Keep section properties (w:sectPr) as they contain page layout
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def create_docx_from_tiptap(
    tiptap_json: dict,
    comments: Optional[list[dict]] = None,
    template_bytes: Optional[bytes] = None,
) -> BytesIO:
    """
    Convert TipTap JSON document to a Word document.

    Args:
        tiptap_json: TipTap document JSON with structure:
            {
                "type": "doc",
                "content": [...]
            }
        comments: Optional list of comment dictionaries with id, author, text, date
        template_bytes: Optional bytes of a .docx file to use as template.
            If provided, the template's styles will be preserved but content
            will be replaced with the TipTap content.

    Returns:
        BytesIO buffer containing the .docx file
    """
    _reset_revision_counter()

    # Create document from template or blank
    if template_bytes:
        template_buffer = BytesIO(template_bytes)
        doc = Document(template_buffer)
        # Clear existing content but keep styles
        _clear_document_content(doc)
    else:
        doc = Document()
    comments_dict = {c["id"]: c for c in (comments or [])}

    # Track which runs correspond to which comment IDs for later linking
    comment_runs_map: dict[str, list] = {}

    content = tiptap_json.get("content", [])
    for node in content:
        process_node(doc, node, comments_dict, comment_runs_map)

    # Add comments after processing all content
    _add_comments_to_document(doc, comments_dict, comment_runs_map)

    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def process_node(
    doc: Document,
    node: dict,
    comments_dict: dict,
    comment_runs_map: dict,
    table_cell=None,
) -> None:
    """
    Process a TipTap node and add it to the document.

    Args:
        doc: The python-docx Document object
        node: A TipTap node dictionary
        comments_dict: Dictionary of comment ID to comment data
        comment_runs_map: Map to track runs for each comment ID
        table_cell: Optional table cell to add content to (for nested content)
    """
    node_type = node.get("type")

    if node_type == "paragraph":
        process_paragraph(
            doc, node, comments_dict, comment_runs_map, table_cell
        )
    elif node_type == "heading":
        process_heading(doc, node, comments_dict, comment_runs_map, table_cell)
    elif node_type == "table":
        process_table(doc, node, comments_dict, comment_runs_map)
    elif node_type == "section":
        process_section(doc, node, comments_dict, comment_runs_map)


def process_paragraph(
    doc: Document,
    node: dict,
    comments_dict: dict,
    comment_runs_map: dict,
    table_cell=None,
) -> None:
    """Process a paragraph node."""
    if table_cell is not None:
        para = (
            table_cell.paragraphs[0]
            if table_cell.paragraphs
            else table_cell.add_paragraph()
        )
        if para.text == "" and len(table_cell.paragraphs) == 1:
            pass
        else:
            para = table_cell.add_paragraph()
    else:
        para = doc.add_paragraph()

    content = node.get("content", [])
    for text_node in content:
        if text_node.get("type") == "text":
            add_text_with_marks(
                para, text_node, comments_dict, comment_runs_map
            )


def process_heading(
    doc: Document,
    node: dict,
    comments_dict: dict,
    comment_runs_map: dict,
    table_cell=None,
) -> None:
    """Process a heading node."""
    level = node.get("attrs", {}).get("level", 1)
    content = node.get("content", [])

    if table_cell is not None:
        para = table_cell.add_paragraph()
        for text_node in content:
            if text_node.get("type") == "text":
                run = para.add_run(text_node.get("text", ""))
                run.bold = True
                _apply_basic_marks(run, text_node.get("marks", []))
    else:
        para = doc.add_heading(level=level)
        for text_node in content:
            if text_node.get("type") == "text":
                add_text_with_marks(
                    para, text_node, comments_dict, comment_runs_map
                )


def process_table(
    doc: Document, node: dict, comments_dict: dict, comment_runs_map: dict
) -> None:
    """Process a table node with support for merged cells (colspan/rowspan)."""
    rows_data = node.get("content", [])
    if not rows_data:
        return

    # Calculate actual grid dimensions accounting for colspan
    num_rows = len(rows_data)
    num_cols = 0
    for row_node in rows_data:
        row_cols = 0
        for cell_node in row_node.get("content", []):
            colspan = cell_node.get("attrs", {}).get("colspan", 1)
            row_cols += colspan
        num_cols = max(num_cols, row_cols)

    if num_rows == 0 or num_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)
    # Try to apply Table Grid style, fall back gracefully if not available
    try:
        table.style = "Table Grid"
    except KeyError:
        try:
            table.style = "TableGrid"
        except KeyError:
            pass  # Use default table style

    # Track which cells are covered by merges (row_idx, col_idx) -> True
    covered_cells: set[tuple[int, int]] = set()

    # Track merges to apply after filling content: list of (start_cell, end_cell)
    merges_to_apply: list[tuple] = []

    for row_idx, row_node in enumerate(rows_data):
        cells_data = row_node.get("content", [])
        grid_col = 0  # Current position in the grid

        for cell_node in cells_data:
            # Skip grid positions that are covered by previous merges
            while (row_idx, grid_col) in covered_cells and grid_col < num_cols:
                grid_col += 1

            if grid_col >= num_cols:
                break

            attrs = cell_node.get("attrs", {})
            colspan = attrs.get("colspan", 1)
            rowspan = attrs.get("rowspan", 1)

            cell = table.rows[row_idx].cells[grid_col]
            if cell.paragraphs:
                cell.paragraphs[0].clear()

            # Fill cell content
            cell_content = cell_node.get("content", [])
            _fill_cell_content(
                doc, cell, cell_content, comments_dict, comment_runs_map
            )

            # Mark covered cells and prepare merge if needed
            if colspan > 1 or rowspan > 1:
                # Mark all cells in the merge range as covered
                for r in range(row_idx, row_idx + rowspan):
                    for c in range(grid_col, grid_col + colspan):
                        if r != row_idx or c != grid_col:  # Don't mark origin
                            covered_cells.add((r, c))

                # Record merge to apply
                end_row = min(row_idx + rowspan - 1, num_rows - 1)
                end_col = min(grid_col + colspan - 1, num_cols - 1)
                if end_row > row_idx or end_col > grid_col:
                    start_cell = table.rows[row_idx].cells[grid_col]
                    end_cell = table.rows[end_row].cells[end_col]
                    merges_to_apply.append((start_cell, end_cell))

            grid_col += colspan

    # Apply all merges after content is filled
    for start_cell, end_cell in merges_to_apply:
        try:
            start_cell.merge(end_cell)
        except Exception:
            pass  # Skip invalid merges gracefully


def _fill_cell_content(
    doc: Document,
    cell,
    cell_content: list,
    comments_dict: dict,
    comment_runs_map: dict,
) -> None:
    """Fill a table cell with content."""
    for i, content_node in enumerate(cell_content):
        if i == 0 and cell.paragraphs:
            if content_node.get("type") == "paragraph":
                for text_node in content_node.get("content", []):
                    if text_node.get("type") == "text":
                        add_text_with_marks(
                            cell.paragraphs[0],
                            text_node,
                            comments_dict,
                            comment_runs_map,
                        )
            elif content_node.get("type") == "heading":
                for text_node in content_node.get("content", []):
                    if text_node.get("type") == "text":
                        run = cell.paragraphs[0].add_run(
                            text_node.get("text", "")
                        )
                        run.bold = True
                        _apply_basic_marks(run, text_node.get("marks", []))
        else:
            process_node(
                doc,
                content_node,
                comments_dict,
                comment_runs_map,
                table_cell=cell,
            )


def process_section(
    doc: Document, node: dict, comments_dict: dict, comment_runs_map: dict
) -> None:
    """Process a section node (just process its content)."""
    content = node.get("content", [])
    for child_node in content:
        process_node(doc, child_node, comments_dict, comment_runs_map)


def add_text_with_marks(
    para, text_node: dict, comments_dict: dict, comment_runs_map: dict
) -> None:
    """
    Add text to a paragraph, handling track changes and comments.

    Args:
        para: The python-docx Paragraph object
        text_node: The TipTap text node
        comments_dict: Dictionary of comment data
        comment_runs_map: Map to track runs for comment linking
    """
    text = text_node.get("text", "")
    if not text:
        return

    marks = text_node.get("marks", [])

    # Check for track change marks
    insertion_mark = None
    deletion_mark = None
    comment_ids = []
    basic_marks = []

    for mark in marks:
        mark_type = mark.get("type")
        if mark_type == "insertion":
            insertion_mark = mark.get("attrs", {})
        elif mark_type == "deletion":
            deletion_mark = mark.get("attrs", {})
        elif mark_type == "comment":
            comment_id = mark.get("attrs", {}).get("commentId")
            if comment_id:
                comment_ids.append(comment_id)
        elif mark_type in ("bold", "italic"):
            basic_marks.append(mark)

    # Handle track changes
    if insertion_mark:
        _add_insertion(para, text, insertion_mark, basic_marks)
    elif deletion_mark:
        _add_deletion(para, text, deletion_mark, basic_marks)
    else:
        # Regular text
        run = para.add_run(text)
        _apply_basic_marks(run, basic_marks)

        # Track runs for comments
        for comment_id in comment_ids:
            if comment_id not in comment_runs_map:
                comment_runs_map[comment_id] = []
            comment_runs_map[comment_id].append(run)


def _add_insertion(para, text: str, attrs: dict, basic_marks: list) -> None:
    """
    Add text as an insertion (tracked change).

    Creates a w:ins element wrapping the run.
    """
    p_elem = para._p

    # Create w:ins element
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), _next_revision_id())
    ins.set(qn("w:author"), attrs.get("author", "Unknown"))
    if attrs.get("date"):
        ins.set(qn("w:date"), attrs["date"])

    # Create run inside insertion
    r = OxmlElement("w:r")

    # Add run properties for formatting
    if basic_marks:
        rPr = OxmlElement("w:rPr")
        for mark in basic_marks:
            if mark.get("type") == "bold":
                rPr.append(OxmlElement("w:b"))
            elif mark.get("type") == "italic":
                rPr.append(OxmlElement("w:i"))
        r.append(rPr)

    # Add text element
    t = OxmlElement("w:t")
    t.text = text
    # Preserve spaces
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)

    ins.append(r)
    p_elem.append(ins)


def _add_deletion(para, text: str, attrs: dict, basic_marks: list) -> None:
    """
    Add text as a deletion (tracked change).

    Creates a w:del element with w:delText inside.
    """
    p_elem = para._p

    # Create w:del element
    del_elem = OxmlElement("w:del")
    del_elem.set(qn("w:id"), _next_revision_id())
    del_elem.set(qn("w:author"), attrs.get("author", "Unknown"))
    if attrs.get("date"):
        del_elem.set(qn("w:date"), attrs["date"])

    # Create run inside deletion
    r = OxmlElement("w:r")

    # Add run properties for formatting
    if basic_marks:
        rPr = OxmlElement("w:rPr")
        for mark in basic_marks:
            if mark.get("type") == "bold":
                rPr.append(OxmlElement("w:b"))
            elif mark.get("type") == "italic":
                rPr.append(OxmlElement("w:i"))
        r.append(rPr)

    # Add deleted text element (w:delText instead of w:t)
    del_text = OxmlElement("w:delText")
    del_text.text = text
    del_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(del_text)

    del_elem.append(r)
    p_elem.append(del_elem)


def _apply_basic_marks(run, marks: list) -> None:
    """Apply basic formatting marks (bold, italic) to a run."""
    for mark in marks:
        mark_type = mark.get("type")
        if mark_type == "bold":
            run.bold = True
        elif mark_type == "italic":
            run.italic = True


def _add_comments_to_document(
    doc: Document, comments_dict: dict, comment_runs_map: dict
) -> None:
    """
    Add comments to the document using python-docx's add_comment API.

    Args:
        doc: The python-docx Document object
        comments_dict: Dictionary of comment ID to comment data
        comment_runs_map: Map of comment ID to list of runs
    """
    for comment_id, runs in comment_runs_map.items():
        if comment_id not in comments_dict:
            continue

        comment_data = comments_dict[comment_id]
        if not runs:
            continue

        # Get first and last run for the comment range
        first_run = runs[0]
        last_run = runs[-1] if len(runs) > 1 else runs[0]

        try:
            # Use python-docx native add_comment API
            doc.add_comment(
                first_run,
                text=comment_data.get("text", ""),
                author=comment_data.get("author", "Unknown"),
                initials=comment_data.get("initials", ""),
            )

            # If comment spans multiple runs, mark the range
            if len(runs) > 1:
                first_run.mark_comment_range(
                    last_run, doc.comments[-1].comment_id
                )
        except Exception as e:
            # If comment creation fails, continue without it
            print(f"Warning: Could not add comment {comment_id}: {e}")
