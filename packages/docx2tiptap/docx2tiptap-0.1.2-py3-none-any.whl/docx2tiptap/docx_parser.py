"""
DOCX Parser - Extracts structured content from Word documents.

Handles:
- Paragraphs with text formatting (bold, italic)
- Numbered/bulleted lists with computed numbering
- Tables with rich cell content
- Headings
- Track changes (insertions/deletions)
- Comments
"""

from io import BytesIO
from typing import Optional

from docx import Document
from docx.oxml.ns import qn

from .comments_parser import (
    extract_comments_from_docx,
    get_text_with_comments,
)
from .models import (
    BorderStyle,
    CellBorders,
    CellStyle,
    Paragraph,
    ParagraphFormatChange,
    Table,
    TableCell,
    TableRow,
    TableStyle,
    TextRun,
)
from .numbering import NumberingTracker
from .revisions_parser import get_text_with_revisions, merge_adjacent_segments
from .serializers import elements_to_dict
from .utils import element_to_base64


def parse_paragraph(
    para,
    numbering_tracker: Optional[NumberingTracker] = None,
    para_index: int = 0,
) -> Paragraph:
    """Parse a python-docx paragraph into our intermediate format."""

    # Get text with revisions (track changes) and comments
    revision_segments = get_text_with_revisions(para._element, para_index)
    revision_segments = merge_adjacent_segments(revision_segments)

    comment_segments = get_text_with_comments(para._element, para_index)

    # Build a map of text positions to comment IDs
    comment_map = {}  # text -> [comment_ids]
    for seg in comment_segments:
        if seg["comments"]:
            comment_map[seg["text"]] = seg["comments"]

    # Convert segments to TextRuns
    runs = []
    for seg in revision_segments:
        if seg.get("break"):
            # Handle break segments (page breaks, line breaks, etc.)
            runs.append(
                TextRun(
                    text="",
                    is_break=True,
                    break_type=seg.get("break_type"),
                    revision=seg.get("revision"),
                )
            )
        elif seg.get("text"):
            # Try to find matching comment info
            comment_ids = comment_map.get(seg["text"], [])

            runs.append(
                TextRun(
                    text=seg["text"],
                    bold=seg.get("bold", False),
                    italic=seg.get("italic", False),
                    revision=seg.get("revision"),
                    comment_ids=comment_ids,
                    raw_rPr=seg.get("raw_rPr"),
                )
            )

    # If no revision-aware segments found, fall back to simple parsing
    if not runs:
        for run in para.runs:
            if run.text:
                runs.append(
                    TextRun(
                        text=run.text,
                        bold=run.bold or False,
                        italic=run.italic or False,
                    )
                )

    # Detect heading level
    level = 0
    style_name = para.style.name if para.style else ""
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.replace("Heading ", ""))
        except ValueError:
            pass

    # Get numbering info - check direct numPr first, then fall back to style
    numbering = None
    raw_pPr = None
    num_id = None
    ilvl = None
    format_change = None

    # Check for direct paragraph formatting (pPr)
    if para._element.pPr is not None:
        pPr = para._element.pPr

        # Check for direct numPr on the paragraph
        num_pr = pPr.find(qn("w:numPr"))
        if num_pr is not None:
            ilvl_elem = num_pr.find(qn("w:ilvl"))
            num_id_elem = num_pr.find(qn("w:numId"))
            if ilvl_elem is not None and num_id_elem is not None:
                ilvl = int(ilvl_elem.get(qn("w:val")))
                num_id = num_id_elem.get(qn("w:val"))

        # Check for pPrChange (tracked formatting change)
        pPr_change = pPr.find(qn("w:pPrChange"))
        if pPr_change is not None:
            format_change = _parse_pPr_change(pPr_change)

        # Preserve pPr if there's ANY direct formatting beyond just pStyle
        # This includes: numPr, ind (indent), spacing, pBdr (borders), rPr, etc.
        # We need to preserve these for lossless round-tripping
        has_direct_formatting = False
        for child in pPr:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag != "pStyle":  # pStyle is just a reference, not direct formatting
                has_direct_formatting = True
                break

        if has_direct_formatting:
            raw_pPr = element_to_base64(pPr)

    # If no direct numPr, check if the style defines numbering
    if numbering_tracker and num_id is None and style_name:
        style_numbering = numbering_tracker.get_numbering_from_style(
            style_name
        )
        if style_numbering:
            num_id, ilvl = style_numbering

    # Compute the actual number if we have numbering info
    if numbering_tracker and num_id and num_id != "0":
        numbering = numbering_tracker.get_number(num_id, ilvl)

    return Paragraph(
        runs=runs,
        style=style_name,
        numbering=numbering,
        level=level,
        raw_pPr=raw_pPr,
        num_id=num_id if num_id and num_id != "0" else None,
        num_ilvl=ilvl if ilvl is not None else 0,
        format_change=format_change,
    )


def _parse_pPr_change(pPr_change) -> Optional[ParagraphFormatChange]:
    """Parse a pPrChange element into a ParagraphFormatChange.

    The pPrChange element tracks formatting changes to a paragraph,
    storing the old formatting inside and the new formatting outside.

    Example XML:
        <w:pPrChange w:id="27" w:author="John" w:date="2025-12-07T15:19:00Z">
            <w:pPr>
                <w:pStyle w:val="SH3Legal"/>
            </w:pPr>
        </w:pPrChange>
    """
    if pPr_change is None:
        return None

    # Extract attributes from pPrChange element
    change_id = pPr_change.get(qn("w:id"))
    author = pPr_change.get(qn("w:author"))
    date = pPr_change.get(qn("w:date"))

    # Get the old formatting from the nested pPr element
    old_pPr = pPr_change.find(qn("w:pPr"))
    old_style = None
    old_num_ilvl = None

    if old_pPr is not None:
        # Extract old style name
        old_pStyle = old_pPr.find(qn("w:pStyle"))
        if old_pStyle is not None:
            old_style = old_pStyle.get(qn("w:val"))

        # Extract old numbering level if present
        old_numPr = old_pPr.find(qn("w:numPr"))
        if old_numPr is not None:
            old_ilvl_elem = old_numPr.find(qn("w:ilvl"))
            if old_ilvl_elem is not None:
                old_num_ilvl = int(old_ilvl_elem.get(qn("w:val")))

    return ParagraphFormatChange(
        id=change_id,
        author=author,
        date=date,
        old_style=old_style,
        old_num_ilvl=old_num_ilvl,
    )


def _parse_border_style(border_elem) -> Optional[BorderStyle]:
    """Parse a border element into a BorderStyle."""
    if border_elem is None:
        return None

    style = border_elem.get(qn("w:val"))
    if style == "nil":
        return None

    width_str = border_elem.get(qn("w:sz"))
    width = int(width_str) if width_str else None
    color = border_elem.get(qn("w:color"))

    return BorderStyle(style=style, width=width, color=color)


def _parse_cell_borders(tcPr) -> Optional[CellBorders]:
    """Parse cell borders from tcPr element."""
    if tcPr is None:
        return None

    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        return None

    top = _parse_border_style(tcBorders.find(qn("w:top")))
    bottom = _parse_border_style(tcBorders.find(qn("w:bottom")))
    left = _parse_border_style(tcBorders.find(qn("w:left")))
    right = _parse_border_style(tcBorders.find(qn("w:right")))

    if not any([top, bottom, left, right]):
        return None

    return CellBorders(top=top, bottom=bottom, left=left, right=right)


def _parse_cell_style(tc, cell) -> Optional[CellStyle]:
    """Parse cell styling from tc element."""
    tcPr = tc.tcPr
    if tcPr is None:
        return None

    style = CellStyle()
    has_style = False

    # Width
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is not None:
        w = tcW.get(qn("w:w"))
        if w:
            style.width = int(w)
            has_style = True

    # Background color (shading)
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        fill = shd.get(qn("w:fill"))
        if fill and fill.lower() not in ("auto", "ffffff"):
            style.background_color = fill
            has_style = True

    # Vertical alignment
    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is not None:
        style.vertical_align = vAlign.get(qn("w:val"))
        has_style = True

    # Borders
    borders = _parse_cell_borders(tcPr)
    if borders:
        style.borders = borders
        has_style = True

    # Text alignment (from first paragraph)
    if cell.paragraphs:
        para = cell.paragraphs[0]
        pPr = para._p.pPr
        if pPr is not None:
            jc = pPr.find(qn("w:jc"))
            if jc is not None:
                style.text_align = jc.get(qn("w:val"))
                has_style = True

    return style if has_style else None


def _parse_table_style(table) -> Optional[TableStyle]:
    """Parse table-level styling."""
    tbl = table._tbl
    style = TableStyle()
    has_style = False

    # Table style name
    if table.style and table.style.name:
        style.style_name = table.style.name
        has_style = True

    # Table alignment
    if table.alignment is not None:
        # Convert enum to string
        alignment_map = {0: "left", 1: "center", 2: "right"}
        style.alignment = alignment_map.get(table.alignment, "left")
        has_style = True

    # Column widths from tblGrid
    tblGrid = tbl.tblGrid
    if tblGrid is not None:
        gridCols = tblGrid.findall(qn("w:gridCol"))
        for col in gridCols:
            w = col.get(qn("w:w"))
            if w:
                style.column_widths.append(int(w))
                has_style = True

    return style if has_style else None


def parse_table(
    table, numbering_tracker: Optional[NumberingTracker] = None
) -> Table:
    """Parse a python-docx table into our intermediate format.

    Handles merged cells by:
    1. Detecting colspan via tc.grid_span
    2. Detecting rowspan via tc.vMerge attribute
    3. Storing merge info in TableCell.colspan and TableCell.rowspan
    4. Skipping continuation cells (cells that are part of a merge but not the origin)

    Also captures raw OOXML for table, row, and cell properties to enable
    lossless round-tripping of complex styles.

    Note: We iterate over raw tc elements instead of row.cells because row.cells
    returns the same _Cell object for vertically merged cells, making it impossible
    to detect which rows are continuations.
    """
    from docx.table import _Cell

    parsed_table = Table()

    # Capture raw table properties (tblPr) and grid (tblGrid)
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    tblGrid = tbl.find(qn("w:tblGrid"))
    parsed_table.raw_tblPr = element_to_base64(tblPr)
    parsed_table.raw_tblGrid = element_to_base64(tblGrid)

    # Parse table-level styles (for backward compatibility)
    parsed_table.style = _parse_table_style(table)

    # Track vertical merges: grid_col -> {"cell": TableCell, "rowspan": int}
    # This helps us calculate rowspan and skip continuation cells
    vmerge_tracking: dict[int, dict] = {}

    for row_idx, row in enumerate(table.rows):
        parsed_row = TableRow()

        # Capture raw row properties (trPr)
        tr = row._tr
        trPr = tr.find(qn("w:trPr"))
        parsed_row.raw_xml = element_to_base64(trPr)

        grid_col = 0  # Track position in the grid

        # Iterate over raw tc elements to properly detect vMerge
        for tc in row._tr.tc_lst:
            colspan = tc.grid_span
            vmerge = (
                tc.vMerge
            )  # None = no merge, "restart" = start, "continue" = continuation

            # Check if this is a continuation of a vertical merge
            if vmerge == "continue":
                # This cell continues a vertical merge from above
                # Increment the rowspan of the origin cell and skip this one
                if grid_col in vmerge_tracking:
                    vmerge_tracking[grid_col]["rowspan"] += 1
                grid_col += colspan
                continue

            # This is either a new cell or the start of a vertical merge
            parsed_cell = TableCell(colspan=colspan, rowspan=1)

            # Capture raw cell properties (tcPr)
            tcPr = tc.find(qn("w:tcPr"))
            parsed_cell.raw_xml = element_to_base64(tcPr)

            # If there was a previous vertical merge in this column, finalize its rowspan
            if grid_col in vmerge_tracking:
                prev_info = vmerge_tracking[grid_col]
                prev_info["cell"].rowspan = prev_info["rowspan"]
                del vmerge_tracking[grid_col]

            # If this starts a new vertical merge, track it
            if vmerge == "restart":
                vmerge_tracking[grid_col] = {"cell": parsed_cell, "rowspan": 1}

            # Create a _Cell wrapper to access paragraphs and tables
            cell = _Cell(tc, table)

            # Parse cell styling
            parsed_cell.style = _parse_cell_style(tc, cell)

            # Parse cell content - cells can contain paragraphs and nested tables
            for para in cell.paragraphs:
                parsed_para = parse_paragraph(para, numbering_tracker)
                if parsed_para.runs:  # Only add non-empty paragraphs
                    parsed_cell.content.append(parsed_para)

            # Check for nested tables
            for nested_table in cell.tables:
                parsed_cell.content.append(
                    parse_table(nested_table, numbering_tracker)
                )

            # Ensure cell has at least one paragraph (empty cell)
            if not parsed_cell.content:
                parsed_cell.content.append(Paragraph(runs=[TextRun(text="")]))

            parsed_row.cells.append(parsed_cell)
            grid_col += colspan

        parsed_table.rows.append(parsed_row)

    # Update rowspan values for all tracked vertical merges
    for col_info in vmerge_tracking.values():
        col_info["cell"].rowspan = col_info["rowspan"]

    return parsed_table


def extract_style_numbering_map(doc) -> dict:
    """
    Extract the mapping of styles to their numbering properties.

    This creates a bidirectional map:
    - style_to_num: {styleName: {numId, ilvl}}
    - num_to_style: {numId: {ilvl: styleName}}

    This is needed because Word uses style-based numbering where different
    styles (SH1Legal, SH2Legal, etc.) correspond to different numbering levels.
    When the user changes the level in TipTap, we need to change the style.

    Returns:
        dict with 'style_to_num' and 'num_to_style' mappings
    """
    style_to_num = {}
    num_to_style = {}

    # Access the styles part
    styles_part = doc.part.styles
    if styles_part is None:
        return {"style_to_num": {}, "num_to_style": {}}

    styles_element = styles_part._element
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    for style in styles_element.findall(".//w:style", ns):
        style_id = style.get(qn("w:styleId"))
        if not style_id:
            continue

        pPr = style.find("w:pPr", ns)
        if pPr is None:
            continue

        numPr = pPr.find("w:numPr", ns)
        if numPr is None:
            continue

        # Get ilvl and numId
        ilvl_elem = numPr.find("w:ilvl", ns)
        numId_elem = numPr.find("w:numId", ns)

        # ilvl might be missing (inherits from style), default to 0
        ilvl = "0"
        if ilvl_elem is not None:
            ilvl = ilvl_elem.get(qn("w:val"), "0")

        # numId might be "inherit" or missing
        if numId_elem is None:
            continue
        numId = numId_elem.get(qn("w:val"))
        if not numId or numId == "inherit":
            continue

        # Store mapping
        style_to_num[style_id] = {"numId": numId, "ilvl": int(ilvl)}

        if numId not in num_to_style:
            num_to_style[numId] = {}
        num_to_style[numId][int(ilvl)] = style_id

    return {"style_to_num": style_to_num, "num_to_style": num_to_style}


def parse_docx(file_content: bytes) -> tuple[list, dict, dict]:
    """
    Parse a DOCX file and return a list of document elements, comments, and style map.

    Args:
        file_content: Raw bytes of the DOCX file

    Returns:
        Tuple of (list of Paragraph/Table/Section objects, dict of comments, style_numbering_map)
    """
    doc = Document(BytesIO(file_content))
    numbering_tracker = NumberingTracker(doc)

    # Extract comments from the document
    comments = extract_comments_from_docx(file_content)

    # Extract style-to-numbering mapping
    style_numbering_map = extract_style_numbering_map(doc)

    elements = []
    para_index = 0

    # Build lookup maps for paragraphs and tables by their XML element
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}

    # Iterate through document body in order
    for element in doc.element.body:
        if element.tag == qn("w:p"):
            # It's a paragraph
            if element in para_map:
                p = para_map[element]
                parsed = parse_paragraph(p, numbering_tracker, para_index)
                if parsed.runs or parsed.numbering:
                    elements.append(parsed)
                para_index += 1
        elif element.tag == qn("w:tbl"):
            # It's a table
            if element in table_map:
                t = table_map[element]
                elements.append(parse_table(t, numbering_tracker))

    return elements, comments, style_numbering_map
