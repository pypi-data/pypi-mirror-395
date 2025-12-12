"""
DOCX Exporter - Converts TipTap JSON back to Word document format.

This module handles the reverse conversion from TipTap editor JSON
back to a Word document (.docx) using python-docx.

Supports:
- Basic text formatting (bold, italic)
- Headings (H1-H6)
- Tables
- Track changes (insertions/deletions) via OOXML
- Comments via python-docx native API
"""

import base64
from io import BytesIO
from typing import Any, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

# Global counter for revision IDs
_revision_id_counter = 0


def _next_revision_id() -> str:
    """Generate next revision ID."""
    global _revision_id_counter
    _revision_id_counter += 1
    return str(_revision_id_counter)


def _reset_revision_counter():
    """Reset revision counter (call at start of each export)."""
    global _revision_id_counter
    _revision_id_counter = 0


def _clear_document_content(doc: Document) -> None:
    """
    Clear all content from a document while preserving styles.

    This removes paragraphs and tables from the document body but keeps
    styles, headers, footers, and other document properties intact.
    """
    # Get the body element
    body = doc.element.body

    # Remove all paragraph and table elements from the body
    # We iterate in reverse to avoid issues with modifying during iteration
    for child in list(body):
        # Keep section properties (w:sectPr) as they contain page layout
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _restore_raw_styles(tiptap_json: dict) -> dict:
    """
    Restore raw OOXML styles from rawStylesStorage node back into content.

    Finds the storage node, deserializes the data, and puts rawTblPr,
    rawTblGrid, rawXml back into the appropriate table/row/cell nodes.
    """
    import copy
    import json

    doc = copy.deepcopy(tiptap_json)
    content = doc.get("content", [])

    # Find and extract rawStylesStorage node
    raw_styles = {}
    new_content = []
    for node in content:
        if node.get("type") == "rawStylesStorage":
            data = node.get("attrs", {}).get("data", "{}")
            raw_styles = json.loads(data)
        else:
            new_content.append(node)

    doc["content"] = new_content

    if not raw_styles:
        return doc

    # Restore styles to nodes
    def process_node(node: dict):
        if node.get("type") == "table":
            attrs = node.get("attrs", {})
            table_id = attrs.get("id")
            if table_id:
                # Restore table-level raw XML
                if f"table:{table_id}:tblPr" in raw_styles:
                    attrs["rawTblPr"] = raw_styles[f"table:{table_id}:tblPr"]
                if f"table:{table_id}:tblGrid" in raw_styles:
                    attrs["rawTblGrid"] = raw_styles[
                        f"table:{table_id}:tblGrid"
                    ]

                # Restore row and cell styles
                for row_idx, row in enumerate(node.get("content", [])):
                    row_key = f"table:{table_id}:row:{row_idx}"
                    if row_key in raw_styles:
                        if "attrs" not in row:
                            row["attrs"] = {}
                        row["attrs"]["rawXml"] = raw_styles[row_key]

                    for cell_idx, cell in enumerate(row.get("content", [])):
                        cell_key = (
                            f"table:{table_id}:row:{row_idx}:cell:{cell_idx}"
                        )
                        if cell_key in raw_styles:
                            if "attrs" not in cell:
                                cell["attrs"] = {}
                            cell["attrs"]["rawXml"] = raw_styles[cell_key]

        for child in node.get("content", []):
            if isinstance(child, dict):
                process_node(child)

    for node in doc.get("content", []):
        process_node(node)

    return doc


def create_docx_from_tiptap(
    tiptap_json: dict,
    comments: Optional[list[dict]] = None,
    template_bytes: Optional[bytes] = None,
) -> BytesIO:
    """
    Convert TipTap JSON document to a Word document.

    Args:
        tiptap_json: TipTap document JSON with structure:
            {
                "type": "doc",
                "content": [...]
            }
        comments: Optional list of comment dictionaries with id, author, text, date
        template_bytes: Optional bytes of a .docx file to use as template.
            If provided, the template's styles will be preserved but content
            will be replaced with the TipTap content.

    Returns:
        BytesIO buffer containing the .docx file
    """
    _reset_revision_counter()

    # Restore raw styles from storage node
    tiptap_json = _restore_raw_styles(tiptap_json)

    # Create document from template or blank
    if template_bytes:
        template_buffer = BytesIO(template_bytes)
        doc = Document(template_buffer)
        # Clear existing content but keep styles
        _clear_document_content(doc)
    else:
        doc = Document()
    comments_dict = {c["id"]: c for c in (comments or [])}

    # Track which runs correspond to which comment IDs for later linking
    comment_runs_map: dict[str, list] = {}

    content = tiptap_json.get("content", [])
    for node in content:
        process_node(doc, node, comments_dict, comment_runs_map)

    # Add comments after processing all content
    _add_comments_to_document(doc, comments_dict, comment_runs_map)

    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def process_node(
    doc: Document,
    node: dict,
    comments_dict: dict,
    comment_runs_map: dict,
    table_cell=None,
) -> None:
    """
    Process a TipTap node and add it to the document.

    Args:
        doc: The python-docx Document object
        node: A TipTap node dictionary
        comments_dict: Dictionary of comment ID to comment data
        comment_runs_map: Map to track runs for each comment ID
        table_cell: Optional table cell to add content to (for nested content)
    """
    node_type = node.get("type")

    if node_type == "paragraph":
        process_paragraph(
            doc, node, comments_dict, comment_runs_map, table_cell
        )
    elif node_type == "heading":
        process_heading(doc, node, comments_dict, comment_runs_map, table_cell)
    elif node_type == "table":
        process_table(doc, node, comments_dict, comment_runs_map)
    elif node_type == "section":
        process_section(doc, node, comments_dict, comment_runs_map)


def process_paragraph(
    doc: Document,
    node: dict,
    comments_dict: dict,
    comment_runs_map: dict,
    table_cell=None,
) -> None:
    """Process a paragraph node."""
    if table_cell is not None:
        para = (
            table_cell.paragraphs[0]
            if table_cell.paragraphs
            else table_cell.add_paragraph()
        )
        if para.text == "" and len(table_cell.paragraphs) == 1:
            pass
        else:
            para = table_cell.add_paragraph()
    else:
        para = doc.add_paragraph()

    content = node.get("content", [])
    for text_node in content:
        if text_node.get("type") == "text":
            add_text_with_marks(
                para, text_node, comments_dict, comment_runs_map
            )


def process_heading(
    doc: Document,
    node: dict,
    comments_dict: dict,
    comment_runs_map: dict,
    table_cell=None,
) -> None:
    """Process a heading node."""
    level = node.get("attrs", {}).get("level", 1)
    content = node.get("content", [])

    if table_cell is not None:
        para = table_cell.add_paragraph()
        for text_node in content:
            if text_node.get("type") == "text":
                run = para.add_run(text_node.get("text", ""))
                run.bold = True
                _apply_basic_marks(run, text_node.get("marks", []))
    else:
        para = doc.add_heading(level=level)
        for text_node in content:
            if text_node.get("type") == "text":
                add_text_with_marks(
                    para, text_node, comments_dict, comment_runs_map
                )


def _base64_to_element(b64_string: str):
    """Deserialize a base64-encoded XML string to an lxml element."""
    if not b64_string:
        return None
    xml_bytes = base64.b64decode(b64_string.encode("ascii"))
    return etree.fromstring(xml_bytes)


def _restore_raw_cell_properties(cell, raw_xml: str) -> None:
    """Restore raw tcPr element from base64-encoded XML."""
    if not raw_xml:
        return

    tc = cell._tc
    new_tcPr = _base64_to_element(raw_xml)
    if new_tcPr is None:
        return

    # Remove existing tcPr if present
    existing_tcPr = tc.find(qn("w:tcPr"))
    if existing_tcPr is not None:
        tc.remove(existing_tcPr)

    # Insert new tcPr at the beginning (it should be first child)
    tc.insert(0, new_tcPr)


def _restore_raw_row_properties(row, raw_xml: str) -> None:
    """Restore raw trPr element from base64-encoded XML."""
    if not raw_xml:
        return

    tr = row._tr
    new_trPr = _base64_to_element(raw_xml)
    if new_trPr is None:
        return

    # Remove existing trPr if present
    existing_trPr = tr.find(qn("w:trPr"))
    if existing_trPr is not None:
        tr.remove(existing_trPr)

    # Insert new trPr at the beginning (it should be first child)
    tr.insert(0, new_trPr)


def _restore_raw_table_properties(
    table, raw_tblPr: str, raw_tblGrid: str
) -> None:
    """Restore raw tblPr and tblGrid elements from base64-encoded XML."""
    tbl = table._tbl

    # Restore tblPr
    if raw_tblPr:
        new_tblPr = _base64_to_element(raw_tblPr)
        if new_tblPr is not None:
            existing_tblPr = tbl.find(qn("w:tblPr"))
            if existing_tblPr is not None:
                tbl.remove(existing_tblPr)
            # tblPr should be first child
            tbl.insert(0, new_tblPr)

    # Restore tblGrid
    if raw_tblGrid:
        new_tblGrid = _base64_to_element(raw_tblGrid)
        if new_tblGrid is not None:
            existing_tblGrid = tbl.find(qn("w:tblGrid"))
            if existing_tblGrid is not None:
                tbl.remove(existing_tblGrid)
            # tblGrid should be after tblPr
            # Find the insert position (after tblPr if present)
            tblPr = tbl.find(qn("w:tblPr"))
            if tblPr is not None:
                tblPr_idx = list(tbl).index(tblPr)
                tbl.insert(tblPr_idx + 1, new_tblGrid)
            else:
                tbl.insert(0, new_tblGrid)


def _cleanup_merged_cells(table) -> None:
    """
    Remove extra tc elements that python-docx leaves behind after merging.

    When python-docx merges cells horizontally, it sets gridSpan on the first cell
    but leaves the subsequent cells. We need to remove cells that push the total
    grid span beyond the table's column count.

    For vertical merges (vMerge), we keep the continuation cells as they're needed
    for the vertical merge structure.
    """
    tbl = table._tbl

    # Get the number of grid columns from tblGrid
    tblGrid = tbl.tblGrid
    if tblGrid is None:
        return
    num_grid_cols = len(tblGrid.findall(qn("w:gridCol")))
    if num_grid_cols == 0:
        return

    for row in table.rows:
        tr = row._tr
        tc_list = list(tr.tc_lst)

        # Track grid position and find cells to remove
        grid_pos = 0
        cells_to_remove = []

        for tc in tc_list:
            tcPr = tc.find(qn("w:tcPr"))

            # Get gridSpan (horizontal span)
            gridSpan = 1
            if tcPr is not None:
                gs = tcPr.find(qn("w:gridSpan"))
                if gs is not None:
                    try:
                        gridSpan = int(gs.get(qn("w:val")))
                    except (ValueError, TypeError):
                        gridSpan = 1

            # If this cell would start beyond the grid, it's an artifact
            if grid_pos >= num_grid_cols:
                cells_to_remove.append(tc)
            else:
                grid_pos += gridSpan

        # Remove artifact cells from the row
        for tc in cells_to_remove:
            tr.remove(tc)


def _apply_cell_style(cell, attrs: dict) -> None:
    """Apply styling to a table cell from TipTap attributes."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Background color
    bg_color = attrs.get("backgroundColor")
    if bg_color:
        # Remove # prefix if present
        color = bg_color.lstrip("#")
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

    # Vertical alignment
    v_align = attrs.get("verticalAlign")
    if v_align:
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), v_align)
        tcPr.append(vAlign)

    # Cell width
    colwidth = attrs.get("colwidth")
    if colwidth and len(colwidth) > 0:
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(colwidth[0]))
        tcW.set(qn("w:type"), "dxa")
        # Remove existing tcW if present
        existing_tcW = tcPr.find(qn("w:tcW"))
        if existing_tcW is not None:
            tcPr.remove(existing_tcW)
        tcPr.insert(0, tcW)  # tcW should be first

    # Borders
    borders = attrs.get("borders")
    if borders:
        tcBorders = OxmlElement("w:tcBorders")
        for side in ["top", "bottom", "left", "right"]:
            border_data = borders.get(side)
            if border_data:
                border_elem = OxmlElement(f"w:{side}")
                if border_data.get("style"):
                    border_elem.set(qn("w:val"), border_data["style"])
                if border_data.get("width"):
                    border_elem.set(qn("w:sz"), str(border_data["width"]))
                if border_data.get("color"):
                    border_elem.set(qn("w:color"), border_data["color"])
                tcBorders.append(border_elem)
        tcPr.append(tcBorders)

    # Text alignment (applied to first paragraph)
    text_align = attrs.get("textAlign")
    if text_align and cell.paragraphs:
        para = cell.paragraphs[0]
        pPr = para._p.get_or_add_pPr()
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), text_align)
        pPr.append(jc)


def _apply_table_style(table, table_attrs: dict) -> None:
    """Apply styling to a table from TipTap attributes."""
    tbl = table._tbl

    # Table alignment
    alignment = table_attrs.get("alignment")
    if alignment:
        from docx.enum.table import WD_TABLE_ALIGNMENT

        alignment_map = {
            "left": WD_TABLE_ALIGNMENT.LEFT,
            "center": WD_TABLE_ALIGNMENT.CENTER,
            "right": WD_TABLE_ALIGNMENT.RIGHT,
        }
        if alignment in alignment_map:
            table.alignment = alignment_map[alignment]

    # Column widths
    colwidths = table_attrs.get("colwidths")
    if colwidths:
        # Set column widths via tblGrid
        tblGrid = tbl.tblGrid
        if tblGrid is not None:
            gridCols = tblGrid.findall(qn("w:gridCol"))
            for i, width in enumerate(colwidths):
                if i < len(gridCols):
                    gridCols[i].set(qn("w:w"), str(width))


def process_table(
    doc: Document, node: dict, comments_dict: dict, comment_runs_map: dict
) -> None:
    """Process a table node with support for merged cells (colspan/rowspan) and styling."""
    rows_data = node.get("content", [])
    if not rows_data:
        return

    # Calculate actual grid dimensions accounting for colspan
    num_rows = len(rows_data)
    num_cols = 0
    for row_node in rows_data:
        row_cols = 0
        for cell_node in row_node.get("content", []):
            colspan = cell_node.get("attrs", {}).get("colspan", 1)
            row_cols += colspan
        num_cols = max(num_cols, row_cols)

    if num_rows == 0 or num_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)

    # Get table-level attributes
    table_attrs = node.get("attrs", {})

    # Check if we have raw XML to restore (lossless round-trip)
    raw_tblPr = table_attrs.get("rawTblPr")
    raw_tblGrid = table_attrs.get("rawTblGrid")

    if raw_tblPr or raw_tblGrid:
        # Restore raw table properties for exact style preservation
        _restore_raw_table_properties(table, raw_tblPr, raw_tblGrid)
    else:
        # Fall back to individual style attributes
        style_name = table_attrs.get("styleName")
        if style_name:
            try:
                table.style = style_name
            except KeyError:
                pass  # Style not found, use default
        else:
            # Try to apply Table Grid style as default
            try:
                table.style = "Table Grid"
            except KeyError:
                try:
                    table.style = "TableGrid"
                except KeyError:
                    pass  # Use default table style

        # Apply table alignment and column widths
        _apply_table_style(table, table_attrs)

    # Track which cells are covered by merges (row_idx, col_idx) -> True
    covered_cells: set[tuple[int, int]] = set()

    # Track merges to apply after filling content: list of (start_cell, end_cell)
    merges_to_apply: list[tuple] = []

    for row_idx, row_node in enumerate(rows_data):
        cells_data = row_node.get("content", [])
        grid_col = 0  # Current position in the grid

        # Restore raw row properties if present
        row_attrs = row_node.get("attrs", {})
        raw_row_xml = row_attrs.get("rawXml")
        if raw_row_xml:
            _restore_raw_row_properties(table.rows[row_idx], raw_row_xml)

        for cell_node in cells_data:
            # Skip grid positions that are covered by previous merges
            while (row_idx, grid_col) in covered_cells and grid_col < num_cols:
                grid_col += 1

            if grid_col >= num_cols:
                break

            attrs = cell_node.get("attrs", {})
            colspan = attrs.get("colspan", 1)
            rowspan = attrs.get("rowspan", 1)

            cell = table.rows[row_idx].cells[grid_col]
            if cell.paragraphs:
                cell.paragraphs[0].clear()

            # Fill cell content
            cell_content = cell_node.get("content", [])
            _fill_cell_content(
                doc, cell, cell_content, comments_dict, comment_runs_map
            )

            # Apply cell styling - prefer raw XML for lossless round-trip
            raw_cell_xml = attrs.get("rawXml")
            if raw_cell_xml:
                _restore_raw_cell_properties(cell, raw_cell_xml)
            else:
                _apply_cell_style(cell, attrs)

            # Mark covered cells and prepare merge if needed
            if colspan > 1 or rowspan > 1:
                # Mark all cells in the merge range as covered
                for r in range(row_idx, row_idx + rowspan):
                    for c in range(grid_col, grid_col + colspan):
                        if r != row_idx or c != grid_col:  # Don't mark origin
                            covered_cells.add((r, c))

                # Record merge to apply
                end_row = min(row_idx + rowspan - 1, num_rows - 1)
                end_col = min(grid_col + colspan - 1, num_cols - 1)
                if end_row > row_idx or end_col > grid_col:
                    start_cell = table.rows[row_idx].cells[grid_col]
                    end_cell = table.rows[end_row].cells[end_col]
                    merges_to_apply.append((start_cell, end_cell))

            grid_col += colspan

    # Apply all merges after content is filled
    for start_cell, end_cell in merges_to_apply:
        try:
            start_cell.merge(end_cell)
        except Exception:
            pass  # Skip invalid merges gracefully

    # Clean up: remove extra tc elements that are artifacts of python-docx merge
    # When raw XML is restored, we need the table structure to match exactly
    _cleanup_merged_cells(table)


def _fill_cell_content(
    doc: Document,
    cell,
    cell_content: list,
    comments_dict: dict,
    comment_runs_map: dict,
) -> None:
    """Fill a table cell with content."""
    for i, content_node in enumerate(cell_content):
        if i == 0 and cell.paragraphs:
            if content_node.get("type") == "paragraph":
                for text_node in content_node.get("content", []):
                    if text_node.get("type") == "text":
                        add_text_with_marks(
                            cell.paragraphs[0],
                            text_node,
                            comments_dict,
                            comment_runs_map,
                        )
            elif content_node.get("type") == "heading":
                for text_node in content_node.get("content", []):
                    if text_node.get("type") == "text":
                        run = cell.paragraphs[0].add_run(
                            text_node.get("text", "")
                        )
                        run.bold = True
                        _apply_basic_marks(run, text_node.get("marks", []))
        else:
            process_node(
                doc,
                content_node,
                comments_dict,
                comment_runs_map,
                table_cell=cell,
            )


def process_section(
    doc: Document, node: dict, comments_dict: dict, comment_runs_map: dict
) -> None:
    """Process a section node (just process its content)."""
    content = node.get("content", [])
    for child_node in content:
        process_node(doc, child_node, comments_dict, comment_runs_map)


def add_text_with_marks(
    para, text_node: dict, comments_dict: dict, comment_runs_map: dict
) -> None:
    """
    Add text to a paragraph, handling track changes and comments.

    Args:
        para: The python-docx Paragraph object
        text_node: The TipTap text node
        comments_dict: Dictionary of comment data
        comment_runs_map: Map to track runs for comment linking
    """
    text = text_node.get("text", "")
    if not text:
        return

    marks = text_node.get("marks", [])

    # Check for track change marks
    insertion_mark = None
    deletion_mark = None
    comment_ids = []
    basic_marks = []

    for mark in marks:
        mark_type = mark.get("type")
        if mark_type == "insertion":
            insertion_mark = mark.get("attrs", {})
        elif mark_type == "deletion":
            deletion_mark = mark.get("attrs", {})
        elif mark_type == "comment":
            comment_id = mark.get("attrs", {}).get("commentId")
            if comment_id:
                comment_ids.append(comment_id)
        elif mark_type in ("bold", "italic"):
            basic_marks.append(mark)

    # Handle track changes
    if insertion_mark:
        _add_insertion(para, text, insertion_mark, basic_marks)
    elif deletion_mark:
        _add_deletion(para, text, deletion_mark, basic_marks)
    else:
        # Regular text
        run = para.add_run(text)
        _apply_basic_marks(run, basic_marks)

        # Track runs for comments
        for comment_id in comment_ids:
            if comment_id not in comment_runs_map:
                comment_runs_map[comment_id] = []
            comment_runs_map[comment_id].append(run)


def _add_insertion(para, text: str, attrs: dict, basic_marks: list) -> None:
    """
    Add text as an insertion (tracked change).

    Creates a w:ins element wrapping the run.
    """
    p_elem = para._p

    # Create w:ins element
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), _next_revision_id())
    ins.set(qn("w:author"), attrs.get("author", "Unknown"))
    if attrs.get("date"):
        ins.set(qn("w:date"), attrs["date"])

    # Create run inside insertion
    r = OxmlElement("w:r")

    # Add run properties for formatting
    if basic_marks:
        rPr = OxmlElement("w:rPr")
        for mark in basic_marks:
            if mark.get("type") == "bold":
                rPr.append(OxmlElement("w:b"))
            elif mark.get("type") == "italic":
                rPr.append(OxmlElement("w:i"))
        r.append(rPr)

    # Add text element
    t = OxmlElement("w:t")
    t.text = text
    # Preserve spaces
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)

    ins.append(r)
    p_elem.append(ins)


def _add_deletion(para, text: str, attrs: dict, basic_marks: list) -> None:
    """
    Add text as a deletion (tracked change).

    Creates a w:del element with w:delText inside.
    """
    p_elem = para._p

    # Create w:del element
    del_elem = OxmlElement("w:del")
    del_elem.set(qn("w:id"), _next_revision_id())
    del_elem.set(qn("w:author"), attrs.get("author", "Unknown"))
    if attrs.get("date"):
        del_elem.set(qn("w:date"), attrs["date"])

    # Create run inside deletion
    r = OxmlElement("w:r")

    # Add run properties for formatting
    if basic_marks:
        rPr = OxmlElement("w:rPr")
        for mark in basic_marks:
            if mark.get("type") == "bold":
                rPr.append(OxmlElement("w:b"))
            elif mark.get("type") == "italic":
                rPr.append(OxmlElement("w:i"))
        r.append(rPr)

    # Add deleted text element (w:delText instead of w:t)
    del_text = OxmlElement("w:delText")
    del_text.text = text
    del_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(del_text)

    del_elem.append(r)
    p_elem.append(del_elem)


def _apply_basic_marks(run, marks: list) -> None:
    """Apply basic formatting marks (bold, italic) to a run."""
    for mark in marks:
        mark_type = mark.get("type")
        if mark_type == "bold":
            run.bold = True
        elif mark_type == "italic":
            run.italic = True


def _add_comments_to_document(
    doc: Document, comments_dict: dict, comment_runs_map: dict
) -> None:
    """
    Add comments to the document using python-docx's add_comment API.

    Args:
        doc: The python-docx Document object
        comments_dict: Dictionary of comment ID to comment data
        comment_runs_map: Map of comment ID to list of runs
    """
    for comment_id, runs in comment_runs_map.items():
        if comment_id not in comments_dict:
            continue

        comment_data = comments_dict[comment_id]
        if not runs:
            continue

        # Get first and last run for the comment range
        first_run = runs[0]
        last_run = runs[-1] if len(runs) > 1 else runs[0]

        try:
            # Use python-docx native add_comment API
            doc.add_comment(
                first_run,
                text=comment_data.get("text", ""),
                author=comment_data.get("author", "Unknown"),
                initials=comment_data.get("initials", ""),
            )

            # If comment spans multiple runs, mark the range
            if len(runs) > 1:
                first_run.mark_comment_range(
                    last_run, doc.comments[-1].comment_id
                )
        except Exception as e:
            # If comment creation fails, continue without it
            print(f"Warning: Could not add comment {comment_id}: {e}")
