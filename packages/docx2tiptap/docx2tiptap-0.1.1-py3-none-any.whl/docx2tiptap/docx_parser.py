"""
DOCX Parser - Extracts structured content from Word documents.

Handles:
- Paragraphs with text formatting (bold, italic)
- Numbered/bulleted lists with computed numbering
- Tables with rich cell content
- Headings
- Track changes (insertions/deletions)
- Comments
"""

import base64
import uuid
from dataclasses import dataclass, field
from io import BytesIO
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from .comments_parser import (
    comments_to_dict,
    extract_comments_from_docx,
    get_text_with_comments,
)
from .revisions_parser import get_text_with_revisions, merge_adjacent_segments


@dataclass
class TextRun:
    """A run of text with formatting and revision/comment info."""

    text: str
    bold: bool = False
    italic: bool = False
    revision: Optional[dict] = (
        None  # {'type': 'insertion'|'deletion', 'id': ..., 'author': ..., 'date': ...}
    )
    comment_ids: list[str] = field(
        default_factory=list
    )  # List of comment IDs this text is part of


@dataclass
class Paragraph:
    """A paragraph containing text runs."""

    runs: list[TextRun] = field(default_factory=list)
    style: Optional[str] = None
    numbering: Optional[str] = None  # Computed numbering like "2a"
    level: int = 0  # Heading level (0 = not a heading)


@dataclass
class BorderStyle:
    """Border styling for a single edge."""

    style: Optional[str] = None  # single, double, dashed, dotted, nil, etc.
    width: Optional[int] = None  # Width in eighths of a point
    color: Optional[str] = None  # Hex color (e.g., "000000")


@dataclass
class CellBorders:
    """Border styles for all edges of a cell."""

    top: Optional[BorderStyle] = None
    bottom: Optional[BorderStyle] = None
    left: Optional[BorderStyle] = None
    right: Optional[BorderStyle] = None


@dataclass
class CellStyle:
    """Styling properties for a table cell."""

    width: Optional[int] = None  # Width in twips (1/1440 inch)
    background_color: Optional[str] = None  # Hex color (e.g., "f0f0f0")
    vertical_align: Optional[str] = None  # top, center, bottom
    borders: Optional[CellBorders] = None
    text_align: Optional[str] = None  # left, center, right, both (justify)


@dataclass
class TableStyle:
    """Styling properties for a table."""

    style_name: Optional[str] = None  # Named style (e.g., "Table Grid")
    alignment: Optional[str] = None  # left, center, right
    column_widths: list[int] = field(default_factory=list)  # Widths in twips


@dataclass
class TableCell:
    """A table cell containing block content."""

    content: list = field(
        default_factory=list
    )  # List of Paragraph or nested Table
    colspan: int = 1  # Number of columns this cell spans
    rowspan: int = 1  # Number of rows this cell spans
    style: Optional[CellStyle] = None
    # Raw OOXML for cell properties (tcPr) - preserves all styling
    raw_xml: Optional[str] = None


@dataclass
class TableRow:
    """A table row containing cells."""

    cells: list[TableCell] = field(default_factory=list)
    # Raw OOXML for row properties (trPr) - preserves row height, header row, etc.
    raw_xml: Optional[str] = None


@dataclass
class Table:
    """A table with rows and cells."""

    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    rows: list[TableRow] = field(default_factory=list)
    style: Optional[TableStyle] = None
    # Raw OOXML for table properties (tblPr) and grid (tblGrid)
    raw_tblPr: Optional[str] = None
    raw_tblGrid: Optional[str] = None


@dataclass
class Section:
    """A document section with content and optional children."""

    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    original_ref: Optional[str] = None
    level: int = 1
    title: str = ""
    content: list = field(default_factory=list)  # Paragraph, Table
    children: list["Section"] = field(default_factory=list)


class NumberingTracker:
    """Tracks list numbering state to compute actual numbers."""

    def __init__(self, document: Document):
        self.document = document
        self.counters: dict[str, list[int]] = (
            {}
        )  # numId -> [level0_count, level1_count, ...]
        self._numbering_formats = self._extract_numbering_formats()
        self._style_numbering = self._extract_style_numbering()

    def _extract_style_numbering(self) -> dict:
        """Extract numbering info from paragraph styles (numId and ilvl)."""
        style_num_info = {}  # styleId -> {'numId': ..., 'ilvl': ...}

        try:
            styles_part = self.document.part.styles
            if styles_part is None:
                return style_num_info
            styles_xml = styles_part._element
        except (KeyError, AttributeError):
            return style_num_info

        # First pass: collect direct numPr from styles
        for style in styles_xml.findall(qn("w:style")):
            style_id = style.get(qn("w:styleId"))
            if not style_id:
                continue

            pPr = style.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    ilvl_elem = numPr.find(qn("w:ilvl"))
                    numId_elem = numPr.find(qn("w:numId"))

                    ilvl = (
                        int(ilvl_elem.get(qn("w:val")))
                        if ilvl_elem is not None
                        else 0
                    )
                    numId = (
                        numId_elem.get(qn("w:val"))
                        if numId_elem is not None
                        else None
                    )

                    if numId and numId != "0":
                        style_num_info[style_id] = {
                            "numId": numId,
                            "ilvl": ilvl,
                        }

        # Second pass: resolve basedOn inheritance for styles without direct numPr
        # We need multiple passes to handle deep inheritance chains
        for _ in range(10):  # Max 10 levels of inheritance
            made_changes = False
            for style in styles_xml.findall(qn("w:style")):
                style_id = style.get(qn("w:styleId"))
                if not style_id or style_id in style_num_info:
                    continue

                basedOn = style.find(qn("w:basedOn"))
                if basedOn is not None:
                    base_style_id = basedOn.get(qn("w:val"))
                    if base_style_id and base_style_id in style_num_info:
                        style_num_info[style_id] = style_num_info[
                            base_style_id
                        ].copy()
                        made_changes = True

            if not made_changes:
                break

        return style_num_info

    def get_numbering_from_style(
        self, style_name: str
    ) -> tuple[str, int] | None:
        """Get numId and ilvl from a style name, resolving inheritance."""
        # Style names may differ from style IDs - try to find a match
        # First try exact match on style ID
        if style_name in self._style_numbering:
            info = self._style_numbering[style_name]
            return info["numId"], info["ilvl"]

        # Try matching by normalizing (remove spaces)
        normalized = style_name.replace(" ", "")
        for style_id, info in self._style_numbering.items():
            if style_id.replace(" ", "") == normalized:
                return info["numId"], info["ilvl"]

        return None

    def _extract_numbering_formats(self) -> dict:
        """Extract numbering format definitions from document."""
        formats = {}
        try:
            numbering_part = self.document.part.numbering_part
        except (KeyError, NotImplementedError):
            # Document has no numbering definitions
            return formats
        if numbering_part is None:
            return formats

        # Parse the numbering definitions
        numbering_xml = numbering_part._element

        # First pass: build a map of style names to their defining abstractNum IDs
        # Some abstractNums use numStyleLink to reference styles defined elsewhere
        style_to_abstract_id = {}
        for abstract_num in numbering_xml.findall(qn("w:abstractNum")):
            abstract_id = abstract_num.get(qn("w:abstractNumId"))
            style_link = abstract_num.find(qn("w:styleLink"))
            if style_link is not None:
                style_name = style_link.get(qn("w:val"))
                style_to_abstract_id[style_name] = abstract_id

        # Second pass: extract level definitions from each abstractNum
        for abstract_num in numbering_xml.findall(qn("w:abstractNum")):
            abstract_id = abstract_num.get(qn("w:abstractNumId"))
            levels = {}
            for lvl in abstract_num.findall(qn("w:lvl")):
                ilvl = int(lvl.get(qn("w:ilvl")))
                num_fmt_elem = lvl.find(qn("w:numFmt"))
                lvl_text_elem = lvl.find(qn("w:lvlText"))

                num_fmt = (
                    num_fmt_elem.get(qn("w:val"))
                    if num_fmt_elem is not None
                    else "decimal"
                )
                lvl_text = (
                    lvl_text_elem.get(qn("w:val"))
                    if lvl_text_elem is not None
                    else "%1."
                )

                levels[ilvl] = {"format": num_fmt, "text": lvl_text}
            formats[abstract_id] = levels

        # Third pass: resolve numStyleLink references
        # If an abstractNum uses numStyleLink, copy levels from the defining abstractNum
        for abstract_num in numbering_xml.findall(qn("w:abstractNum")):
            abstract_id = abstract_num.get(qn("w:abstractNumId"))
            num_style_link = abstract_num.find(qn("w:numStyleLink"))
            if num_style_link is not None:
                style_name = num_style_link.get(qn("w:val"))
                # Find the abstractNum that defines this style
                if style_name in style_to_abstract_id:
                    defining_abstract_id = style_to_abstract_id[style_name]
                    if defining_abstract_id in formats:
                        formats[abstract_id] = formats[defining_abstract_id]

        # Map numId to abstractNumId and extract startOverride values
        self._num_to_abstract = {}
        self._start_overrides = {}  # numId -> {ilvl: start_value}
        for num in numbering_xml.findall(qn("w:num")):
            num_id = num.get(qn("w:numId"))
            abstract_ref = num.find(qn("w:abstractNumId"))
            if abstract_ref is not None:
                self._num_to_abstract[num_id] = abstract_ref.get(qn("w:val"))

            # Check for lvlOverride with startOverride
            for lvl_override in num.findall(qn("w:lvlOverride")):
                ilvl = int(lvl_override.get(qn("w:ilvl")))
                start_override = lvl_override.find(qn("w:startOverride"))
                if start_override is not None:
                    start_val = int(start_override.get(qn("w:val")))
                    if num_id not in self._start_overrides:
                        self._start_overrides[num_id] = {}
                    # Store start-1 because get_number increments before returning
                    self._start_overrides[num_id][ilvl] = start_val - 1

        return formats

    def get_number(self, num_id: str, ilvl: int) -> str:
        """Get the computed number for a list item."""
        if num_id not in self.counters:
            self.counters[num_id] = [0] * 10  # Support up to 10 levels
            # Apply startOverride values if present
            if num_id in self._start_overrides:
                for override_ilvl, start_val in self._start_overrides[
                    num_id
                ].items():
                    self.counters[num_id][override_ilvl] = start_val

        # Increment current level, reset deeper levels
        self.counters[num_id][ilvl] += 1
        for i in range(ilvl + 1, 10):
            self.counters[num_id][i] = 0

        # Get format info
        abstract_id = self._num_to_abstract.get(num_id)
        if abstract_id and abstract_id in self._numbering_formats:
            level_info = self._numbering_formats[abstract_id].get(ilvl, {})
            num_fmt = level_info.get("format", "decimal")
            lvl_text = level_info.get("text", "%1.")
        else:
            num_fmt = "decimal"
            lvl_text = "%1."

        # Build the number string
        result = lvl_text
        for i in range(ilvl + 1):
            count = self.counters[num_id][i]
            formatted = self._format_number(
                count, num_fmt if i == ilvl else "decimal"
            )
            result = result.replace(f"%{i+1}", formatted)

        return result

    def _format_number(self, n: int, fmt: str) -> str:
        """Format a number according to the numbering format."""
        if fmt == "decimal":
            return str(n)
        elif fmt == "lowerLetter":
            return chr(ord("a") + n - 1) if 1 <= n <= 26 else str(n)
        elif fmt == "upperLetter":
            return chr(ord("A") + n - 1) if 1 <= n <= 26 else str(n)
        elif fmt == "lowerRoman":
            return self._to_roman(n).lower()
        elif fmt == "upperRoman":
            return self._to_roman(n)
        elif fmt == "bullet":
            return "•"
        else:
            return str(n)

    def _to_roman(self, n: int) -> str:
        """Convert integer to Roman numerals."""
        if n <= 0:
            return str(n)
        result = ""
        for value, numeral in [
            (1000, "M"),
            (900, "CM"),
            (500, "D"),
            (400, "CD"),
            (100, "C"),
            (90, "XC"),
            (50, "L"),
            (40, "XL"),
            (10, "X"),
            (9, "IX"),
            (5, "V"),
            (4, "IV"),
            (1, "I"),
        ]:
            while n >= value:
                result += numeral
                n -= value
        return result


def parse_paragraph(
    para,
    numbering_tracker: Optional[NumberingTracker] = None,
    para_index: int = 0,
) -> Paragraph:
    """Parse a python-docx paragraph into our intermediate format."""

    # Get text with revisions (track changes) and comments
    revision_segments = get_text_with_revisions(para._element, para_index)
    revision_segments = merge_adjacent_segments(revision_segments)

    comment_segments = get_text_with_comments(para._element, para_index)

    # Build a map of text positions to comment IDs
    comment_map = {}  # text -> [comment_ids]
    for seg in comment_segments:
        if seg["comments"]:
            comment_map[seg["text"]] = seg["comments"]

    # Convert segments to TextRuns
    runs = []
    for seg in revision_segments:
        if seg["text"]:
            # Try to find matching comment info
            comment_ids = comment_map.get(seg["text"], [])

            runs.append(
                TextRun(
                    text=seg["text"],
                    bold=seg.get("bold", False),
                    italic=seg.get("italic", False),
                    revision=seg.get("revision"),
                    comment_ids=comment_ids,
                )
            )

    # If no revision-aware segments found, fall back to simple parsing
    if not runs:
        for run in para.runs:
            if run.text:
                runs.append(
                    TextRun(
                        text=run.text,
                        bold=run.bold or False,
                        italic=run.italic or False,
                    )
                )

    # Detect heading level
    level = 0
    style_name = para.style.name if para.style else ""
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.replace("Heading ", ""))
        except ValueError:
            pass

    # Get numbering info - check direct numPr first, then fall back to style
    numbering = None
    if numbering_tracker:
        num_id = None
        ilvl = None

        # First check for direct numPr on the paragraph
        if para._element.pPr is not None:
            num_pr = para._element.pPr.find(qn("w:numPr"))
            if num_pr is not None:
                ilvl_elem = num_pr.find(qn("w:ilvl"))
                num_id_elem = num_pr.find(qn("w:numId"))
                if ilvl_elem is not None and num_id_elem is not None:
                    ilvl = int(ilvl_elem.get(qn("w:val")))
                    num_id = num_id_elem.get(qn("w:val"))

        # If no direct numPr, check if the style defines numbering
        if num_id is None and style_name:
            style_numbering = numbering_tracker.get_numbering_from_style(
                style_name
            )
            if style_numbering:
                num_id, ilvl = style_numbering

        # Compute the actual number if we have numbering info
        if num_id and num_id != "0":
            numbering = numbering_tracker.get_number(num_id, ilvl)

    return Paragraph(
        runs=runs, style=style_name, numbering=numbering, level=level
    )


def _element_to_base64(element) -> Optional[str]:
    """Serialize an lxml element to base64-encoded XML string."""
    if element is None:
        return None
    xml_bytes = etree.tostring(element, encoding="unicode")
    return base64.b64encode(xml_bytes.encode("utf-8")).decode("ascii")


def _base64_to_element(b64_string: str):
    """Deserialize a base64-encoded XML string to an lxml element."""
    if not b64_string:
        return None
    xml_bytes = base64.b64decode(b64_string.encode("ascii"))
    return etree.fromstring(xml_bytes)


def _parse_border_style(border_elem) -> Optional[BorderStyle]:
    """Parse a border element into a BorderStyle."""
    if border_elem is None:
        return None

    style = border_elem.get(qn("w:val"))
    if style == "nil":
        return None

    width_str = border_elem.get(qn("w:sz"))
    width = int(width_str) if width_str else None
    color = border_elem.get(qn("w:color"))

    return BorderStyle(style=style, width=width, color=color)


def _parse_cell_borders(tcPr) -> Optional[CellBorders]:
    """Parse cell borders from tcPr element."""
    if tcPr is None:
        return None

    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        return None

    top = _parse_border_style(tcBorders.find(qn("w:top")))
    bottom = _parse_border_style(tcBorders.find(qn("w:bottom")))
    left = _parse_border_style(tcBorders.find(qn("w:left")))
    right = _parse_border_style(tcBorders.find(qn("w:right")))

    if not any([top, bottom, left, right]):
        return None

    return CellBorders(top=top, bottom=bottom, left=left, right=right)


def _parse_cell_style(tc, cell) -> Optional[CellStyle]:
    """Parse cell styling from tc element."""
    tcPr = tc.tcPr
    if tcPr is None:
        return None

    style = CellStyle()
    has_style = False

    # Width
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is not None:
        w = tcW.get(qn("w:w"))
        if w:
            style.width = int(w)
            has_style = True

    # Background color (shading)
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        fill = shd.get(qn("w:fill"))
        if fill and fill.lower() not in ("auto", "ffffff"):
            style.background_color = fill
            has_style = True

    # Vertical alignment
    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is not None:
        style.vertical_align = vAlign.get(qn("w:val"))
        has_style = True

    # Borders
    borders = _parse_cell_borders(tcPr)
    if borders:
        style.borders = borders
        has_style = True

    # Text alignment (from first paragraph)
    if cell.paragraphs:
        para = cell.paragraphs[0]
        pPr = para._p.pPr
        if pPr is not None:
            jc = pPr.find(qn("w:jc"))
            if jc is not None:
                style.text_align = jc.get(qn("w:val"))
                has_style = True

    return style if has_style else None


def _parse_table_style(table) -> Optional[TableStyle]:
    """Parse table-level styling."""
    tbl = table._tbl
    style = TableStyle()
    has_style = False

    # Table style name
    if table.style and table.style.name:
        style.style_name = table.style.name
        has_style = True

    # Table alignment
    if table.alignment is not None:
        # Convert enum to string
        alignment_map = {0: "left", 1: "center", 2: "right"}
        style.alignment = alignment_map.get(table.alignment, "left")
        has_style = True

    # Column widths from tblGrid
    tblGrid = tbl.tblGrid
    if tblGrid is not None:
        gridCols = tblGrid.findall(qn("w:gridCol"))
        for col in gridCols:
            w = col.get(qn("w:w"))
            if w:
                style.column_widths.append(int(w))
                has_style = True

    return style if has_style else None


def parse_table(
    table, numbering_tracker: Optional[NumberingTracker] = None
) -> Table:
    """Parse a python-docx table into our intermediate format.

    Handles merged cells by:
    1. Detecting colspan via tc.grid_span
    2. Detecting rowspan via tc.vMerge attribute
    3. Storing merge info in TableCell.colspan and TableCell.rowspan
    4. Skipping continuation cells (cells that are part of a merge but not the origin)

    Also captures raw OOXML for table, row, and cell properties to enable
    lossless round-tripping of complex styles.

    Note: We iterate over raw tc elements instead of row.cells because row.cells
    returns the same _Cell object for vertically merged cells, making it impossible
    to detect which rows are continuations.
    """
    from docx.table import _Cell

    parsed_table = Table()

    # Capture raw table properties (tblPr) and grid (tblGrid)
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    tblGrid = tbl.find(qn("w:tblGrid"))
    parsed_table.raw_tblPr = _element_to_base64(tblPr)
    parsed_table.raw_tblGrid = _element_to_base64(tblGrid)

    # Parse table-level styles (for backward compatibility)
    parsed_table.style = _parse_table_style(table)

    # Track vertical merges: grid_col -> {"cell": TableCell, "rowspan": int}
    # This helps us calculate rowspan and skip continuation cells
    vmerge_tracking: dict[int, dict] = {}

    for row_idx, row in enumerate(table.rows):
        parsed_row = TableRow()

        # Capture raw row properties (trPr)
        tr = row._tr
        trPr = tr.find(qn("w:trPr"))
        parsed_row.raw_xml = _element_to_base64(trPr)

        grid_col = 0  # Track position in the grid

        # Iterate over raw tc elements to properly detect vMerge
        for tc in row._tr.tc_lst:
            colspan = tc.grid_span
            vmerge = (
                tc.vMerge
            )  # None = no merge, "restart" = start, "continue" = continuation

            # Check if this is a continuation of a vertical merge
            if vmerge == "continue":
                # This cell continues a vertical merge from above
                # Increment the rowspan of the origin cell and skip this one
                if grid_col in vmerge_tracking:
                    vmerge_tracking[grid_col]["rowspan"] += 1
                grid_col += colspan
                continue

            # This is either a new cell or the start of a vertical merge
            parsed_cell = TableCell(colspan=colspan, rowspan=1)

            # Capture raw cell properties (tcPr)
            tcPr = tc.find(qn("w:tcPr"))
            parsed_cell.raw_xml = _element_to_base64(tcPr)

            # If there was a previous vertical merge in this column, finalize its rowspan
            if grid_col in vmerge_tracking:
                prev_info = vmerge_tracking[grid_col]
                prev_info["cell"].rowspan = prev_info["rowspan"]
                del vmerge_tracking[grid_col]

            # If this starts a new vertical merge, track it
            if vmerge == "restart":
                vmerge_tracking[grid_col] = {"cell": parsed_cell, "rowspan": 1}

            # Create a _Cell wrapper to access paragraphs and tables
            cell = _Cell(tc, table)

            # Parse cell styling
            parsed_cell.style = _parse_cell_style(tc, cell)

            # Parse cell content - cells can contain paragraphs and nested tables
            for para in cell.paragraphs:
                parsed_para = parse_paragraph(para, numbering_tracker)
                if parsed_para.runs:  # Only add non-empty paragraphs
                    parsed_cell.content.append(parsed_para)

            # Check for nested tables
            for nested_table in cell.tables:
                parsed_cell.content.append(
                    parse_table(nested_table, numbering_tracker)
                )

            # Ensure cell has at least one paragraph (empty cell)
            if not parsed_cell.content:
                parsed_cell.content.append(Paragraph(runs=[TextRun(text="")]))

            parsed_row.cells.append(parsed_cell)
            grid_col += colspan

        parsed_table.rows.append(parsed_row)

    # Update rowspan values for all tracked vertical merges
    for col_info in vmerge_tracking.values():
        col_info["cell"].rowspan = col_info["rowspan"]

    return parsed_table


def parse_docx(file_content: bytes) -> tuple[list, dict]:
    """
    Parse a DOCX file and return a list of document elements and comments.

    Args:
        file_content: Raw bytes of the DOCX file

    Returns:
        Tuple of (list of Paragraph/Table/Section objects, dict of comments)
    """
    doc = Document(BytesIO(file_content))
    numbering_tracker = NumberingTracker(doc)

    # Extract comments from the document
    comments = extract_comments_from_docx(file_content)

    elements = []
    para_index = 0

    # Build lookup maps for paragraphs and tables by their XML element
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}

    # Iterate through document body in order
    for element in doc.element.body:
        if element.tag == qn("w:p"):
            # It's a paragraph
            if element in para_map:
                p = para_map[element]
                parsed = parse_paragraph(p, numbering_tracker, para_index)
                if parsed.runs or parsed.numbering:
                    elements.append(parsed)
                para_index += 1
        elif element.tag == qn("w:tbl"):
            # It's a table
            if element in table_map:
                t = table_map[element]
                elements.append(parse_table(t, numbering_tracker))

    return elements, comments


def _border_style_to_dict(border: Optional[BorderStyle]) -> Optional[dict]:
    """Convert BorderStyle to dict."""
    if border is None:
        return None
    return {
        "style": border.style,
        "width": border.width,
        "color": border.color,
    }


def _cell_borders_to_dict(borders: Optional[CellBorders]) -> Optional[dict]:
    """Convert CellBorders to dict."""
    if borders is None:
        return None
    result = {}
    if borders.top:
        result["top"] = _border_style_to_dict(borders.top)
    if borders.bottom:
        result["bottom"] = _border_style_to_dict(borders.bottom)
    if borders.left:
        result["left"] = _border_style_to_dict(borders.left)
    if borders.right:
        result["right"] = _border_style_to_dict(borders.right)
    return result if result else None


def _cell_style_to_dict(style: Optional[CellStyle]) -> Optional[dict]:
    """Convert CellStyle to dict."""
    if style is None:
        return None
    result = {}
    if style.width is not None:
        result["width"] = style.width
    if style.background_color:
        result["backgroundColor"] = style.background_color
    if style.vertical_align:
        result["verticalAlign"] = style.vertical_align
    if style.borders:
        borders = _cell_borders_to_dict(style.borders)
        if borders:
            result["borders"] = borders
    if style.text_align:
        result["textAlign"] = style.text_align
    return result if result else None


def _table_style_to_dict(style: Optional[TableStyle]) -> Optional[dict]:
    """Convert TableStyle to dict."""
    if style is None:
        return None
    result = {}
    if style.style_name:
        result["styleName"] = style.style_name
    if style.alignment:
        result["alignment"] = style.alignment
    if style.column_widths:
        result["columnWidths"] = style.column_widths
    return result if result else None


def elements_to_dict(elements: list) -> list[dict]:
    """Convert parsed elements to JSON-serializable dictionaries."""
    result = []

    for elem in elements:
        if isinstance(elem, Paragraph):
            runs_data = []
            for r in elem.runs:
                run_dict = {"text": r.text, "bold": r.bold, "italic": r.italic}
                if r.revision:
                    run_dict["revision"] = r.revision
                if r.comment_ids:
                    run_dict["commentIds"] = r.comment_ids
                runs_data.append(run_dict)

            result.append(
                {
                    "type": "paragraph",
                    "runs": runs_data,
                    "style": elem.style,
                    "numbering": elem.numbering,
                    "level": elem.level,
                }
            )
        elif isinstance(elem, Table):
            # Build rows with raw XML preservation
            rows_data = []
            for row in elem.rows:
                row_dict = {
                    "cells": [
                        {
                            "content": elements_to_dict(cell.content),
                            "colspan": cell.colspan,
                            "rowspan": cell.rowspan,
                            **(
                                {"style": _cell_style_to_dict(cell.style)}
                                if cell.style
                                else {}
                            ),
                            **(
                                {"rawXml": cell.raw_xml} if cell.raw_xml else {}
                            ),
                        }
                        for cell in row.cells
                    ]
                }
                # Add raw row XML if present
                if row.raw_xml:
                    row_dict["rawXml"] = row.raw_xml
                rows_data.append(row_dict)

            table_dict = {
                "type": "table",
                "id": elem.id,
                "rows": rows_data,
            }
            if elem.style:
                style_dict = _table_style_to_dict(elem.style)
                if style_dict:
                    table_dict["style"] = style_dict
            # Add raw table XML if present
            if elem.raw_tblPr:
                table_dict["rawTblPr"] = elem.raw_tblPr
            if elem.raw_tblGrid:
                table_dict["rawTblGrid"] = elem.raw_tblGrid
            result.append(table_dict)
        elif isinstance(elem, Section):
            result.append(
                {
                    "type": "section",
                    "id": elem.id,
                    "originalRef": elem.original_ref,
                    "level": elem.level,
                    "title": elem.title,
                    "content": elements_to_dict(elem.content),
                    "children": elements_to_dict(elem.children),
                }
            )

    return result
