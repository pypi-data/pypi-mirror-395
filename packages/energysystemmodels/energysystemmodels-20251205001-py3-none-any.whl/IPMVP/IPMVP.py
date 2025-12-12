import pandas as pd
from datetime import datetime
from sklearn.linear_model import LinearRegression
from scipy.stats import t
from sklearn.metrics import mean_squared_error
from sklearn.preprocessing import PolynomialFeatures



import numpy as np
#pip install statsmodels 
import statsmodels.api as sm

import docx
from docx.shared import Inches,RGBColor
import matplotlib.pyplot as plt

def docx_report(y_bl,y_report,df,df_report,conformite,conformite_report,table_incertitude,table_incertitude_report,df_savings,report_prediction,bl_prediction,start_baseline_period,end_baseline_period,start_reporting_period,end_reporting_period,site="****"):
    #pip install python-docx
    
    # Créer un nouveau document
    doc = docx.Document()

    # Ajouter un titre
    doc.add_heading('Rapport de mesure et vérification', 0)

    #Ajoutez un premier chapitre à votre rapport :
    doc.add_heading("1. Contexte du projet", level=1)
    # Ajouter un paragraphe de texte
    doc.add_paragraph("Ce projet s’inscrit dans le cadre des objectifs de réduction des consommations d’énergie sur le site de "+str(site)+ ". Il vise à réduire la consommation d’énergie liée au chauffage. La solution repose sur la mise en œuvre d’un système de pilotage des vannes thermostatiques en fonction des besoins thermiques des salles .")
    #doc.add_page_break()
    doc.add_heading("2. Description de la ou des AAPE", level=1)
    doc.add_paragraph("L’AAPE (Action d’Amélioration de la Performance Energétique) consiste à réguler les températures pièce par pièce et heure par heure en fonction des plannings d’occupation et de l’occupation réelle en agissant sur des vannes connectées. L’algorithme de régulation apprend également du comportement thermique du bâtiment afin d’adapterau plus juste la mise en marche et l’arrêt du chauffage.")
    doc.add_paragraph("Comment l’AAPE agit sur les consommations ?")
    doc.add_paragraph("Cette AAPE réduit la consommation d’énergie des zones inoccupées et améliore le confort. Par rapport à un planning de chauffe traditionnel pour lequel toutes les zones d’un bâtiment sont chauffées de manière quasiment égale voire continue en journée, les températures de consignes sont ici adaptées zone par zone en fonction de leur occupation prévisionnelle, de l’occupation réelle mesurée et du comportement de la pièce. L’occupation réelle mesurée pemet de réaliser des réduits de consigne pour générer une économie au-delà du respect du planning de consigne défini initialement. ")
    doc.add_paragraph("Cette AAPE génère en moyenne une économie de 25% sur la consommation de gaz.")


    #doc.add_page_break()
    doc.add_heading("3. Option M&V choisie pour l’AAPE", level=1)
    doc.add_paragraph("L’option choisie est :")
    doc.add_paragraph("l’option B, si le gain envisagé ne concerne qu'un sous ensemble du site et non l’ensemble du site. Il apparaît plus judicieux d’isoler l’AAPE sur son périmètre et de mesurer les consommations de gaz liées au chauffage avec l’ensemble des paramètres influents.")
    doc.add_paragraph("l’option C, si le gain envisagé concerne l’ensemble du site avec un compteur général d'un fournisseur d'énergie. Sont utilisées, ici les relevés de consommation de gaz de l’ensemble du site ainsi que les mesures de l'ensemble des facteurs infulents .")

    #doc.add_page_break()
    doc.add_heading("4. Rappel de la période de référence", level=1)
    doc.add_paragraph("La période de référence est du "+str(start_baseline_period)+" au "+str(end_baseline_period))
    #tracer la consommation de la periode de référence
    fig, ax = plt.subplots()
    ax.plot(y_bl.index, y_bl.iloc[:, 0], label=y_bl.columns[0])
    ax.set_xlabel('Date')
    ax.set_ylabel(y_bl.columns[0])
    ax.legend(loc='upper right')
    fig.savefig('conso_ref.png')
    doc.add_picture('conso_ref.png', width=Inches(6))
    # Ajouter une légende ou un titre à l'image
    caption = doc.add_paragraph('Figure 1. Consommation de la période de référence')
    caption.alignment = 1  # Aligner le texte au centre
    caption.bold = True  # Mettre en gras le texte de la légende

        ############ Créer une table avec la même forme que le dataframe#################
    y_bl_describe=y_bl.describe()
    y_bl_describe.loc["sum"]=y_bl.sum()
    y_bl_describe=y_bl_describe.round(2)
    table = doc.add_table(rows=y_bl_describe.shape[0]+1, cols=y_bl_describe.shape[1]+1)
    # Ajouter l'en-tête de colonne
    for j in range(y_bl_describe.shape[-1]):
        table.cell(0,j+1).text = y_bl_describe.columns[j]
    # Ajouter l'en-tête de ligne
    table.cell(0,0).text = 'Index'
    # Ajouter les données du dataframe
    for i in range(y_bl_describe.shape[0]):
        table.cell(i+1,0).text = str(y_bl_describe.index[i])
        for j in range(y_bl_describe.shape[-1]):
            table.cell(i+1,j+1).text = str(y_bl_describe.iloc[i,j])
    table.style = 'Colorful List Accent 1'
    ############################################################################

    #doc.add_page_break()
    doc.add_heading("5. La période de suivi", level=1)
    doc.add_paragraph("La période de suivi de ce rapport est du "+str(start_reporting_period)+" au "+str(end_reporting_period))
    
    fig, ax = plt.subplots()
    ax.plot(y_report.index, y_report.iloc[:, 0],label=y_report.columns[0])
    ax.set_xlabel('Date')
    ax.set_ylabel(y_report.columns[0])
    ax.legend(loc='upper right')

    fig.savefig('conso_suivi.png')
    doc.add_picture('conso_suivi.png', width=Inches(6))
    # Ajouter une légende ou un titre à l'image
    caption = doc.add_paragraph('Figure 2. Consommation de la période de suivi')
    caption.alignment = 1  # Aligner le texte au centre
    caption.bold = True  # Mettre en gras le texte de la légende

    ############ Créer une table avec la même forme que le dataframe#################
    y_report_describe=y_report.describe()
    y_report_describe.loc["sum"]=y_report.sum()
    y_report_describe=y_report_describe.round(2)
    table = doc.add_table(rows=y_report_describe.shape[0]+1, cols=y_report_describe.shape[1]+1)
    # Ajouter l'en-tête de colonne
    for j in range(y_report_describe.shape[-1]):
        table.cell(0,j+1).text = y_report_describe.columns[j]
    # Ajouter l'en-tête de ligne
    table.cell(0,0).text = 'Index'
    # Ajouter les données du dataframe
    for i in range(y_report_describe.shape[0]):
        table.cell(i+1,0).text = str(y_report_describe.index[i])
        for j in range(y_report_describe.shape[-1]):
            table.cell(i+1,j+1).text = str(y_report_describe.iloc[i,j])
    table.style = 'Colorful List Accent 1'
    ############################################################################
     
     
    #doc.add_page_break()
    doc.add_heading("6. Modèle de calcul des économies d'énergie", level=1)
    doc.add_heading("6.1. Modèle ANTE", level=2)
    doc.add_paragraph("Nous proposons ici d'évaluer un modèle sur la base la pèriode de référence qui sera appliqué sur la période de suivi pour le calcul des économies")

        ############ Créer une table avec la même forme que le dataframe#################
    table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1]+1)
    # Ajouter l'en-tête de colonne
    for j in range(df.shape[-1]):
        table.cell(0,j+1).text = df.columns[j]
    # Ajouter l'en-tête de ligne
    table.cell(0,0).text = 'Index'
    # Ajouter les données du dataframe
    for i in range(df.shape[0]):
        table.cell(i+1,0).text = str(df.index[i])
        for j in range(df.shape[-1]):
            table.cell(i+1,j+1).text = str(df.iloc[i,j])
    table.style = 'Colorful List Accent 1'
    ############################################################################
    doc.add_heading("6.1.1. conformité du Modèle ANTE", level=3)
    ############ Créer une table avec la même forme que le dataframe#################
    table = doc.add_table(rows=conformite.shape[0]+1, cols=conformite.shape[1]+1)
    # Ajouter l'en-tête de colonne
    for j in range(conformite.shape[-1]):
        table.cell(0,j+1).text = conformite.columns[j]
    # Ajouter l'en-tête de ligne
    table.cell(0,0).text = 'Index'
    # Ajouter les données du dataframe
    for i in range(conformite.shape[0]):
        table.cell(i+1,0).text = str(conformite.index[i])
        for j in range(conformite.shape[-1]):
            table.cell(i+1,j+1).text = str(conformite.iloc[i,j])
    table.style = 'Colorful List Accent 1'
    ############################################################################

    doc.add_heading("6.1.2. Incertitude du Modèle ANTE", level=3) 
        ############ Créer une table avec la même forme que le dataframe#################
    table = doc.add_table(rows=table_incertitude.shape[0]+1, cols=table_incertitude.shape[1]+1)
    # Ajouter l'en-tête de colonne
    for j in range(table_incertitude.shape[-1]):
        table.cell(0,j+1).text = table_incertitude.columns[j]
    # Ajouter l'en-tête de ligne
    table.cell(0,0).text = 'Index'
    # Ajouter les données du dataframe
    for i in range(table_incertitude.shape[0]):
        table.cell(i+1,0).text = str(table_incertitude.index[i])
        for j in range(table_incertitude.shape[-1]):
            table.cell(i+1,j+1).text = str(table_incertitude.iloc[i,j])
    table.style = 'Colorful List Accent 1'
    ############################################################################


  

    doc.add_heading("6.2. Modèle POST", level=2)
    doc.add_paragraph("Nous proposons ici d'évaluer un modèle sur la base la pèriode de suivi qui sera appliqué sur la période de référence dans le cas ou les données de la période de référence ne sont pas suffisantes pour établir un modèle valide")

        ############ Créer une table avec la même forme que le dataframe#################
    table = doc.add_table(rows=df_report.shape[0]+1, cols=df_report.shape[1]+1)
    # Ajouter l'en-tête de colonne
    for j in range(df_report.shape[-1]):
        table.cell(0,j+1).text = df_report.columns[j]
    # Ajouter l'en-tête de ligne
    table.cell(0,0).text = 'Index'
    # Ajouter les données du dataframe
    for i in range(df_report.shape[0]):
        table.cell(i+1,0).text = str(df_report.index[i])
        for j in range(df_report.shape[-1]):
            table.cell(i+1,j+1).text = str(df_report.iloc[i,j])
    table.style = 'Colorful List Accent 1'
    ############################################################################
    doc.add_heading("6.2.1. conformité du Modèle POST", level=3)
           ############ Créer une table avec la même forme que le dataframe#################
    table = doc.add_table(rows=conformite_report.shape[0]+1, cols=conformite_report.shape[1]+1)
    # Ajouter l'en-tête de colonne
    for j in range(conformite_report.shape[-1]):
        table.cell(0,j+1).text = conformite_report.columns[j]
    # Ajouter l'en-tête de ligne
    table.cell(0,0).text = 'Index'
    # Ajouter les données du dataframe
    for i in range(conformite_report.shape[0]):
        table.cell(i+1,0).text = str(conformite_report.index[i])
        for j in range(conformite_report.shape[-1]):
            table.cell(i+1,j+1).text = str(conformite_report.iloc[i,j])
    table.style = 'Colorful List Accent 1'
    ############################################################################

    doc.add_heading("6.2.2. Incertitude du Modèle POST", level=3)
        ############ Créer une table avec la même forme que le dataframe#################
    table = doc.add_table(rows=table_incertitude_report.shape[0]+1, cols=table_incertitude_report.shape[1]+1)
    # Ajouter l'en-tête de colonne
    for j in range(table_incertitude_report.shape[-1]):
        table.cell(0,j+1).text = table_incertitude_report.columns[j]
    # Ajouter l'en-tête de ligne
    table.cell(0,0).text = 'Index'
    # Ajouter les données du dataframe
    for i in range(table_incertitude_report.shape[0]):
        table.cell(i+1,0).text = str(table_incertitude_report.index[i])
        for j in range(table_incertitude_report.shape[-1]):
            table.cell(i+1,j+1).text = str(table_incertitude_report.iloc[i,j])
    table.style = 'Colorful List Accent 1'
    ############################################################################


    doc.add_heading("7. Calcul des économies d'énergie : option ANTE-POST et option POST-ANTE", level=1)
              ############ Créer une table avec la même forme que le dataframe#################
    table = doc.add_table(rows=df_savings.shape[0]+1, cols=df_savings.shape[1]+1)
  
  
            

    # Ajouter l'en-tête de colonne
    for j in range(df_savings.shape[-1]):
        table.cell(0,j+1).text = df_savings.columns[j]
    # Ajouter l'en-tête de ligne
    table.cell(0,0).text = 'Méthode de calcul'
    # Ajouter les données du dataframe
    for i in range(df_savings.shape[0]):
        table.cell(i+1,0).text = str(df_savings.index[i])
        for j in range(df_savings.shape[-1]):
            table.cell(i+1,j+1).text = str(df_savings.iloc[i,j])
    # Adding style to a table
    table.style = 'Colorful List Accent 1'
    ############################################################################


    #Economies ANTE-POST###############""
    fig, ax = plt.subplots(figsize=(8, 6))
    ax.plot(y_bl.index, y_bl.iloc[:, 0],label="relevé base de référence")
    ax.plot(y_report.index, y_report.iloc[:, 0],label="relevé période de suivi")
    ax.plot(report_prediction.index, report_prediction.iloc[:, 0],label="modèle IPMVP")

    ax.set_xlabel('Date')
    ax.set_ylabel(y_bl.columns[0])
    ax.legend(loc='upper right')

    fig.savefig('eco-ANTE-POST.png')
    doc.add_picture('eco-ANTE-POST.png', width=Inches(6))
    # Ajouter une légende ou un titre à l'image
    caption = doc.add_paragraph('Figure 3. Application du modèle sur la période POST')
    caption.alignment = 1  # Aligner le texte au centre
    caption.bold = True  # Mettre en gras le texte de la légende
    #########################""""

    #Economies POST-ANTE###############""
    fig, ax = plt.subplots(figsize=(8, 6))
    ax.plot(y_bl.index, y_bl.iloc[:, 0], label='relevé base de référence')
    ax.plot(y_report.index, y_report.iloc[:, 0], label='relevé période de suivi')
    ax.plot(bl_prediction.index, bl_prediction.iloc[:, 0],label="modèle IPMVP")
    ax.set_xlabel('Date')
    ax.set_ylabel(y_bl.columns[0])
    ax.legend(loc='upper right')

    fig.savefig('eco-POST-ANTE.png')
    doc.add_picture('eco-POST-ANTE.png', width=Inches(6))
    # Ajouter une légende ou un titre à l'image
    caption = doc.add_paragraph('Figure 4. Application du modèle sur la période ANTE')
    caption.alignment = 1  # Aligner le texte au centre
    caption.bold = True  # Mettre en gras le texte de la légende
    #########################""""

    #Ajoutez un deuxième chapitre à votre rapport :
    # doc.add_page_break()
    # doc.add_heading("Chapitre 2", level=1)

      # ajouter un nouveau style Equation
    # equation_style = doc.styles.add_style('Equation', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    # equation_style.base_style = doc.styles['Normal']
    # equation_style.font.size = docx.shared.Pt(12)
    # equation_style.font.name = 'Cambria Math'

    # ajouter un paragraphe contenant une équation avec des exposants
    # ajouter un paragraphe contenant l'équation
    # eq = doc.add_paragraph('', style='Equation')
    # eq.add_run('a').font.superscript = True
    # eq.add_run('x').font.subscript = True
    # eq.add_run('² + b = 0').font.subscript = True



    
    # Enregistrer le document
    
    doc.save(str(site)+'_rapport_M&V_'+datetime.now().strftime("%Y-%m-%d %H%M%S")+'.docx')

def creation_plynome(X,degree=1):
    if degree==1:
        X=X

    else: 
        # créer un objet PolynomialFeatures avec un degré de 2
        poly = PolynomialFeatures(degree=degree, include_bias=False)
        # transformer les données du DataFrame en un ensemble de données polynomiales de degré 2
        X_poly = poly.fit_transform(X)

        # Obtenir les noms des colonnes pour le nouveau dataframe
        col_names = poly.get_feature_names(X.columns)
        # Créer un nouveau dataframe avec les noms de colonnes
        X_poly = pd.DataFrame(X_poly, columns=col_names,index=X.index)
        # afficher les données transformées
        X=X_poly
    print("X_poly=",X)

    return X



# supp les val aber.
def drop_outliers (df,seuil_z_scores):
    # Select only the numeric columns
    df = df.select_dtypes(include=np.number)    
    # Calculate the z-score for the numeric columns
    z_scores = (df - df.mean()) / df.std()
    print("z_scores----------------------------",z_scores.max())
    # Get the location of the outliers
    outlier_location = np.argwhere((z_scores > seuil_z_scores).values)
    #valeur à supprimer
    df_out = df[(z_scores >= seuil_z_scores).any(axis=1)]
    print(df_out.index,type(df_out.index))

    if df_out.empty or df_out.isnull().all(axis=1).any():
        pass
    else:
        print("valeurs à supprimer",df_out,outlier_location )

    # # Exclude rows where any z-score is greater than 3
    df = df[(z_scores < seuil_z_scores).all(axis=1)]
    #print(df)
    return df,df_out


def conformite_ipmvp(r2,cv_remse,stat_t,stat_t_normale95): #,,
    conformite=[]
    if r2>=0.75:
        conformite.append([r2,True])
    else:
        conformite.append([r2,False])
    if cv_remse<=0.2:
        conformite.append([cv_remse,True])
    else:
        conformite.append([cv_remse,False])

    conformite=(pd.DataFrame(conformite,index=["r2","cv_remse"],columns=["valeur","conformité IPMVP"]))
    stat_t = stat_t.rename(index={stat_t.index[0]: 'valeur'})
    stat_t_T=stat_t.T
    stat_t_T["conformité IPMVP"] = stat_t_T["valeur"].apply(lambda x: True if (x is not None and float(x) >= stat_t_normale95) else False) 
    
    conformite=pd.concat([conformite,stat_t_T])
    # print(conformite)

    return conformite

# stat_t=pd.DataFrame([[-0.072971,12.51099,None,None]],columns=["stat_t_const",  "stat_t_DJU",  "stat_t_x2" , "stat_t_x3"],index=["valeur"])    
# conformite_ipmvp(0.9,0.1,stat_t)
def regression_model(X_bl,y_bl,approache):
    #approache="ANTE-POST"
    # Modeling :
    model = LinearRegression()
    model.fit(X_bl,y_bl)
    y_pred = model.predict(X_bl) # predicted consumption
    
    n = len(y_bl) # nombre de population
    p=X_bl.shape[1] #nb variables explicatives

    # Calcul du R2
    r2 = model.score(X_bl, y_bl)
    #print('R2:', r2)

    # Affichage des coefficients et l'intercept
    coefficients = []
    for p_ in range(p):
        coefficients.append(model.coef_[0, p_])
    
    const = model.intercept_[0]



    # Calcul du RMSE
    mse = mean_squared_error(y_bl, y_pred, squared=True)
    ddof=max(n-p-1,0) # le degrès de liberté (comme celui d'excel)
    rmse=(mse*((n)/ddof))**0.5 #comme celui d'excel

    stat_t_normale95=t.ppf(1-0.05/2, ddof)
    print("statistique ************************************* loi normale",ddof,stat_t_normale95)
    
   
    #calcul du CV_RMSE
    cv_rmse=rmse/y_bl.mean()[0]

    ######################calculs d'incertitude##################################################
    niveau_confiance=0.8
    gamma=(1-niveau_confiance)/2
    stat_t_normale=t.ppf(niveau_confiance+gamma, ddof)
    
    print("statistique t loi normale",ddof,stat_t_normale)
    precision_absolue=stat_t_normale*rmse
    precision_relative=precision_absolue/y_bl.mean()[0]


    table_incertitude=pd.DataFrame({"valeurs":[gamma,niveau_confiance,stat_t_normale,rmse,precision_absolue,precision_relative]},index=["gamma","niveau_confiance","stat_t_normale","Erreur type (rmse)","precision_absolue +/-","precision_relative"]).round(2)
    print(table_incertitude)
    #############################################################################################################################################################################################################
   

    # calcul des erreurs standard...................................................
    # Calculer les erreurs types des coefficients
    # add a constant term to the input data for statsmodels
    X_bl_withConst = sm.add_constant(X_bl)
    # fit a linear regression model using statsmodels
    model_sm = sm.OLS(y_bl, X_bl_withConst).fit()
    # extract the standard errors of the coefficients from the statsmodels model
    serr = pd.DataFrame(model_sm.bse).T
    serr = serr.rename(index={0: approache})
    #serr=serr.add_prefix("SE_")

    # print the standard errors
    print("serr====================",serr)

    const_coef=(pd.DataFrame([np.concatenate(([const], coefficients))],columns=np.concatenate((['const'], X_bl.columns)).tolist()).rename(index={0: approache})).astype(float)
    print("const_coef=========================",const_coef)

    #calcul de la statistique t de chaque coef
    #stat_t=const_coef.astype(float).div(serr.astype(float))
   
    print(list(serr.columns),list(const_coef.columns))
    stat_t=const_coef.div(serr)
    
    #print("stat_t===============",stat_t)

    stat_t=stat_t.add_prefix("stat_t_")

    print("stat_t===============",stat_t)
    
    #arrondir
    serr=serr.round(1)
    stat_t=stat_t.round(1)
    r2=round(r2,2)
    rmse=round(rmse,1)
    cv_rmse=round(cv_rmse,2)
    ddof=round(ddof,0)
    const_coef=const_coef.round(10)

   
    

    # Modèle y_predicted_arr = modèle y avec arrondis [const,DJU,planning]:
    df=pd.DataFrame([[r2,rmse,cv_rmse,ddof]], columns=["r2","rmse","cv_rmse","ddof"], index=[approache])
  
    df=pd.merge( const_coef.add_prefix("coef_"),df, left_index=True, right_index=True)
    df=pd.merge(df, serr.add_prefix("serr_"), left_index=True, right_index=True)
    df=pd.merge(df,stat_t, left_index=True, right_index=True).T

    y_pred=pd.DataFrame(y_pred)   
    y_pred = y_pred.set_index(y_bl.index)  

    conformite=conformite_ipmvp(r2,cv_rmse,stat_t,stat_t_normale95)

    return model,y_pred,df,conformite,table_incertitude


def Mathematical_Models(y,X,start_baseline_period,end_baseline_period,start_reporting_period,end_reporting_period,print_report=False,seuil_z_scores=8,degree=1,site="****"):
    #remettre les données dans l'ordre
    y=pd.DataFrame(y).sort_index(axis=0, ascending=True, inplace=False, kind="quicksort")
    X=pd.DataFrame(X).sort_index(axis=0, ascending=True, inplace=False, kind="quicksort")
    y,y_out=drop_outliers(y,seuil_z_scores)
    #X,X_out=drop_outliers(X,3)
    X= X.drop(y_out.index, axis=0)

    ###################Ajouter etude des corrélations pol
    X=creation_plynome(X,degree=degree)


    #calculer de delta t
    dt = y.index.to_series().diff().astype('timedelta64[h]')
    curent_delta_time=dt.median()
    print("curent_delta_time (heures)",curent_delta_time)


    # if end_reporting_period is None :
    #     end_reporting_period = df.tail(1).index.strftime("%Y/%m/%d")[0]
    #     print("the end of the reporting period is the current date",end_reporting_period)

    #le jeu de données de la periode de référence
    X_bl = X.loc[(X.index >= start_baseline_period) & (X.index <= end_baseline_period)]
    #print("X base line", X_bl.head)

    y_bl = y.loc[(y.index >= start_baseline_period) & (y.index <= end_baseline_period)]
    print("y base line", y_bl)

    # jeu de données de la periode de suivi
    X_report = X.loc[(X.index >= start_reporting_period) & (X.index <= end_reporting_period)]
    y_report = y.loc[(y.index >= start_reporting_period) & (y.index <= end_reporting_period)]

    
##########################################################################################################################""
    # # Modeling :
   

    model_bl,y_pred,df,conformite,table_incertitude=regression_model(X_bl,y_bl,"ANTE-POST") 

    model_report,y_pred_report,df_report,conformite_report,table_incertitude_report=regression_model(X_report,y_report,"POST-ANTE") 

    report_prediction=model_bl.predict(X_report)
    sum_report_prediction= sum(report_prediction)[0]
    bl_prediction=model_report.predict(X_bl)
    sum_bl_prediction= sum(bl_prediction)[0]

    report_prediction=pd.DataFrame(report_prediction,index=X_report.index)
    print("report_prediction",report_prediction)

    bl_prediction=pd.DataFrame(bl_prediction,index=X_bl.index)
    print("bl_prediction",bl_prediction)

    sum_bl=y_bl.sum().values[0]
    sum_report=y_report.sum().values[0]

    savings_post=(sum_report_prediction-sum_report)/sum_report*100
    saving_ante=(-sum_bl_prediction+sum_bl)/sum_bl*100 #inversé pour avoir des éco positives

    df_savings=pd.DataFrame({"ANTE-POST":[sum_report,sum_report_prediction,savings_post],"POST-ANTE":[sum_bl,sum_bl_prediction,saving_ante]},index=["Relevé de consommation","Prédiction","pourcentage d'économie>0"]).round(2)
    
   
    if print_report==True:
        x=docx_report(y_bl,y_report,df,df_report,conformite,conformite_report,table_incertitude,table_incertitude_report,df_savings,report_prediction,bl_prediction,start_baseline_period,end_baseline_period,start_reporting_period,end_reporting_period,site=site)

    return  y_pred,df,conformite,table_incertitude,y_pred_report,df_report,conformite_report,table_incertitude_report,df_savings
   
  

