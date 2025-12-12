"""Format converter for report generation (Markdown, HTML, PDF, DOCX)."""

import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Any

import csv
import markdown
from docx import Document
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)

try:
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl not available. Excel generation will be disabled.")

try:
    from markdownify import markdownify as md
    MARKDOWNIFY_AVAILABLE = True
except ImportError:
    MARKDOWNIFY_AVAILABLE = False
    logger.warning("markdownify not available. HTML to Markdown conversion will be limited.")

WEASYPRINT_AVAILABLE = False
WEASYPRINT_HTML = None
WEASYPRINT_CSS = None

try:
    import sys
    import io
    
    stderr_backup = sys.stderr
    sys.stderr = io.StringIO()
    
    try:
        from weasyprint import HTML, CSS
        WEASYPRINT_AVAILABLE = True
        WEASYPRINT_HTML = HTML
        WEASYPRINT_CSS = CSS
    finally:
        sys.stderr = stderr_backup
except (ImportError, OSError) as e:
    WEASYPRINT_AVAILABLE = False
    error_msg = str(e)
    if "libgobject" in error_msg.lower() or "dlopen" in error_msg.lower():
        logger.debug("WeasyPrint not available (system library missing). PDF generation will be disabled.")
    else:
        logger.debug("WeasyPrint not available. PDF generation will be disabled. Error: %s", error_msg)


class FormatConverter:
    """Converts content between formats (Markdown, HTML, PDF, DOCX)."""

    def __init__(self):
        """Initialize format converter."""
        self.md_extensions = [
            "extra",
            "codehilite",
            "tables",
            "fenced_code",
            "toc",
        ]

    def markdown_to_html(
        self,
        markdown_content: str,
        styles: dict[str, Any] | None = None,
    ) -> str:
        """Convert Markdown to HTML.

        Args:
            markdown_content: Markdown content
            styles: CSS styles dictionary

        Returns:
            HTML content
        """
        html_content = markdown.markdown(
            markdown_content,
            extensions=self.md_extensions,
        )

        css_styles = self._build_css(styles or {})

        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        {css_styles}
    </style>
</head>
<body>
    {html_content}
</body>
</html>"""

        return html

    def markdown_to_pdf(
        self,
        markdown_content: str,
        styles: dict[str, Any] | None = None,
        branding: dict[str, Any] | None = None,
    ) -> bytes:
        """Convert Markdown to PDF.

        Args:
            markdown_content: Markdown content
            styles: CSS styles dictionary
            branding: Branding configuration

        Returns:
            PDF bytes

        Raises:
            ValueError: If WeasyPrint not available
        """
        if not WEASYPRINT_AVAILABLE or WEASYPRINT_HTML is None or WEASYPRINT_CSS is None:
            raise ValueError(
                "WeasyPrint not available. Install with: pip install weasyprint "
                "and system dependencies (see https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#installation)"
            )

        html_content = self.markdown_to_html(markdown_content, styles)

        css_styles = self._build_css(styles or {}, branding)

        try:
            pdf_bytes = WEASYPRINT_HTML(string=html_content).write_pdf(
                stylesheets=[WEASYPRINT_CSS(string=css_styles)]
            )
            return pdf_bytes
        except (RuntimeError, OSError, ValueError) as e:
            logger.error("PDF generation failed: %s", e)
            raise ValueError(f"PDF generation failed: {e}") from e
        except Exception as e:
            logger.error("Unexpected error during PDF generation: %s", e)
            raise ValueError(f"PDF generation failed: Unexpected error") from e

    def _parse_inline_markdown(self, paragraph, text: str) -> None:
        """Parse inline markdown formatting and add runs to paragraph.

        Handles: **bold**, *italic*, ***bold italic***, `code`, ~~strikethrough~~

        Args:
            paragraph: The docx paragraph to add runs to
            text: The text containing inline markdown
        """
        # Pattern to match inline markdown: **bold**, *italic*, `code`, ~~strike~~
        # Order matters: check *** before ** before *
        pattern = re.compile(
            r'(\*\*\*(.+?)\*\*\*)'  # ***bold italic***
            r'|(\*\*(.+?)\*\*)'      # **bold**
            r'|(\*(.+?)\*)'          # *italic*
            r'|(`(.+?)`)'            # `code`
            r'|(~~(.+?)~~)'          # ~~strikethrough~~
        )

        last_end = 0
        for match in pattern.finditer(text):
            # Add any text before this match as plain text
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            # Determine which group matched and apply formatting
            if match.group(2):  # ***bold italic***
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif match.group(4):  # **bold**
                run = paragraph.add_run(match.group(4))
                run.bold = True
            elif match.group(6):  # *italic*
                run = paragraph.add_run(match.group(6))
                run.italic = True
            elif match.group(8):  # `code`
                run = paragraph.add_run(match.group(8))
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            elif match.group(10):  # ~~strikethrough~~
                run = paragraph.add_run(match.group(10))
                run.font.strike = True

            last_end = match.end()

        # Add any remaining text after the last match
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def _add_formatted_paragraph(self, doc, text: str, style: str | None = None) -> None:
        """Add a paragraph with inline markdown formatting parsed.

        Args:
            doc: The docx Document
            text: Text that may contain inline markdown
            style: Optional paragraph style (e.g., "List Bullet")
        """
        if style:
            paragraph = doc.add_paragraph(style=style)
        else:
            paragraph = doc.add_paragraph()

        self._parse_inline_markdown(paragraph, text)

    def markdown_to_docx(
        self,
        markdown_content: str,
        branding: dict[str, Any] | None = None,
    ) -> bytes:
        """Convert Markdown to DOCX.

        Args:
            markdown_content: Markdown content
            branding: Branding configuration

        Returns:
            DOCX bytes
        """
        doc = Document()

        if branding:
            if branding.get("logo_path"):
                try:
                    logo_path = Path(branding["logo_path"])
                    if logo_path.exists():
                        doc.add_picture(str(logo_path), width=Inches(2))
                except Exception as e:
                    logger.warning("Failed to add logo: %s", e)

            if branding.get("company_name"):
                title = doc.add_heading(branding["company_name"], 0)
                title.alignment = 1

        lines = markdown_content.split("\n")
        in_code_block = False

        for line in lines:
            stripped_line = line.strip()

            if not stripped_line:
                continue

            # Handle code blocks
            if stripped_line.startswith("```"):
                in_code_block = not in_code_block
                continue

            if in_code_block:
                # Add code block content with monospace font
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(line)  # Preserve original indentation
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                continue

            # Handle headings (strip inline formatting for headings)
            if stripped_line.startswith("# "):
                heading_text = re.sub(r'\*+|`|~~', '', stripped_line[2:])
                doc.add_heading(heading_text, level=1)
            elif stripped_line.startswith("## "):
                heading_text = re.sub(r'\*+|`|~~', '', stripped_line[3:])
                doc.add_heading(heading_text, level=2)
            elif stripped_line.startswith("### "):
                heading_text = re.sub(r'\*+|`|~~', '', stripped_line[4:])
                doc.add_heading(heading_text, level=3)
            elif stripped_line.startswith("#### "):
                heading_text = re.sub(r'\*+|`|~~', '', stripped_line[5:])
                doc.add_heading(heading_text, level=4)
            elif stripped_line.startswith("##### "):
                heading_text = re.sub(r'\*+|`|~~', '', stripped_line[6:])
                doc.add_heading(heading_text, level=5)
            # Handle bullet lists
            elif stripped_line.startswith("- ") or stripped_line.startswith("* "):
                self._add_formatted_paragraph(doc, stripped_line[2:], style="List Bullet")
            # Handle numbered lists
            elif re.match(r'^\d+\.\s', stripped_line):
                list_text = re.sub(r'^\d+\.\s', '', stripped_line)
                self._add_formatted_paragraph(doc, list_text, style="List Number")
            # Handle blockquotes
            elif stripped_line.startswith("> "):
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.5)
                self._parse_inline_markdown(paragraph, stripped_line[2:])
            # Regular paragraph with inline formatting
            else:
                self._add_formatted_paragraph(doc, stripped_line)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def html_to_markdown(self, html_content: str) -> str:
        """Convert HTML to Markdown.

        Args:
            html_content: HTML content

        Returns:
            Markdown content
        """
        if not MARKDOWNIFY_AVAILABLE:
            logger.warning("markdownify not available. Returning HTML content.")
            return html_content
        return md(html_content)

    def _build_css(
        self,
        styles: dict[str, Any],
        branding: dict[str, Any] | None = None,
    ) -> str:
        """Build CSS from styles dictionary.

        Args:
            styles: Styles dictionary
            branding: Branding configuration

        Returns:
            CSS string
        """
        css_parts = []

        primary_color = branding.get("primary_color", "#2563eb") if branding else "#2563eb"
        secondary_color = branding.get("secondary_color", "#64748b") if branding else "#64748b"
        font_family = branding.get("font_family", "Arial, sans-serif") if branding else "Arial, sans-serif"

        css_parts.append(f"""
            @page {{
                size: A4;
                margin: 2cm;
            }}
            body {{
                font-family: {font_family};
                line-height: 1.6;
                color: #333;
                font-size: 11pt;
            }}
            h1 {{
                color: {primary_color};
                border-bottom: 2px solid {primary_color};
                padding-bottom: 10px;
                font-size: 22pt;
                margin-top: 0;
            }}
            h2 {{
                color: {secondary_color};
                margin-top: 24px;
                font-size: 16pt;
                border-bottom: 1px solid #e5e7eb;
                padding-bottom: 6px;
            }}
            h3 {{
                color: #374151;
                margin-top: 18px;
                font-size: 13pt;
            }}
            h4 {{
                color: #4b5563;
                margin-top: 14px;
                font-size: 11pt;
            }}
            p {{
                margin: 8px 0;
                text-align: justify;
            }}
            ul, ol {{
                margin: 10px 0;
                padding-left: 24px;
            }}
            li {{
                margin: 4px 0;
                line-height: 1.5;
            }}
            code {{
                background-color: #f3f4f6;
                padding: 2px 6px;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
                font-size: 10pt;
            }}
            pre {{
                background-color: #f3f4f6;
                padding: 12px;
                border-radius: 5px;
                overflow-x: auto;
                font-size: 9pt;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin: 16px 0;
                font-size: 10pt;
            }}
            th, td {{
                border: 1px solid #d1d5db;
                padding: 8px 12px;
                text-align: left;
            }}
            th {{
                background-color: {primary_color};
                color: white;
                font-weight: bold;
            }}
            tr:nth-child(even) {{
                background-color: #f9fafb;
            }}
            hr {{
                border: none;
                border-top: 1px solid #e5e7eb;
                margin: 20px 0;
            }}
            strong {{
                color: #1f2937;
            }}
            .severity-critical {{
                color: #dc2626;
                font-weight: bold;
            }}
            .severity-high {{
                color: #ea580c;
                font-weight: bold;
            }}
            .severity-medium {{
                color: #ca8a04;
            }}
            .severity-low {{
                color: #16a34a;
            }}
        """)

        if styles:
            for selector, properties in styles.items():
                props_str = "; ".join(f"{k}: {v}" for k, v in properties.items())
                css_parts.append(f"{selector} {{ {props_str} }}")

        return "\n".join(css_parts)

    def _strip_markdown(self, text: str) -> str:
        """Strip markdown formatting from text.

        Args:
            text: Text with markdown formatting

        Returns:
            Plain text without markdown
        """
        # Remove bold/italic markers
        text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        # Remove code backticks
        text = re.sub(r'`(.+?)`', r'\1', text)
        # Remove strikethrough
        text = re.sub(r'~~(.+?)~~', r'\1', text)
        return text

    def _parse_compliance_data(self, markdown_content: str) -> list[dict]:
        """Parse compliance report markdown into structured data.

        Args:
            markdown_content: Markdown content of compliance report

        Returns:
            List of control dictionaries with structured data
        """
        controls = []
        lines = markdown_content.split("\n")

        current_resource = ""
        current_control = {}

        for line in lines:
            line = line.strip()
            if not line or line == "---":
                continue

            # Detect resource sections (e.g., "### Cloud SQL (Cloud SQL)")
            if line.startswith("### ") or line.startswith("## "):
                prefix = "### " if line.startswith("### ") else "## "
                resource = self._strip_markdown(line[len(prefix):])
                if "(" in resource:
                    current_resource = resource.split("(")[0].strip()
                else:
                    current_resource = resource

            # Detect control headers (e.g., "#### PCI-DSS PCI-DSS-1.3.7:")
            elif line.startswith("#### "):
                if current_control:
                    controls.append(current_control)
                control_text = self._strip_markdown(line[5:])
                # Parse control ID (e.g., "PCI-DSS PCI-DSS-1.3.7: Description")
                if ":" in control_text:
                    parts = control_text.split(":", 1)
                    control_id = parts[0].strip()
                    description = parts[1].strip() if len(parts) > 1 else ""
                else:
                    control_id = control_text
                    description = ""
                current_control = {
                    "resource": current_resource,
                    "control_id": control_id,
                    "description": description,
                    "severity": "",
                    "details": ""
                }

            # Detect severity lines
            elif line.startswith("**Severity**:") or line.startswith("*Severity*:"):
                severity = self._strip_markdown(line.split(":", 1)[1].strip()) if ":" in line else ""
                if current_control:
                    current_control["severity"] = severity

            # Additional details (paragraphs after control header)
            elif current_control and not line.startswith("#") and not line.startswith("|"):
                if current_control["details"]:
                    current_control["details"] += " " + self._strip_markdown(line)
                else:
                    current_control["details"] = self._strip_markdown(line)

        # Don't forget the last control
        if current_control:
            controls.append(current_control)

        return controls

    def markdown_to_excel(
        self,
        markdown_content: str,
        branding: dict[str, Any] | None = None,
    ) -> bytes:
        """Convert Markdown to Excel (.xlsx).

        Args:
            markdown_content: Markdown content
            branding: Branding configuration

        Returns:
            Excel bytes

        Raises:
            ValueError: If openpyxl not available
        """
        if not OPENPYXL_AVAILABLE:
            raise ValueError(
                "openpyxl not available. Install with: pip install openpyxl"
            )

        from openpyxl.utils import get_column_letter
        from openpyxl.styles import Border, Side

        wb = Workbook()

        # Sheet 1: Summary
        ws_summary = wb.active
        ws_summary.title = "Summary"

        # Sheet 2: Controls Detail
        ws_controls = wb.create_sheet("Controls")

        # Style definitions
        title_font = Font(bold=True, size=16, color="2563EB")
        header_fill = PatternFill(start_color="2563EB", end_color="2563EB", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        section_font = Font(bold=True, size=12, color="374151")
        thin_border = Border(
            left=Side(style='thin', color='D1D5DB'),
            right=Side(style='thin', color='D1D5DB'),
            top=Side(style='thin', color='D1D5DB'),
            bottom=Side(style='thin', color='D1D5DB')
        )

        # Parse the markdown
        lines = markdown_content.split("\n")
        summary_row = 1

        # Extract title
        for line in lines:
            if line.startswith("# "):
                title = self._strip_markdown(line[2:])
                ws_summary.merge_cells("A1:E1")
                cell = ws_summary.cell(row=1, column=1, value=title)
                cell.font = title_font
                cell.alignment = Alignment(horizontal="center", vertical="center")
                ws_summary.row_dimensions[1].height = 30
                summary_row = 3
                break

        # Extract summary info and resources
        current_section = ""
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue

            if line_stripped.startswith("## "):
                current_section = self._strip_markdown(line_stripped[3:])
                ws_summary.cell(row=summary_row, column=1, value=current_section)
                ws_summary.cell(row=summary_row, column=1).font = section_font
                summary_row += 1
            elif line_stripped.startswith("**Generated**:") or line_stripped.startswith("*Generated*:"):
                ws_summary.cell(row=summary_row, column=1, value="Generated")
                ws_summary.cell(row=summary_row, column=2, value=self._strip_markdown(line_stripped.split(":", 1)[1].strip()))
                summary_row += 1
            elif (line_stripped.startswith("- ") or line_stripped.startswith("* ")) and current_section:
                item = self._strip_markdown(line_stripped[2:])
                ws_summary.cell(row=summary_row, column=1, value=f"  • {item}")
                summary_row += 1
            elif "Total Controls" in line_stripped or "CRITICAL" in line_stripped or "HIGH" in line_stripped:
                # Skip these - they go in the controls sheet
                pass

        # Controls sheet - structured table
        controls = self._parse_compliance_data(markdown_content)

        # Header row
        headers = ["Resource", "Control ID", "Severity", "Description", "Details"]
        for col, header in enumerate(headers, 1):
            cell = ws_controls.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_border
        ws_controls.row_dimensions[1].height = 25

        # Severity colors
        severity_fills = {
            "CRITICAL": PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid"),
            "HIGH": PatternFill(start_color="FFEDD5", end_color="FFEDD5", fill_type="solid"),
            "MEDIUM": PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid"),
            "LOW": PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid"),
        }

        # Data rows
        for row_idx, control in enumerate(controls, 2):
            ws_controls.cell(row=row_idx, column=1, value=control.get("resource", ""))
            ws_controls.cell(row=row_idx, column=2, value=control.get("control_id", ""))

            severity = control.get("severity", "").upper()
            severity_cell = ws_controls.cell(row=row_idx, column=3, value=severity)
            if severity in severity_fills:
                severity_cell.fill = severity_fills[severity]
            severity_cell.alignment = Alignment(horizontal="center")

            ws_controls.cell(row=row_idx, column=4, value=control.get("description", ""))
            ws_controls.cell(row=row_idx, column=5, value=control.get("details", "")[:500])  # Limit details

            # Apply borders
            for col in range(1, 6):
                ws_controls.cell(row=row_idx, column=col).border = thin_border

        # Auto-size columns
        for ws in [ws_summary, ws_controls]:
            for col_idx in range(1, ws.max_column + 1):
                max_length = 0
                column_letter = get_column_letter(col_idx)
                for row_idx in range(1, min(ws.max_row + 1, 100)):  # Sample first 100 rows
                    cell = ws.cell(row=row_idx, column=col_idx)
                    try:
                        if hasattr(cell, 'value') and cell.value:
                            max_length = max(max_length, min(len(str(cell.value)), 60))
                    except Exception:
                        pass
                ws.column_dimensions[column_letter].width = max(max_length + 2, 12)

        # Set specific widths for controls sheet
        ws_controls.column_dimensions['A'].width = 20
        ws_controls.column_dimensions['B'].width = 35
        ws_controls.column_dimensions['C'].width = 12
        ws_controls.column_dimensions['D'].width = 50
        ws_controls.column_dimensions['E'].width = 80

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def markdown_to_csv(
        self,
        markdown_content: str,
    ) -> bytes:
        """Convert Markdown to CSV with structured compliance data.

        Args:
            markdown_content: Markdown content

        Returns:
            CSV bytes with proper tabular structure
        """
        from io import StringIO
        output = StringIO()
        writer = csv.writer(output)

        # Parse the markdown into structured data
        controls = self._parse_compliance_data(markdown_content)

        if controls:
            # Write header
            writer.writerow(["Resource", "Control ID", "Severity", "Description", "Details"])

            # Write control data
            for control in controls:
                writer.writerow([
                    control.get("resource", ""),
                    control.get("control_id", ""),
                    control.get("severity", ""),
                    control.get("description", ""),
                    control.get("details", "")[:500]  # Limit details length
                ])
        else:
            # Fallback: If no controls parsed, extract tables and basic structure
            lines = markdown_content.split("\n")
            in_table = False
            table_header_written = False

            for line in lines:
                line = line.strip()
                if not line or line == "---":
                    continue

                # Handle markdown tables
                if line.startswith("|") and "|" in line[1:]:
                    if "---" in line:
                        continue
                    parts = [self._strip_markdown(p.strip()) for p in line.split("|")[1:-1]]
                    writer.writerow(parts)
                    in_table = True
                    table_header_written = True
                elif in_table and not line.startswith("|"):
                    in_table = False
                # Handle headings and content
                elif not in_table:
                    if line.startswith("#"):
                        # Strip heading markers and write as single cell
                        clean_line = self._strip_markdown(re.sub(r'^#+\s*', '', line))
                        if clean_line:
                            writer.writerow([clean_line])
                    elif line.startswith("- ") or line.startswith("* "):
                        clean_line = self._strip_markdown(line[2:])
                        writer.writerow([f"• {clean_line}"])
                    elif not line.startswith("**") or ":" in line:
                        # Regular content or key-value pairs
                        clean_line = self._strip_markdown(line)
                        if ":" in clean_line and not clean_line.startswith("http"):
                            parts = clean_line.split(":", 1)
                            writer.writerow([parts[0].strip(), parts[1].strip() if len(parts) > 1 else ""])
                        elif clean_line:
                            writer.writerow([clean_line])

        # Return as UTF-8 encoded bytes with BOM for Excel compatibility
        return ('\ufeff' + output.getvalue()).encode('utf-8')

    def markdown_to_xml(
        self,
        markdown_content: str,
        branding: dict[str, Any] | None = None,
    ) -> bytes:
        """Convert Markdown to XML.

        Args:
            markdown_content: Markdown content
            branding: Branding configuration

        Returns:
            XML bytes
        """
        import xml.etree.ElementTree as ET
        from xml.dom import minidom

        root = ET.Element("report")

        if branding:
            if branding.get("company_name"):
                root.set("company", branding["company_name"])

        lines = markdown_content.split("\n")
        current_section = root

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith("# "):
                section = ET.SubElement(root, "section")
                section.set("level", "1")
                section.set("title", line[2:])
                current_section = section
            elif line.startswith("## "):
                subsection = ET.SubElement(current_section, "subsection")
                subsection.set("level", "2")
                subsection.set("title", line[3:])
            elif line.startswith("### "):
                subsubsection = ET.SubElement(current_section, "subsection")
                subsubsection.set("level", "3")
                subsubsection.set("title", line[4:])
            elif line.startswith("|") and "|" in line[1:]:
                if "---" not in line:
                    row = ET.SubElement(current_section, "row")
                    parts = [p.strip() for p in line.split("|")[1:-1]]
                    for i, part in enumerate(parts):
                        cell = ET.SubElement(row, "cell")
                        cell.set("index", str(i))
                        cell.text = part
            elif line.startswith("- ") or line.startswith("* "):
                item = ET.SubElement(current_section, "item")
                item.text = line[2:]
            else:
                para = ET.SubElement(current_section, "paragraph")
                para.text = line

        xml_str = ET.tostring(root, encoding="unicode")
        pretty_xml = minidom.parseString(xml_str).toprettyxml(indent="  ")
        return pretty_xml.encode("utf-8")

    def convert_format(
        self,
        content: str | bytes,
        source_format: str,
        target_format: str,
        styles: dict[str, Any] | None = None,
        branding: dict[str, Any] | None = None,
    ) -> bytes:
        """Convert content from one format to another.

        Args:
            content: Source content (string for text formats, bytes for binary)
            source_format: Source format (markdown, html, pdf, docx)
            target_format: Target format (markdown, html, pdf, docx, xlsx, csv)
            styles: CSS styles dictionary
            branding: Branding configuration

        Returns:
            Converted content as bytes

        Raises:
            ValueError: If conversion not supported or fails
        """
        if source_format == target_format:
            if isinstance(content, bytes):
                return content
            return content.encode("utf-8")

        markdown_content = None

        if source_format == "markdown":
            markdown_content = content if isinstance(content, str) else content.decode("utf-8")
        elif source_format == "html":
            html_content = content if isinstance(content, str) else content.decode("utf-8")
            markdown_content = self.html_to_markdown(html_content)
        elif source_format in ("pdf", "docx"):
            raise ValueError(f"Conversion from {source_format} to {target_format} not supported")

        if not markdown_content:
            raise ValueError(f"Unable to extract markdown from {source_format}")

        if target_format == "markdown":
            return markdown_content.encode("utf-8")
        elif target_format == "html":
            html_content = self.markdown_to_html(markdown_content, styles)
            return html_content.encode("utf-8")
        elif target_format == "pdf":
            return self.markdown_to_pdf(markdown_content, styles, branding)
        elif target_format == "docx":
            return self.markdown_to_docx(markdown_content, branding)
        elif target_format == "xlsx":
            return self.markdown_to_excel(markdown_content, branding)
        elif target_format == "csv":
            return self.markdown_to_csv(markdown_content)
        elif target_format == "xml":
            return self.markdown_to_xml(markdown_content, branding)
        else:
            raise ValueError(f"Unsupported target format: {target_format}")

