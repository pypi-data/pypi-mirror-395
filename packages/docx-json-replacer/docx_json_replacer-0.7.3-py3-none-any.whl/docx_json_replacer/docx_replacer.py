"""
DocxReplacer with table support for docx-json-replacer
"""
import json
import re
from typing import Dict, Any, Tuple, List
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from .utility.html_parse import clean_html_content
    from .table_handler import TableHandler
    from .formatting_handler import FormattingHandler
except ImportError:
    from utility.html_parse import clean_html_content
    from table_handler import TableHandler
    from formatting_handler import FormattingHandler


class DocxReplacer:
    """Replace placeholders in DOCX files with JSON data, including table support"""

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.table_handler = TableHandler()
        self.formatting_handler = FormattingHandler()
        self.table_placeholders = {}
        self.multi_table_placeholders = {}  # For multiple tables
        self.formatted_content_placeholders = {}  # For formatted content (titles, headings, bullets)
        self._value_cache = {}
        self._table_check_cache = {}

    def replace_from_json(self, json_data: Dict[str, Any]) -> None:
        """Replace placeholders in paragraphs AND tables"""

        # Pre-compile patterns
        patterns = self._compile_patterns(json_data)

        # Pre-process values
        processed_values = self._preprocess_values(json_data)

        # Process regular paragraphs
        self._process_paragraphs(patterns, processed_values)

        # Process table cells
        self._process_tables(patterns, processed_values)

        # Insert dynamic tables for table data
        self._batch_insert_tables()

        # Insert multiple tables for multi-table data
        self._batch_insert_multi_tables()

        # Insert formatted content (titles, headings, bullets)
        self._batch_insert_formatted_content()

    def _compile_patterns(self, json_data: Dict[str, Any]) -> Dict[str, Tuple[re.Pattern, re.Pattern]]:
        """Pre-compile regex patterns for all placeholders"""
        patterns = {}
        for key in json_data.keys():
            escaped_key = re.escape(key)
            pattern = re.compile(r'\{\{' + escaped_key + r'\}\}')
            pattern_spaced = re.compile(r'\{\{ ' + escaped_key + r' \}\}')
            patterns[key] = (pattern, pattern_spaced)
        return patterns

    def _preprocess_values(self, json_data: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
        """Pre-process all values"""
        processed = {}

        for key, value in json_data.items():
            is_multi_table = self.table_handler.is_multi_table_data(value)
            is_table = self._check_is_table(value)
            is_formatted = self.formatting_handler.has_formatting_tags(value)

            if is_multi_table:
                processed[key] = {
                    'is_table': False,  # Don't treat as single table
                    'is_multi_table': True,
                    'is_formatted': False,
                    'original': value,
                    'processed': None
                }
            elif is_table:
                processed[key] = {
                    'is_table': True,
                    'is_multi_table': False,
                    'is_formatted': False,
                    'original': value,
                    'processed': None
                }
            elif is_formatted:
                processed[key] = {
                    'is_table': False,
                    'is_multi_table': False,
                    'is_formatted': True,
                    'original': value,
                    'processed': None
                }
            else:
                cleaned = clean_html_content(value) if isinstance(value, str) else str(value)
                # Replace [dx-tab] with actual tab characters
                cleaned = self.formatting_handler.replace_inline_tabs(cleaned)
                processed[key] = {
                    'is_table': False,
                    'is_multi_table': False,
                    'is_formatted': False,
                    'original': value,
                    'processed': cleaned
                }

        return processed

    def _check_is_table(self, value: Any) -> bool:
        """Check if value is table data with caching"""
        value_id = id(value)
        if value_id in self._table_check_cache:
            return self._table_check_cache[value_id]

        result = (self.table_handler.is_table_data(value) or
                  (isinstance(value, str) and '<table' in value.lower()))

        self._table_check_cache[value_id] = result
        return result

    def _process_paragraphs(self, patterns: Dict, processed_values: Dict) -> None:
        """Process all document paragraphs while preserving formatting"""
        for paragraph in self.doc.paragraphs:
            text = paragraph.text
            if not text or '{{' not in text:
                continue

            # Check if we need to handle special placeholders
            table_placeholder_found = False
            multi_table_placeholder_found = False
            formatted_placeholder_found = False

            for key, (pattern, pattern_spaced) in patterns.items():
                if pattern.search(text) or pattern_spaced.search(text):
                    value_data = processed_values[key]

                    if value_data.get('is_multi_table'):
                        # Store for multiple table insertion
                        self.multi_table_placeholders[paragraph] = (key, value_data['original'])
                        multi_table_placeholder_found = True
                        break
                    elif value_data.get('is_table'):
                        # Store for table insertion
                        self.table_placeholders[paragraph] = (key, value_data['original'])
                        table_placeholder_found = True
                        break
                    elif value_data.get('is_formatted'):
                        # Capture formatting BEFORE clearing the paragraph
                        formatting_info = self._capture_paragraph_formatting(paragraph)
                        # Store for formatted content insertion (titles, headings, bullets)
                        self.formatted_content_placeholders[paragraph] = (key, value_data['original'], formatting_info)
                        formatted_placeholder_found = True
                        break

            # If it's a special placeholder, clear the paragraph
            if table_placeholder_found or multi_table_placeholder_found or formatted_placeholder_found:
                paragraph.text = ''
                continue

            # Process runs to preserve formatting for non-table replacements
            self._process_paragraph_runs(paragraph, patterns, processed_values)

    def _capture_paragraph_formatting(self, paragraph) -> Dict[str, Any]:
        """Capture formatting info from a paragraph's runs and paragraph format"""
        formatting = {
            'font_size': None,
            'font_name': None,
            'bold': None,
            'italic': None,
            'underline': None,
            'font_color': None,
            'left_indent': None,
            'first_line_indent': None
        }

        # Try to get formatting from runs in the paragraph
        if paragraph.runs:
            first_run = paragraph.runs[0]
            formatting['font_size'] = first_run.font.size
            formatting['font_name'] = first_run.font.name
            formatting['bold'] = first_run.bold
            formatting['italic'] = first_run.italic
            formatting['underline'] = first_run.underline
            if first_run.font.color.rgb:
                formatting['font_color'] = first_run.font.color.rgb

        # Capture paragraph indentation
        formatting['left_indent'] = paragraph.paragraph_format.left_indent
        formatting['first_line_indent'] = paragraph.paragraph_format.first_line_indent

        return formatting

    def _process_paragraph_runs(self, paragraph, patterns: Dict, processed_values: Dict) -> None:
        """Process runs within a paragraph to preserve formatting"""
        # Collect all runs and their text
        runs = paragraph.runs
        if not runs:
            # If there are no runs but there is text, create one run with the text
            if paragraph.text:
                text = paragraph.text
                paragraph.text = ''
                run = paragraph.add_run(text)
                runs = [run]
            else:
                return

        # First, check if placeholders are split across runs
        # This can happen when Word splits text into multiple runs
        full_text = paragraph.text
        if '{{' in full_text:
            # Check if any complete placeholders exist in the full text
            has_complete_placeholders = False
            for key, (pattern, pattern_spaced) in patterns.items():
                if pattern.search(full_text) or pattern_spaced.search(full_text):
                    has_complete_placeholders = True
                    break

            # If we have complete placeholders and multiple runs, merge and re-split
            if has_complete_placeholders and len(runs) > 1:
                # Save the formatting from the first run that has a placeholder start
                formatting_run = None
                for run in runs:
                    if '{{' in run.text:
                        formatting_run = run
                        break
                if not formatting_run:
                    formatting_run = runs[0]

                # Store the original formatting
                original_bold = formatting_run.bold
                original_italic = formatting_run.italic
                original_underline = formatting_run.underline
                original_font_size = formatting_run.font.size
                original_font_name = formatting_run.font.name
                original_font_color = None
                if formatting_run.font.color.rgb:
                    original_font_color = formatting_run.font.color.rgb

                # Process the full text
                new_text = full_text
                for key, (pattern, pattern_spaced) in patterns.items():
                    if key not in processed_values:
                        continue
                    value_data = processed_values[key]
                    # Skip table, multi-table, and formatted content (they're handled separately)
                    if not value_data.get('is_table') and not value_data.get('is_multi_table') and not value_data.get('is_formatted'):
                        # Regular text replacement
                        replacement = value_data['processed']
                        new_text = pattern.sub(replacement, new_text)
                        new_text = pattern_spaced.sub(replacement, new_text)

                # Remove any remaining unmatched placeholders
                unmatched_pattern = re.compile(r'\{\{[^}]+\}\}')
                new_text = unmatched_pattern.sub('', new_text)

                # Clear existing runs and create a new one with the replaced text
                paragraph.clear()
                new_run = paragraph.add_run(new_text)

                # Apply the preserved formatting
                if original_bold is not None:
                    new_run.bold = original_bold
                if original_italic is not None:
                    new_run.italic = original_italic
                if original_underline is not None:
                    new_run.underline = original_underline
                if original_font_size is not None:
                    new_run.font.size = original_font_size
                if original_font_name is not None:
                    new_run.font.name = original_font_name
                if original_font_color is not None:
                    new_run.font.color.rgb = original_font_color
                return

        # Process each run individually (for cases where placeholders are within single runs)
        for run in runs:
            if not run.text:
                continue

            original_text = run.text
            new_text = original_text
            modified = False

            # Store the original formatting
            original_bold = run.bold
            original_italic = run.italic
            original_underline = run.underline
            original_font_size = run.font.size
            original_font_name = run.font.name
            original_font_color = None
            if run.font.color.rgb:
                original_font_color = run.font.color.rgb

            # Replace placeholders in this run
            for key, (pattern, pattern_spaced) in patterns.items():
                if key not in processed_values:
                    continue

                if pattern.search(new_text) or pattern_spaced.search(new_text):
                    value_data = processed_values[key]

                    # Skip table, multi-table, and formatted content (they're handled separately)
                    if not value_data.get('is_table') and not value_data.get('is_multi_table') and not value_data.get('is_formatted'):
                        # Regular text replacement
                        replacement = value_data['processed']
                        new_text = pattern.sub(replacement, new_text)
                        new_text = pattern_spaced.sub(replacement, new_text)
                        modified = True

            # Remove any remaining unmatched placeholders
            unmatched_pattern = re.compile(r'\{\{[^}]+\}\}')
            if unmatched_pattern.search(new_text):
                new_text = unmatched_pattern.sub('', new_text)
                modified = True

            if modified:
                # Update the run text while preserving formatting
                run.text = new_text

                # Restore the original formatting
                if original_bold is not None:
                    run.bold = original_bold
                if original_italic is not None:
                    run.italic = original_italic
                if original_underline is not None:
                    run.underline = original_underline
                if original_font_size is not None:
                    run.font.size = original_font_size
                if original_font_name is not None:
                    run.font.name = original_font_name
                if original_font_color is not None:
                    run.font.color.rgb = original_font_color

    def _process_tables(self, patterns: Dict, processed_values: Dict) -> None:
        """Process all table cells in the document while preserving formatting"""
        for table in self.doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Process each paragraph in the cell
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if not text or '{{' not in text:
                            continue

                        # Check if we have table placeholders (which can't be inserted in cells)
                        has_table_placeholder = False
                        for key, (pattern, pattern_spaced) in patterns.items():
                            if key not in processed_values:
                                continue
                            if pattern.search(text) or pattern_spaced.search(text):
                                value_data = processed_values[key]
                                if value_data.get('is_table') or value_data.get('is_multi_table'):
                                    has_table_placeholder = True
                                    # Replace with placeholder text
                                    paragraph.text = '[Table data - see document]'
                                    break

                        # If no table placeholder, process runs to preserve formatting
                        if not has_table_placeholder:
                            # Use the same run processing method we use for regular paragraphs
                            self._process_paragraph_runs(paragraph, patterns, processed_values)

    def _batch_insert_tables(self) -> None:
        """Insert all tables in batch"""
        if not self.table_placeholders:
            return

        processed_tables = {}
        split_tables = {}  # For tables that should be split into multiple tables

        for paragraph, (key, value) in self.table_placeholders.items():
            table_data = self.table_handler.process_table_data(value)

            # Check if this should be split into multiple tables
            if table_data.get('should_split'):
                split_tables[paragraph] = table_data.get('split_tables', [])
            elif table_data.get('rows'):
                processed_tables[paragraph] = table_data

        # Insert regular tables
        for paragraph, table_data in processed_tables.items():
            self._insert_table_fast(paragraph, table_data)

        # Insert split tables (multiple tables for each placeholder)
        for paragraph, tables_list in split_tables.items():
            self._insert_multiple_tables(paragraph, tables_list)

    def _insert_table_fast(self, paragraph, table_data: Dict[str, Any]) -> None:
        """Fast table insertion"""
        rows = table_data['rows']
        num_rows = len(rows)
        num_cols = len(rows[0]['cells']) if rows and 'cells' in rows[0] else 0

        if num_rows == 0 or num_cols == 0:
            return

        parent = paragraph._element.getparent()
        index = parent.index(paragraph._element)

        # Set paragraph properties to keep with next (table)
        self._set_paragraph_keep_with_next(paragraph)

        table = self.doc.add_table(rows=num_rows, cols=num_cols)

        # Prevent page break before table
        table.allow_autofit = False
        table.autofit = False  # Disable autofit to preserve column widths

        # Set table properties to prevent page breaks
        self._set_table_no_page_break(table)

        # Set table grid column widths from first row's cell styles
        if rows and rows[0]:
            first_row = rows[0]

            # Extract cell styles from the first row's cells
            cell_styles = []
            if 'cells' in first_row:
                for cell in first_row['cells']:
                    if isinstance(cell, dict) and 'style' in cell:
                        cell_styles.append(cell['style'])
                    else:
                        cell_styles.append({})

            if cell_styles:
                self._set_table_grid_widths(table, cell_styles)

        # Try to apply Table Grid style, fall back to default if not available
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If 'Table Grid' style doesn't exist, try alternative styles or use default
            try:
                # Try common table style alternatives
                for style_name in ['TableGrid', 'Light Grid', 'Light List', 'Normal Table']:
                    try:
                        table.style = style_name
                        break
                    except KeyError:
                        continue
            except:
                # If all else fails, leave the default style
                pass

        # Collect column widths from first row if available
        column_widths = []
        if rows and rows[0].get('cell_styles'):
            for cell_style in rows[0]['cell_styles']:
                if 'width' in cell_style:
                    column_widths.append(cell_style['width'])

        table_rows = table.rows
        for row_idx, row_data in enumerate(rows):
            cells = row_data.get('cells', [])
            default_row_style = row_data.get('style', {})
            row_cells = table_rows[row_idx].cells

            # Prevent page break in the first row
            if row_idx == 0:
                self._prevent_row_page_break(table_rows[row_idx])

                # Set column widths based on first row
                if column_widths:
                    from docx.shared import Cm, Inches, Pt
                    for col_idx, width in enumerate(column_widths):
                        if col_idx < len(table.columns):
                            try:
                                if isinstance(width, str):
                                    if width.endswith('cm'):
                                        table.columns[col_idx].width = Cm(float(width[:-2]))
                                    elif width.endswith('in'):
                                        table.columns[col_idx].width = Inches(float(width[:-2]))
                                    elif width.endswith('pt'):
                                        table.columns[col_idx].width = Pt(float(width[:-2]))
                                    else:
                                        # Assume cm if no unit
                                        table.columns[col_idx].width = Cm(float(width))
                                else:
                                    table.columns[col_idx].width = Cm(float(width))
                            except Exception as e:
                                print(f"Error setting column width: {e}")

            for col_idx, cell_data in enumerate(cells):
                if col_idx < len(row_cells):
                    cell = row_cells[col_idx]

                    # Handle both old format (string) and new format (dict with content and style)
                    if isinstance(cell_data, dict):
                        cell_text = cell_data.get('content', '')
                        cell_style = cell_data.get('style', default_row_style)
                    else:
                        cell_text = str(cell_data)
                        cell_style = default_row_style

                    # Process HTML content in cell text
                    self._set_cell_content_with_formatting(cell, str(cell_text))

                    # Apply cell-specific style or row style
                    if cell_style:
                        self._apply_cell_style_fast(cell, cell_style)

        # Clear the placeholder paragraph to avoid spacing issues
        if paragraph.text.strip() == '':
            # If paragraph is empty after placeholder removal, delete it
            parent.remove(paragraph._element)
            parent.insert(index, table._element)
        else:
            # Otherwise insert table after the paragraph
            parent.insert(index + 1, table._element)

    def _insert_multiple_tables(self, paragraph, tables_list: List[Dict[str, Any]]) -> None:
        """
        Insert multiple tables for a single placeholder (for split data)

        Args:
            paragraph: The placeholder paragraph
            tables_list: List of table data structures to insert
        """
        if not tables_list:
            return

        parent = paragraph._element.getparent()
        base_index = parent.index(paragraph._element)

        # Insert each table with spacing between them
        for i, table_data in enumerate(tables_list):
            if table_data.get('rows'):
                # Add spacing paragraph between tables (except before first)
                if i > 0:
                    spacing_para = self.doc.add_paragraph()
                    parent.insert(base_index + (i * 2), spacing_para._element)

                # Insert the table
                self._insert_table_at_index(parent, base_index + (i * 2) + 1, table_data)

        # Remove the placeholder paragraph if it's empty
        if paragraph.text.strip() == '':
            parent.remove(paragraph._element)

    def _batch_insert_multi_tables(self) -> None:
        """Insert multiple tables for each multi-table placeholder"""
        if not self.multi_table_placeholders:
            return

        for paragraph, (key, value) in self.multi_table_placeholders.items():
            tables_data = self.table_handler.process_multi_table_data(value)

            if tables_data:
                parent = paragraph._element.getparent()
                base_index = parent.index(paragraph._element)

                # Insert each table with optional spacing between them
                for i, table_data in enumerate(tables_data):
                    if table_data.get('rows'):
                        # Add spacing paragraph between tables (except before first)
                        if i > 0:
                            spacing_para = self.doc.add_paragraph()
                            parent.insert(base_index + (i * 2), spacing_para._element)

                        # Insert the table
                        self._insert_table_at_index(parent, base_index + (i * 2) + 1, table_data)

    def _batch_insert_formatted_content(self) -> None:
        """Insert formatted paragraphs (titles, headings, bullets) in batch"""
        if not self.formatted_content_placeholders:
            return

        for paragraph, placeholder_data in self.formatted_content_placeholders.items():
            # Unpack the stored data (key, value, formatting_info)
            if len(placeholder_data) == 3:
                _, value, formatting_info = placeholder_data
            else:
                # Fallback for old format without formatting info
                _, value = placeholder_data
                formatting_info = {}

            # Parse the formatted content
            content_specs = self.formatting_handler.parse_formatted_content(value)

            if not content_specs:
                continue

            parent = paragraph._element.getparent()
            base_index = parent.index(paragraph._element)

            # Get the captured formatting from the placeholder
            original_font_size = formatting_info.get('font_size')
            original_font_name = formatting_info.get('font_name')
            original_bold = formatting_info.get('bold')
            original_italic = formatting_info.get('italic')
            original_underline = formatting_info.get('underline')
            original_font_color = formatting_info.get('font_color')
            original_left_indent = formatting_info.get('left_indent')
            original_first_line_indent = formatting_info.get('first_line_indent')

            # Track list numbering state for consecutive bullets of same type
            current_list_type = None
            list_num_id = None

            # Insert each formatted paragraph
            for i, spec in enumerate(content_specs):
                style_name = spec['style']
                text = spec['text']
                is_bullet = spec.get('is_bullet', False)
                bullet_type = spec.get('bullet_type')

                # Create new paragraph
                try:
                    new_para = self.doc.add_paragraph(style=style_name)
                except KeyError:
                    # Style doesn't exist, fall back to Normal
                    new_para = self.doc.add_paragraph(style='Normal')

                # Add the text as a run and apply original formatting
                run = new_para.add_run(text)

                # Apply the original formatting from the placeholder
                if original_font_size is not None:
                    run.font.size = original_font_size
                if original_font_name is not None:
                    run.font.name = original_font_name
                if original_bold is not None:
                    run.bold = original_bold
                if original_italic is not None:
                    run.italic = original_italic
                if original_underline is not None:
                    run.underline = original_underline
                if original_font_color is not None:
                    run.font.color.rgb = original_font_color

                # Handle custom numbering for alpha/roman bullets
                if is_bullet and bullet_type:
                    # Check if we need to start a new list or continue existing
                    if bullet_type != current_list_type:
                        current_list_type = bullet_type
                        list_num_id = self._create_custom_numbering(
                            bullet_type, original_font_size, original_left_indent
                        )

                    if list_num_id is not None:
                        self._apply_numbering_to_paragraph(new_para, list_num_id)

                # Insert at correct position
                parent.insert(base_index + i, new_para._element)

            # Remove the placeholder paragraph
            parent.remove(paragraph._element)

    def _create_custom_numbering(self, bullet_type: str, font_size=None, left_indent=None) -> int:
        """Create custom numbering definition for alpha/roman lists"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Get the numbering format
        num_fmt = self.formatting_handler.get_numbering_format(bullet_type)

        if num_fmt is None:
            return None  # Standard bullet, no custom numbering needed

        # Access the numbering part of the document
        numbering_part = self.doc.part.numbering_part
        if numbering_part is None:
            # Create numbering part if it doesn't exist
            from docx.parts.numbering import NumberingPart
            numbering_part = NumberingPart.new()
            self.doc.part.relate_to(numbering_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering')

        numbering = numbering_part.numbering_definitions._numbering

        # Find max abstractNumId and numId
        abstract_nums = numbering.findall(qn('w:abstractNum'))
        max_abstract_id = max([int(a.get(qn('w:abstractNumId'))) for a in abstract_nums], default=-1)

        nums = numbering.findall(qn('w:num'))
        max_num_id = max([int(n.get(qn('w:numId'))) for n in nums], default=0)

        new_abstract_id = max_abstract_id + 1
        new_num_id = max_num_id + 1

        # Create abstractNum element
        abstract_num = OxmlElement('w:abstractNum')
        abstract_num.set(qn('w:abstractNumId'), str(new_abstract_id))

        # Create level 0 (single level list)
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')

        # Start value
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        # Number format
        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), num_fmt)
        lvl.append(numFmt)

        # Level text (e.g., "a)" or "1.")
        lvlText = OxmlElement('w:lvlText')
        if num_fmt in ['lowerLetter', 'upperLetter']:
            lvlText.set(qn('w:val'), '%1)')
        elif num_fmt in ['lowerRoman', 'upperRoman']:
            lvlText.set(qn('w:val'), '%1)')
        else:
            lvlText.set(qn('w:val'), '%1.')
        lvl.append(lvlText)

        # Level justification
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)

        # Paragraph properties (indentation)
        # Use original paragraph indent or default to document standard
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')

        # Calculate left indent in twips (1/20 of a point)
        # left_indent from python-docx is in EMUs (914400 per inch, or 635 per twip)
        if left_indent is not None:
            # Convert EMUs to twips: EMUs / 635 = twips
            left_twips = int(left_indent / 635) + 360  # Add hanging space for bullet
        else:
            # Default: match document's standard lowerLetter indent (981 twips left, 360 hanging)
            left_twips = 981

        hanging_twips = 360  # Space for bullet marker (about 0.25 inch)

        ind.set(qn('w:left'), str(left_twips))
        ind.set(qn('w:hanging'), str(hanging_twips))
        pPr.append(ind)
        lvl.append(pPr)

        # Run properties for the bullet/number marker (font size)
        if font_size is not None:
            rPr = OxmlElement('w:rPr')
            sz = OxmlElement('w:sz')
            # font_size is in EMUs (12700 per point), convert to half-points
            # 1 point = 12700 EMU, w:sz uses half-points (1pt = 2 half-points)
            size_in_half_points = int((font_size / 12700) * 2)
            sz.set(qn('w:val'), str(size_in_half_points))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(size_in_half_points))
            rPr.append(szCs)
            lvl.append(rPr)

        abstract_num.append(lvl)

        # Insert abstractNum before num elements
        first_num = numbering.find(qn('w:num'))
        if first_num is not None:
            first_num.addprevious(abstract_num)
        else:
            numbering.append(abstract_num)

        # Create num element that references the abstractNum
        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(new_num_id))
        abstractNumId = OxmlElement('w:abstractNumId')
        abstractNumId.set(qn('w:val'), str(new_abstract_id))
        num.append(abstractNumId)
        numbering.append(num)

        return new_num_id

    def _apply_numbering_to_paragraph(self, paragraph, num_id: int) -> None:
        """Apply numbering to a paragraph"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p = paragraph._element
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p.insert(0, pPr)

        # Remove existing numPr if any
        existing_numPr = pPr.find(qn('w:numPr'))
        if existing_numPr is not None:
            pPr.remove(existing_numPr)

        # Create numPr element
        numPr = OxmlElement('w:numPr')

        # Add ilvl (indent level)
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numPr.append(ilvl)

        # Add numId
        numId_elem = OxmlElement('w:numId')
        numId_elem.set(qn('w:val'), str(num_id))
        numPr.append(numId_elem)

        pPr.append(numPr)

    def _insert_table_at_index(self, parent, index: int, table_data: Dict[str, Any]) -> None:
        """Insert a table at a specific index in the parent element"""
        rows = table_data['rows']
        num_rows = len(rows)
        num_cols = len(rows[0]['cells']) if rows and 'cells' in rows[0] else 0

        if num_rows == 0 or num_cols == 0:
            return

        table = self.doc.add_table(rows=num_rows, cols=num_cols)

        # Prevent page break before table
        table.allow_autofit = False
        table.autofit = False  # Disable autofit to preserve column widths

        # Set table properties to prevent page breaks
        self._set_table_no_page_break(table)

        # Set table grid column widths from first row's cell styles
        if rows and rows[0]:
            first_row = rows[0]

            # Extract cell styles from the first row's cells
            cell_styles = []
            if 'cells' in first_row:
                for cell in first_row['cells']:
                    if isinstance(cell, dict) and 'style' in cell:
                        cell_styles.append(cell['style'])
                    else:
                        cell_styles.append({})

            if cell_styles:
                self._set_table_grid_widths(table, cell_styles)

        # Try to apply Table Grid style
        try:
            table.style = 'Table Grid'
        except KeyError:
            try:
                for style_name in ['TableGrid', 'Light Grid', 'Light List', 'Normal Table']:
                    try:
                        table.style = style_name
                        break
                    except KeyError:
                        continue
            except:
                pass

        # Collect column widths from first row if available
        column_widths = []
        if rows and rows[0].get('cell_styles'):
            for cell_style in rows[0]['cell_styles']:
                if 'width' in cell_style:
                    column_widths.append(cell_style['width'])

        table_rows = table.rows
        for row_idx, row_data in enumerate(rows):
            cells = row_data.get('cells', [])
            default_row_style = row_data.get('style', {})
            row_cells = table_rows[row_idx].cells

            # Prevent page break in the first row
            if row_idx == 0:
                self._prevent_row_page_break(table_rows[row_idx])

                # Set column widths based on first row
                if column_widths:
                    from docx.shared import Cm, Inches, Pt
                    for col_idx, width in enumerate(column_widths):
                        if col_idx < len(table.columns):
                            try:
                                if isinstance(width, str):
                                    if width.endswith('cm'):
                                        table.columns[col_idx].width = Cm(float(width[:-2]))
                                    elif width.endswith('in'):
                                        table.columns[col_idx].width = Inches(float(width[:-2]))
                                    elif width.endswith('pt'):
                                        table.columns[col_idx].width = Pt(float(width[:-2]))
                                    else:
                                        # Assume cm if no unit
                                        table.columns[col_idx].width = Cm(float(width))
                                else:
                                    table.columns[col_idx].width = Cm(float(width))
                            except Exception as e:
                                print(f"Error setting column width: {e}")

            for col_idx, cell_data in enumerate(cells):
                if col_idx < len(row_cells):
                    cell = row_cells[col_idx]

                    # Handle both old format (string) and new format (dict with content and style)
                    if isinstance(cell_data, dict):
                        cell_text = cell_data.get('content', '')
                        cell_style = cell_data.get('style', default_row_style)
                    else:
                        cell_text = str(cell_data)
                        cell_style = default_row_style

                    # Process HTML content in cell text
                    self._set_cell_content_with_formatting(cell, str(cell_text))

                    # Apply cell-specific style or row style
                    if cell_style:
                        self._apply_cell_style_fast(cell, cell_style)

        parent.insert(index, table._element)

    def _set_cell_content_with_formatting(self, cell, content: str) -> None:
        """Set cell content with HTML formatting support"""
        import re

        # Clear existing content
        cell.text = ""
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

        # Process the content to handle HTML tags
        parts = self._parse_html_for_cell(content)

        for part in parts:
            run = paragraph.add_run(part['text'])
            if part.get('bold'):
                run.bold = True
            if part.get('italic'):
                run.italic = True
            if part.get('underline'):
                run.underline = True

    def _parse_html_for_cell(self, html_text: str) -> list:
        """Parse HTML text and return list of formatted text parts"""
        import re

        # Handle line breaks
        html_text = html_text.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')

        # Handle paragraphs
        html_text = re.sub(r'<p[^>]*>', '', html_text)
        html_text = html_text.replace('</p>', '\n')

        # Parse HTML more carefully by processing character by character
        parts = []
        current_text = []
        current_format = {'bold': False, 'italic': False, 'underline': False}
        i = 0

        while i < len(html_text):
            if html_text[i] == '<':
                # Find the end of the tag
                tag_end = html_text.find('>', i)
                if tag_end == -1:
                    # Malformed tag, treat as text
                    current_text.append(html_text[i])
                    i += 1
                    continue

                tag = html_text[i:tag_end+1]
                tag_content = tag[1:-1].lower()

                # Check if we have accumulated text to save
                if current_text:
                    text = ''.join(current_text)
                    if text.strip():
                        part = {'text': text}
                        if current_format['bold']:
                            part['bold'] = True
                        if current_format['italic']:
                            part['italic'] = True
                        if current_format['underline']:
                            part['underline'] = True
                        parts.append(part)
                    current_text = []

                # Handle opening tags
                if tag_content in ['b', 'strong']:
                    # If bold is already active, treat duplicate opening tag as CLOSING tag
                    if current_format['bold']:
                        # Save current bold text
                        if current_text:
                            text = ''.join(current_text)
                            if text.strip():
                                part = {'text': text, 'bold': True}
                                if current_format['italic']:
                                    part['italic'] = True
                                if current_format['underline']:
                                    part['underline'] = True
                                parts.append(part)
                            current_text = []
                        # Turn OFF bold (duplicate opening tag acts as closing)
                        current_format['bold'] = False
                    else:
                        # Turn ON bold
                        current_format['bold'] = True
                elif tag_content in ['i', 'em']:
                    # If italic is already active, treat duplicate opening tag as CLOSING tag
                    if current_format['italic']:
                        # Save current italic text
                        if current_text:
                            text = ''.join(current_text)
                            if text.strip():
                                part = {'text': text}
                                if current_format['bold']:
                                    part['bold'] = True
                                part['italic'] = True
                                if current_format['underline']:
                                    part['underline'] = True
                                parts.append(part)
                            current_text = []
                        # Turn OFF italic
                        current_format['italic'] = False
                    else:
                        # Turn ON italic
                        current_format['italic'] = True
                elif tag_content == 'u':
                    # If underline is already active, treat duplicate opening tag as CLOSING tag
                    if current_format['underline']:
                        # Save current underline text
                        if current_text:
                            text = ''.join(current_text)
                            if text.strip():
                                part = {'text': text}
                                if current_format['bold']:
                                    part['bold'] = True
                                if current_format['italic']:
                                    part['italic'] = True
                                part['underline'] = True
                                parts.append(part)
                            current_text = []
                        # Turn OFF underline
                        current_format['underline'] = False
                    else:
                        # Turn ON underline
                        current_format['underline'] = True

                # Handle closing tags
                elif tag_content in ['/b', '/strong']:
                    current_format['bold'] = False
                elif tag_content in ['/i', '/em']:
                    current_format['italic'] = False
                elif tag_content == '/u':
                    current_format['underline'] = False

                i = tag_end + 1

            else:
                # Regular text
                current_text.append(html_text[i])
                i += 1

        # Add any remaining text
        if current_text:
            text = ''.join(current_text)
            if text.strip():
                part = {'text': text}
                if current_format['bold']:
                    part['bold'] = True
                if current_format['italic']:
                    part['italic'] = True
                if current_format['underline']:
                    part['underline'] = True
                parts.append(part)

        # If no parts were created, just return the cleaned text
        if not parts:
            cleaned_text = re.sub(r'<[^>]+>', '', html_text)
            return [{'text': cleaned_text}]

        return parts

    def _apply_cell_style_fast(self, cell, style: Dict[str, Any]) -> None:
        """Fast cell styling including width, height, borders, font size, and alignment"""
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Apply width if specified
        if width := style.get('width'):
            try:
                if isinstance(width, str):
                    if width.endswith('in'):
                        cell.width = Inches(float(width[:-2]))
                    elif width.endswith('pt'):
                        cell.width = Pt(float(width[:-2]))
                    elif width.endswith('cm'):
                        cell.width = Cm(float(width[:-2]))
                    elif width.endswith('px'):
                        # Convert pixels to points (1px ≈ 0.75pt)
                        cell.width = Pt(float(width[:-2]) * 0.75)
                    elif width.endswith('%'):
                        # For percentage, we'd need table width - for now skip
                        pass
                    else:
                        # Assume inches if no unit
                        cell.width = Inches(float(width))
                else:
                    # Numeric value - assume inches
                    cell.width = Inches(float(width))
            except (ValueError, AttributeError) as e:
                pass  # Ignore invalid width values

        # Apply height (for row height, applied at cell level)
        if height := style.get('height'):
            try:
                # Get the row through the cell
                tc = cell._tc
                tr = tc.getparent()  # Get the table row element

                # Create height property
                # Check if row already has properties
                trPr = tr.find(qn('w:trPr'))
                if trPr is None:
                    trPr = OxmlElement('w:trPr')
                    tr.insert(0, trPr)

                # Remove existing height if present
                for child in trPr:
                    if child.tag.endswith('trHeight'):
                        trPr.remove(child)

                trHeight = OxmlElement('w:trHeight')

                # Convert height value to twips (1/20 of a point)
                if isinstance(height, str):
                    if height.endswith('pt'):
                        height_val = int(float(height[:-2]) * 20)
                    elif height.endswith('in'):
                        height_val = int(float(height[:-2]) * 1440)
                    elif height.endswith('cm'):
                        height_val = int(float(height[:-2]) * 567)
                    elif height.endswith('px'):
                        height_val = int(float(height[:-2]) * 15)  # 1px ≈ 15 twips
                    else:
                        # Assume points if no unit
                        height_val = int(float(height) * 20)
                else:
                    # Numeric value - assume points
                    height_val = int(float(height) * 20)

                trHeight.set(qn('w:val'), str(height_val))
                trHeight.set(qn('w:hRule'), 'auto')  # Use 'auto' to allow flexible height without forcing page breaks
                trPr.append(trHeight)

            except (ValueError, AttributeError) as e:
                pass  # Ignore invalid height values

        # Apply padding/margins if specified
        if padding := style.get('padding'):
            self._set_cell_padding(cell, padding)

        # Apply borders if specified
        if borders := style.get('borders'):
            self._set_cell_borders(cell, borders)

        if bg := style.get('bg'):
            self._set_cell_bg_fast(cell, bg)

        # Apply horizontal text alignment
        # Handle 'align' for horizontal alignment
        if align := style.get('align'):
            align_lower = align.lower()
            for paragraph in cell.paragraphs:
                if align_lower == 'center':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align_lower == 'right':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align_lower == 'left':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif align_lower == 'justify':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # If align is not recognized, don't change it

        # Apply vertical alignment
        # Handle 'valign' for vertical alignment
        if valign := style.get('valign'):
            valign_lower = valign.lower()

            # Get cell's XML properties
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # Remove existing vAlign if present
            for child in list(tcPr):
                if child.tag.endswith('vAlign'):
                    tcPr.remove(child)

            # Create vAlign element with proper value
            vAlign = OxmlElement('w:vAlign')

            # Map valign values to Word's vAlign values
            if valign_lower in ['center', 'middle']:
                vAlign.set(qn('w:val'), 'center')
            elif valign_lower in ['top']:
                vAlign.set(qn('w:val'), 'top')
            elif valign_lower in ['bottom']:
                vAlign.set(qn('w:val'), 'bottom')
            else:
                # For any unrecognized value (including 'left'), default to top
                # since 'left' is not a valid vertical alignment
                vAlign.set(qn('w:val'), 'top')

            tcPr.append(vAlign)

        # Apply font properties
        font_size = style.get('font_size')
        if any(style.get(k) for k in ['bold', 'italic', 'color']) or font_size:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if style.get('bold'):
                        run.bold = True
                    if style.get('italic'):
                        run.italic = True

                    # Apply font size
                    if font_size:
                        run.font.size = Pt(font_size)

                    if color_hex := style.get('color'):
                        color_hex = color_hex.replace('#', '')
                        if len(color_hex) == 6:
                            rgb = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
                            run.font.color.rgb = RGBColor(*rgb)

    def _prevent_row_page_break(self, row) -> None:
        """Prevent page break in table row"""
        try:
            # Get the table row element
            tr = row._tr

            # Get or create table row properties
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                tr.insert(0, trPr)

            # Remove existing cantSplit if present
            for child in trPr:
                if child.tag.endswith('cantSplit'):
                    trPr.remove(child)

            # Add cantSplit property to prevent row from splitting across pages
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '1')  # Explicitly set to true
            trPr.append(cantSplit)

        except Exception:
            pass  # Silently ignore if we can't set the property

    def _set_paragraph_keep_with_next(self, paragraph) -> None:
        """Set paragraph to keep with next element"""
        try:
            p = paragraph._element
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p.insert(0, pPr)

            # Add keepNext property
            keepNext = OxmlElement('w:keepNext')
            keepNext.set(qn('w:val'), '1')
            pPr.append(keepNext)

            # Also add keepLines to prevent splitting
            keepLines = OxmlElement('w:keepLines')
            keepLines.set(qn('w:val'), '1')
            pPr.append(keepLines)

            # Set page break before to false
            pageBreakBefore = OxmlElement('w:pageBreakBefore')
            pageBreakBefore.set(qn('w:val'), '0')
            pPr.append(pageBreakBefore)

        except Exception:
            pass

    def _set_table_grid_widths(self, table, cell_styles) -> None:
        """Set table grid column widths"""
        try:
            from docx.shared import Cm, Inches, Pt
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            tbl = table._tbl

            # First, ensure we have tblPr element
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            # Find or create tblGrid - it should be right after tblPr
            tblGrid = tbl.find(qn('w:tblGrid'))
            if tblGrid is None:
                tblGrid = OxmlElement('w:tblGrid')
                # Insert after tblPr (position 1 if tblPr exists at 0)
                tbl.insert(1, tblGrid)

            # Clear existing gridCol elements
            for child in list(tblGrid):
                if child.tag.endswith('gridCol'):
                    tblGrid.remove(child)

            # Add new gridCol elements with widths
            # Also set column widths directly on the table columns
            for i, style in enumerate(cell_styles):
                gridCol = OxmlElement('w:gridCol')
                width_set = False

                if 'width' in style:
                    width = style['width']
                    # Convert to twips (1/20 point, 1 cm = 567 twips)
                    if isinstance(width, str):
                        if width.endswith('cm'):
                            width_twips = int(float(width[:-2]) * 567)
                            # Also set column width if we have access to columns
                            if i < len(table.columns):
                                table.columns[i].width = Cm(float(width[:-2]))
                        elif width.endswith('in'):
                            width_twips = int(float(width[:-2]) * 1440)
                            if i < len(table.columns):
                                table.columns[i].width = Inches(float(width[:-2]))
                        elif width.endswith('pt'):
                            width_twips = int(float(width[:-2]) * 20)
                            if i < len(table.columns):
                                table.columns[i].width = Pt(float(width[:-2]))
                        else:
                            # Assume cm
                            width_twips = int(float(width) * 567)
                            if i < len(table.columns):
                                table.columns[i].width = Cm(float(width))
                    else:
                        # Assume cm
                        width_twips = int(float(width) * 567)
                        if i < len(table.columns):
                            table.columns[i].width = Cm(float(width))

                    gridCol.set(qn('w:w'), str(width_twips))
                    width_set = True

                tblGrid.append(gridCol)

        except Exception as e:
            pass  # Silently fail to avoid debug output

    def _set_table_no_page_break(self, table) -> None:
        """Set table properties to prevent page breaks"""
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            tbl = table._tbl
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            # Set table to not allow row to break across pages
            tblPrEx = OxmlElement('w:tblPrEx')
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '1')
            tblPrEx.append(cantSplit)
            tbl.append(tblPrEx)

            # Remove existing tblLayout if present
            for child in list(tblPr):
                if child.tag.endswith('tblLayout'):
                    tblPr.remove(child)

            # Set table layout to fixed to preserve column widths
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)

        except Exception:
            pass

    def _set_cell_padding(self, cell, padding_config: Dict[str, Any]) -> None:
        """Set cell padding/margins with configuration"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Remove existing margins
        for child in list(tcPr):
            if child.tag.endswith('tcMar'):
                tcPr.remove(child)

        # Create new margins element
        tcMar = OxmlElement('w:tcMar')

        # Convert padding values to twips (1 point = 20 twips)
        # Default Word padding is usually around 0.08" (115 twips)

        # Helper function to convert various units to twips
        def to_twips(value):
            if isinstance(value, str):
                if value.endswith('pt'):
                    return int(float(value[:-2]) * 20)
                elif value.endswith('cm'):
                    return int(float(value[:-2]) * 567)
                elif value.endswith('in'):
                    return int(float(value[:-2]) * 1440)
                elif value.endswith('px'):
                    return int(float(value[:-2]) * 15)
                else:
                    # Assume points if no unit
                    return int(float(value) * 20)
            else:
                # Numeric value - assume points
                return int(float(value) * 20)

        # Set each margin
        sides = {
            'top': 'top',
            'bottom': 'bottom',
            'left': 'start',  # In Word XML, left is 'start'
            'right': 'end'     # In Word XML, right is 'end'
        }

        for padding_side, word_side in sides.items():
            if padding_side in padding_config:
                margin_val = to_twips(padding_config[padding_side])
                margin_element = OxmlElement(f'w:{word_side}')
                margin_element.set(qn('w:w'), str(margin_val))
                margin_element.set(qn('w:type'), 'dxa')  # dxa = twips
                tcMar.append(margin_element)

        # If we have any margins set, add the tcMar element
        if len(tcMar):
            tcPr.append(tcMar)

    def _set_cell_borders(self, cell, border_config) -> None:
        """Set cell borders with configuration"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Remove existing borders
        for child in list(tcPr):
            if child.tag.endswith('tcBorders'):
                tcPr.remove(child)

        # Create new borders element
        tcBorders = OxmlElement('w:tcBorders')

        # Handle "borders": "none" - remove all borders
        if isinstance(border_config, str) and border_config.lower() == 'none':
            # Set all sides to nil
            for side in ['top', 'bottom', 'left', 'right']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)
            return

        # Check if we have individual border specifications
        has_individual_borders = any(key in border_config for key in ['top', 'bottom', 'left', 'right'])

        if has_individual_borders:
            # Handle individual border specifications
            for side in ['top', 'bottom', 'left', 'right']:
                if side in border_config:
                    side_config = border_config[side]
                    # Handle both dict and direct value formats
                    if isinstance(side_config, dict):
                        color = side_config.get('color', '000000').replace('#', '')
                        size_raw = side_config.get('size', 0.5)
                        size = str(int(size_raw * 8))
                        style = side_config.get('style', 'single')
                    else:
                        # If it's just a number, use it as size
                        color = '000000'
                        size_raw = side_config if isinstance(side_config, (int, float)) else 0.5
                        size = str(int(size_raw * 8))
                        style = 'single'

                    # Create border element
                    border = OxmlElement(f'w:{side}')

                    if size == '0':
                        # Explicitly remove border by setting style to nil
                        border.set(qn('w:val'), 'nil')
                    else:
                        # Add border with specified style, size, and color
                        border.set(qn('w:val'), style)
                        border.set(qn('w:sz'), size)
                        border.set(qn('w:color'), color)
                        border.set(qn('w:space'), '0')

                    tcBorders.append(border)
        else:
            # Use the old format (uniform borders)
            border_color = border_config.get('color', '000000').replace('#', '')
            border_size = str(int(border_config.get('size', 0.5) * 8))
            border_style = border_config.get('style', 'single')

            # Define which borders to apply
            borders_to_apply = border_config.get('sides', ['top', 'bottom', 'left', 'right'])

            # Create border elements
            for side in borders_to_apply:
                if side in ['top', 'bottom', 'left', 'right']:
                    border = OxmlElement(f'w:{side}')
                    border.set(qn('w:val'), border_style)
                    border.set(qn('w:sz'), border_size)
                    border.set(qn('w:color'), border_color)
                    border.set(qn('w:space'), '0')
                    tcBorders.append(border)

        # Only append tcBorders if it has child elements
        if len(tcBorders) > 0:
            tcPr.append(tcBorders)

    def _set_cell_bg_fast(self, cell, color: str) -> None:
        """Fast background setting"""
        color = color.replace('#', '')
        tc_pr = cell._tc.get_or_add_tcPr()

        if existing := tc_pr.find(qn('w:shd')):
            tc_pr.remove(existing)

        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        tc_pr.append(shd)

    def save(self, output_path: str) -> None:
        """Save the document"""
        self.doc.save(output_path)

    def replace_from_json_file(self, json_path: str) -> None:
        """Load JSON and replace"""
        with open(json_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        self.replace_from_json(json_data)


def replace_docx_template(docx_path: str, json_data: Dict[str, Any], output_path: str) -> None:
    """Utility function to replace template and save in one step"""
    replacer = DocxReplacer(docx_path)
    replacer.replace_from_json(json_data)
    replacer.save(output_path)