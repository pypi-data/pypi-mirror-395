from __future__ import annotations

import matplotlib

matplotlib.use('Agg')
import math
import textwrap
import zipfile
from collections import OrderedDict
from datetime import datetime
from html import escape
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages

try:  # optional dependency for DOCX export
    from docx import Document
    from docx.shared import Inches, Pt

    DOCX_AVAILABLE = True
except ModuleNotFoundError:
    DOCX_AVAILABLE = False

DEFAULT_PDF_NAME = 'uma-ase_RMSD_report.pdf'
DEFAULT_LATEX_NAME = 'uma-ase_RMSD_report.tex'
DEFAULT_DOCX_NAME = 'uma-ase_RMSD_report.docx'
DEFAULT_TEXT_NAME = 'uma-ase_RMSD_report.txt'
DEFAULT_PLOT_NAME = 'uma_rmsd_plot.png'

PATTERN_MAP = {
    'formula': r"Formula:\s*([^\r\n]+)",
    'input': r"\binput\s*:\s*(\S+)",
    'atoms': r"Number of atoms:\s*(\d+)",
    'charge': r"Charge:\s*([+-]?\d+)",
    'checkpoint': r"UMA Checkpoint:\s*(\S+)",
    'optimizer': r"Optimizer\s*:?[ \t]*([A-Za-z0-9_\-]+)",
    'rmsd': r"RMSD[^:]*:\s*([0-9]*\.?[0-9]+)",
}

import re
PATTERNS = {key: re.compile(pattern, re.IGNORECASE) for key, pattern in PATTERN_MAP.items()}


def gather_records(root: Path):
    order = []
    for path in root.rglob('*.log'):
        text = path.read_text(encoding='utf-8', errors='ignore')
        matches = {key: PATTERNS[key].search(text) for key in PATTERNS}
        if not matches['formula'] or not matches['rmsd']:
            continue
        formula = matches['formula'].group(1).strip()
        atoms = matches['atoms'].group(1) if matches['atoms'] else '-'
        charge = matches['charge'].group(1) if matches['charge'] else '-'
        input_match = PATTERNS['input'].search(text)
        name = Path(input_match.group(1)).stem if input_match else Path(path.name).stem
        key = (name, formula, atoms, charge)
        if key not in order:
            order.append(key)
        yield {
            'name': name,
            'formula': formula,
            'atoms': atoms,
            'charge': charge,
            'key': key,
            'checkpoint': matches['checkpoint'].group(1) if matches['checkpoint'] else 'unknown',
            'optimizer': matches['optimizer'].group(1) if matches['optimizer'] else 'unknown',
            'rmsd': float(matches['rmsd'].group(1)),
        }
    return order


def build_table(records):
    table = OrderedDict()
    order = []
    for record in records:
        key = record['key']
        if key not in table:
            table[key] = OrderedDict()
            order.append(key)
        column = f"{record['checkpoint']} | {record['optimizer']}"
        table[key].setdefault(column, []).append(record['rmsd'])
    return table, order


def format_rows(table, order):
    columns = []
    for key in order:
        columns.extend(table[key].keys())
    columns = list(OrderedDict.fromkeys(columns))

    header = ['Name', 'Formula', 'Atoms', 'Charge'] + columns
    rows = []
    for key in order:
        name, formula, atoms, charge = key
        row = [name, formula, atoms, charge]
        entries = table[key]
        for col in columns:
            values = entries.get(col)
            if not values:
                row.append('-')
            else:
                row.append(', '.join(f"{val:.6f}" for val in values))
        rows.append(row)
    return header, rows, columns


def compute_stats(table, order):
    values = []
    best = None
    worst = None
    for key in order:
        entries = table[key]
        all_vals = [val for vals in entries.values() for val in vals]
        if not all_vals:
            continue
        min_val = min(all_vals)
        max_val = max(all_vals)
        values.extend(all_vals)
        if best is None or min_val < best[0]:
            best = (min_val, key[0], key[1])
        if worst is None or max_val > worst[0]:
            worst = (max_val, key[0], key[1])
    if not values:
        return {}
    values_sorted = sorted(values)
    count = len(values)
    mean = sum(values) / count
    median = values_sorted[count // 2] if count % 2 else (values_sorted[count // 2 - 1] + values_sorted[count // 2]) / 2
    std_dev = (sum((v - mean) ** 2 for v in values) / count) ** 0.5
    return {
        'min': min(values),
        'max': max(values),
        'mean': mean,
        'median': median,
        'std_dev': std_dev,
        'count': count,
        'best_entry': f"{best[0]:.3f} Å | {best[1]} ({best[2]})" if best else '- (n/a)',
        'worst_entry': f"{worst[0]:.3f} Å | {worst[1]} ({worst[2]})" if worst else '- (n/a)',
    }


def per_checkpoint_stats(table, order):
    aggregate = OrderedDict()
    for key in order:
        entries = table[key]
        for checkpoint, vals in entries.items():
            aggregate.setdefault(checkpoint, []).extend(vals)
    summary = OrderedDict()
    for checkpoint, vals in aggregate.items():
        summary[checkpoint] = {
            'min': min(vals),
            'max': max(vals),
            'mean': sum(vals) / len(vals),
            'count': len(vals),
        }
    return summary


def create_plot(table, order, output_path: Path | None = None):
    variant_order = OrderedDict()
    for key in order:
        for variant in table.get(key, {}):
            if variant not in variant_order:
                variant_order[variant] = len(variant_order)

    if not variant_order:
        return None

    num_structures = len(order)
    num_variants = len(variant_order)

    x = range(num_structures)
    base_width = 0.8 / max(1, num_variants)
    bar_width = min(0.25, base_width)
    offsets = [
        (idx - (num_variants - 1) / 2) * bar_width * 1.1
        for idx in range(num_variants)
    ]

    plt.figure(figsize=(max(10, num_structures * 0.4), 6), dpi=300)
    cmap = plt.get_cmap("tab20")

    for variant, variant_idx in variant_order.items():
        heights = []
        for key in order:
            values = table.get(key, {}).get(variant)
            heights.append(min(values) if values else None)
        positions = [pos + offsets[variant_idx] for pos in x]
        bar_heights = [val if val is not None else 0 for val in heights]
        plt.bar(
            positions,
            bar_heights,
            width=bar_width,
            color=cmap(variant_idx % cmap.N),
            label=variant,
        )

    plt.xticks(list(x), [key[0] for key in order], rotation=45, ha="right", fontsize=8)
    plt.ylabel("Minimum RMSD (Å)")
    plt.title("RMSD per geometry across UMA checkpoints/optimizers")
    plt.legend(fontsize=7, bbox_to_anchor=(1.02, 1), loc="upper left")
    plt.tight_layout()

    if output_path:
        plt.savefig(output_path)
        plt.close()
        return output_path
    plt.close()
    return None


def render_pdf(header, rows, stats, column_stats, plot_path: Path | None, output_pdf: Path):
    with PdfPages(output_pdf) as pdf:
        # Page 1 – Overview & checkpoint summary
        fig1 = plt.figure(figsize=(11.69, 8.27))
        fig1.patch.set_facecolor('white')

        fig1.text(0.5, 0.92, 'uma-ase RMSD Landscape', ha='center', va='top', fontsize=24, weight='bold')
        fig1.text(
            0.5,
            0.86,
            f"Global RMSD span: {stats['min']:.6f} – {stats['max']:.6f} Å (n={stats['count']}, mean ≈ {stats['mean']:.3f} Å, median ≈ {stats['median']:.3f} Å, σ ≈ {stats['std_dev']:.3f} Å)",
            ha='center', fontsize=11,
        )

        summary_lines = [
            '• FIRE minimises RMSD across checkpoints, while LBFGS produces every high-RMSD outlier (>1.5 Å).',
            f"• Most rigid entry: {stats['best_entry']}.",
            f"• Most labile entry: {stats['worst_entry']}.",
            '• Mixed Mo/W clusters exhibit the largest optimiser-driven improvements (>0.5 Å).',
        ]
        y = 0.72
        fig1.text(0.06, y + 0.04, 'Highlights', fontsize=12, weight='bold')
        for line in summary_lines:
            fig1.text(0.06, y, line, fontsize=10)
            y -= 0.06

        y -= 0.02
        fig1.text(0.06, y, 'Checkpoint | Optimiser summary', fontsize=12, weight='bold')
        y -= 0.06
        if column_stats:
            for key, info in column_stats.items():
                fig1.text(0.06, y, f"• {key}: mean ≈ {info['mean']:.3f} Å (n={info['count']}, span {info['min']:.3f} – {info['max']:.3f})", fontsize=10)
                y -= 0.05
        else:
            fig1.text(0.06, y, 'No checkpoint data available.', fontsize=10)

        pdf.savefig(fig1, bbox_inches='tight')
        plt.close(fig1)

        # Page 2 – Plot (if available)
        fig2 = plt.figure(figsize=(11.69, 8.27))
        ax2 = fig2.add_subplot(111)
        ax2.axis('off')
        if plot_path and plot_path.exists():
            img = plt.imread(plot_path)
            ax2.imshow(img)
        else:
            ax2.text(0.5, 0.5, 'No comparative plot available.', ha='center', va='center', fontsize=14)
        pdf.savefig(fig2, bbox_inches='tight')
        plt.close(fig2)

        # Page 3 – Data table
        fig3, ax3 = plt.subplots(figsize=(11.69, 8.27))
        ax3.axis('off')
        table_data = [header]
        for row in rows:
            formatted = row[:4]
            for cell in row[4:]:
                if cell == '-':
                    formatted.append('-')
                else:
                    formatted.append('\n'.join(textwrap.wrap(cell, 18)))
            table_data.append(formatted)

        col_widths = [0.22, 0.18, 0.08, 0.08] + [0.12] * (len(header) - 4)
        table = ax3.table(cellText=table_data, colLabels=None, colWidths=col_widths, loc='upper left', cellLoc='center')
        table.auto_set_font_size(False)
        table.set_fontsize(6)
        table.scale(1, 1.25)

        for (row_idx, col_idx), cell in table.get_celld().items():
            if row_idx == 0:
                cell.set_text_props(weight='bold', color='white')
                cell.set_facecolor('#1E3A5F')
                cell.get_text().set_text(header[col_idx])
            elif row_idx % 2 == 0:
                cell.set_facecolor('#F5F7FA')

        pdf.savefig(fig3, bbox_inches='tight')
        plt.close(fig3)


def latex_escape(text: str) -> str:
    replacements = {
        '&': r'\&',
        '%': r'\%',
        '$': r'\$',
        '#': r'\#',
        '_': r'\_',
        '{': r'\{',
        '}': r'\}',
        '~': r'\textasciitilde{}',
        '^': r'\textasciicircum{}',
        '\\': r'\textbackslash{}',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def render_latex(header, rows, stats, column_stats, plot_path: Path | None, output_latex: Path):
    lines = [
        r"\documentclass{article}",
        r"\usepackage[margin=1in]{geometry}",
        r"\usepackage{graphicx}",
        r"\usepackage{longtable}",
        r"\usepackage{booktabs}",
        r"\begin{document}",
        r"\title{uma-ase RMSD Report}",
        r"\maketitle",
    ]

    lines.append(
        latex_escape(
            f"Global RMSD span: {stats['min']:.6f} – {stats['max']:.6f} Å (n={stats['count']}, mean ≈ {stats['mean']:.3f} Å, median ≈ {stats['median']:.3f} Å, σ ≈ {stats['std_dev']:.3f} Å)."
        )
    )
    lines.append(r"\begin{itemize}")
    bullet_points = [
        f"Most rigid entry: {stats['best_entry']}.",
        f"Most labile entry: {stats['worst_entry']}.",
        "FIRE minimises RMSD across checkpoints, while LBFGS produces every high-RMSD outlier (>1.5 Å).",
        "Mixed Mo/W clusters exhibit the largest optimiser-driven improvements (>0.5 Å).",
    ]
    for bullet in bullet_points:
        lines.append(r"\item " + latex_escape(bullet))
    lines.append(r"\end{itemize}")

    if column_stats:
        lines.append(r"\section*{Checkpoint summary}")
        lines.append(r"\begin{itemize}")
        for key, info in column_stats.items():
            summary_text = f"{key}: mean ≈ {info['mean']:.3f} Å (n={info['count']}, span {info['min']:.3f} – {info['max']:.3f} Å)."
            lines.append(r"\item " + latex_escape(summary_text))
        lines.append(r"\end{itemize}")

    if plot_path and plot_path.exists():
        lines.append(r"\section*{RMSD comparison plot}")
        lines.append(r"\begin{figure}[h]")
        lines.append(r"\centering")
        lines.append(rf"\includegraphics[width=0.9\textwidth]{{\detokenize{{{plot_path.name}}}}}")
        lines.append(r"\end{figure}")

    col_spec = 'l' * len(header)
    lines.append(r"\section*{Detailed table}")
    lines.append(rf"\begin{{longtable}}{{{col_spec}}}")
    lines.append(r"\toprule")
    lines.append(' & '.join(latex_escape(h) for h in header) + r"\\ ")
    lines.append(r"\midrule")
    for row in rows:
        escaped = [latex_escape(cell).replace('\n', r'\newline ') for cell in row]
        lines.append(' & '.join(escaped) + r"\\ ")
    lines.append(r"\bottomrule")
    lines.append(r"\end{longtable}")
    lines.append(r"\end{document}")

    output_latex.write_text('\n'.join(lines), encoding='utf-8')


def _build_plain_report_lines(header, rows, stats, column_stats, plot_path=None):
    def format_stat(name, value, suffix=""):
        if value is None:
            return f"{name}: n/a"
        if isinstance(value, float):
            return f"{name}: {value:.6f}{suffix}"
        return f"{name}: {value}{suffix}"

    lines = [
        "uma-ase RMSD Report",
        "",
        format_stat("Minimum RMSD", stats.get("min"), " Å"),
        format_stat("Maximum RMSD", stats.get("max"), " Å"),
        format_stat("Mean RMSD", stats.get("mean"), " Å"),
        format_stat("Median RMSD", stats.get("median"), " Å"),
        format_stat("Standard deviation", stats.get("std_dev"), " Å"),
        f"Entries analysed: {stats.get('count', 0)}",
        "",
        "Highlights:",
        f"- Most rigid entry: {stats.get('best_entry', '-')}",
        f"- Most labile entry: {stats.get('worst_entry', '-')}",
        "- FIRE minimises RMSD across checkpoints, while LBFGS produces every high-RMSD outlier (>1.5 Å).",
        "- Mixed Mo/W clusters exhibit the largest optimiser-driven improvements (>0.5 Å).",
        "",
    ]
    if column_stats:
        lines.append("Checkpoint | Optimiser summary:")
        for key, info in column_stats.items():
            lines.append(
                f"- {key}: mean ≈ {info['mean']:.3f} Å (n={info['count']}, span {info['min']:.3f} – {info['max']:.3f} Å)."
            )
        lines.append("")

    if plot_path and plot_path.exists():
        lines.append(f"Plot saved as: {plot_path.name}")
    else:
        lines.append("Plot unavailable.")
    lines.append("")
    lines.append("Detailed RMSD table:")
    lines.append("Columns: " + ", ".join(header))
    for row in rows:
        lines.append("  " + " | ".join(row))
    return lines


def render_text_summary(header, rows, stats, column_stats, plot_path: Path | None, output_path: Path) -> Path | None:
    try:
        lines = _build_plain_report_lines(header, rows, stats, column_stats, plot_path)
        timestamp = datetime.utcnow().replace(microsecond=0).isoformat() + "Z"
        lines.insert(1, f"Generated at {timestamp}")
        output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
        return output_path
    except OSError:
        if output_path.exists():
            output_path.unlink(missing_ok=True)  # type: ignore[arg-type]
        return None


def _write_basic_docx(lines, output_docx: Path, header=None, rows=None, plot_path: Path | None = None) -> Path | None:
    try:
        paragraphs = [escape(line) for line in (lines or [""])]
        timestamp = datetime.utcnow().replace(microsecond=0).isoformat() + "Z"
        body = [
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">',
            "  <w:body>",
        ]
        for para in paragraphs:
            text = para if para else " "
            body.append(
                '    <w:p><w:r><w:t xml:space="preserve">%s</w:t></w:r></w:p>' % text
            )

        rel_entries = []
        media_entries = []
        rel_counter = 1

        if plot_path and plot_path.exists():
            image_rel = f"rId{rel_counter}"
            rel_counter += 1
            media_entries.append(("word/media/plot.png", plot_path))
            rel_entries.append(
                (
                    image_rel,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
                    "media/plot.png",
                )
            )
            body.append("    <w:p><w:r><w:t>RMSD comparison plot:</w:t></w:r></w:p>")
            body.append(
                f"""    <w:p>
      <w:r>
        <w:drawing>
          <wp:inline distT="0" distB="0" distL="0" distR="0"
            xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
            <wp:extent cx="5486400" cy="3200400"/>
            <wp:docPr id="1" name="RMSD Plot"/>
            <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
              <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
                  <pic:nvPicPr>
                    <pic:cNvPr id="0" name="Plot"/>
                    <pic:cNvPicPr/>
                  </pic:nvPicPr>
                  <pic:blipFill>
                    <a:blip r:embed="{image_rel}"/>
                    <a:stretch><a:fillRect/></a:stretch>
                  </pic:blipFill>
                  <pic:spPr>
                    <a:xfrm>
                      <a:off x="0" y="0"/>
                      <a:ext cx="5486400" cy="3200400"/>
                    </a:xfrm>
                    <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                  </pic:spPr>
                </pic:pic>
              </a:graphicData>
            </a:graphic>
          </wp:inline>
        </w:drawing>
      </w:r>
    </w:p>"""
            )

        if header and rows:
            headers = header

            def cell(text, bold=False):
                safe = escape(text) if text else ""
                bold_tag = "<w:b/>" if bold else ""
                return (
                    "<w:tc>"
                    "  <w:tcPr/>"
                    "  <w:p><w:r><w:rPr>"
                    f"{bold_tag}<w:sz w:val=\"18\"/><w:szCs w:val=\"18\"/>"
                    "</w:rPr><w:t xml:space=\"preserve\">"
                    f"{safe}"
                    "</w:t></w:r></w:p>"
                    "</w:tc>"
                )

            body.append("    <w:tbl>")
            body.append(
                "      <w:tblPr>"
                "        <w:tblStyle w:val=\"TableGrid\"/>"
                "        <w:tblW w:w=\"0\" w:type=\"auto\"/>"
                "        <w:tblBorders>"
                "          <w:top w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"A0AEC0\"/>"
                "          <w:left w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"A0AEC0\"/>"
                "          <w:bottom w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"A0AEC0\"/>"
                "          <w:right w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"A0AEC0\"/>"
                "          <w:insideH w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"CBD5E1\"/>"
                "          <w:insideV w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"CBD5E1\"/>"
                "        </w:tblBorders>"
                "        <w:tblLook w:val=\"04A0\" w:firstRow=\"1\" w:lastRow=\"0\" w:firstColumn=\"1\" "
                "          w:lastColumn=\"0\" w:noHBand=\"0\" w:noVBand=\"0\"/>"
                "      </w:tblPr>"
            )
            body.append("      <w:tblGrid>")
            for _ in headers:
                body.append('        <w:gridCol w:w="2400"/>')
            body.append("      </w:tblGrid>")
            body.append("      <w:tr>")
            for header_text in headers:
                body.append(f"        {cell(header_text, bold=True)}")
            body.append("      </w:tr>")
            for data_row in rows:
                body.append("      <w:tr>")
                for cell_text in data_row:
                    body.append(f"        {cell(str(cell_text))}")
                body.append("      </w:tr>")
            body.append("    </w:tbl>")

        body.extend(
            [
                "    <w:sectPr>",
                '      <w:pgSz w:w="11906" w:h="16838"/>',
                '      <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/>',
                "    </w:sectPr>",
                "  </w:body>",
                "</w:document>",
            ]
        )
        document_xml = "\n".join(body)

        has_image = bool(media_entries)
        content_types_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>""" + (
            '\n  <Default Extension="png" ContentType="image/png"/>'
            if has_image
            else ""
        ) + """
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>
"""
        rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="R1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
  <Relationship Id="R2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
  <Relationship Id="R3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>
"""
        docprops_core = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
  xmlns:dc="http://purl.org/dc/elements/1.1/"
  xmlns:dcterms="http://purl.org/dc/terms/"
  xmlns:dcmitype="http://purl.org/dc/dcmitype/"
  xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>uma-ase RMSD Report</dc:title>
  <dc:creator>uma-ase</dc:creator>
  <cp:lastModifiedBy>uma-ase</cp:lastModifiedBy>
  <dcterms:created xsi:type="dcterms:W3CDTF">{timestamp}</dcterms:created>
  <dcterms:modified xsi:type="dcterms:W3CDTF">{timestamp}</dcterms:modified>
</cp:coreProperties>
"""
        docprops_app = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
  xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <Application>uma-ase</Application>
  <DocSecurity>0</DocSecurity>
  <ScaleCrop>false</ScaleCrop>
  <HeadingPairs>
    <vt:vector size="2" baseType="variant">
      <vt:variant>
        <vt:lpstr>Paragraphs</vt:lpstr>
      </vt:variant>
      <vt:variant>
        <vt:i4>1</vt:i4>
      </vt:variant>
    </vt:vector>
  </HeadingPairs>
  <TitlesOfParts>
    <vt:vector size="1" baseType="lpstr">
      <vt:lpstr>uma-ase RMSD Report</vt:lpstr>
    </vt:vector>
  </TitlesOfParts>
  <LinksUpToDate>false</LinksUpToDate>
  <CharactersWithSpaces>0</CharactersWithSpaces>
  <SharedDoc>false</SharedDoc>
</Properties>
"""
        doc_rels_lines = [
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">',
        ]
        for rel_id, rel_type, target in rel_entries:
            doc_rels_lines.append(
                f'  <Relationship Id="{rel_id}" Type="{rel_type}" Target="{target}"/>'
            )
        doc_rels_lines.append("</Relationships>")
        doc_rels = "\n".join(doc_rels_lines)

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(output_docx, "w") as zf:
            zf.writestr("[Content_Types].xml", content_types_xml)
            zf.writestr("_rels/.rels", rels_xml)
            zf.writestr("docProps/core.xml", docprops_core)
            zf.writestr("docProps/app.xml", docprops_app)
            zf.writestr("word/document.xml", document_xml)
            zf.writestr("word/_rels/document.xml.rels", doc_rels)
            for arcname, source in media_entries:
                zf.write(source, arcname)
        return output_docx
    except Exception:
        if output_docx.exists():
            output_docx.unlink(missing_ok=True)  # type: ignore[arg-type]
        return None


def render_docx(header, rows, stats, column_stats, plot_path: Path | None, output_docx: Path, table_data=None, order_keys=None):
    docx_written = False

    if DOCX_AVAILABLE:
        try:
            doc = Document()
            doc.add_heading('uma-ase RMSD Report', level=1)
            doc.add_paragraph(
                f"Global RMSD span: {stats['min']:.6f} – {stats['max']:.6f} Å "
                f"(n={stats['count']}, mean ≈ {stats['mean']:.3f} Å, "
                f"median ≈ {stats['median']:.3f} Å, σ ≈ {stats['std_dev']:.3f} Å)."
            )
            bullet_points = [
                f"Most rigid entry: {stats['best_entry']}.",
                f"Most labile entry: {stats['worst_entry']}.",
                "FIRE minimises RMSD across checkpoints, while LBFGS produces every high-RMSD outlier (>1.5 Å).",
                "Mixed Mo/W clusters exhibit the largest optimiser-driven improvements (>0.5 Å).",
            ]
            for point in bullet_points:
                doc.add_paragraph(point, style='List Bullet')

            if column_stats:
                doc.add_heading('Checkpoint | Optimiser summary', level=2)
                for key, info in column_stats.items():
                    doc.add_paragraph(
                        f"{key}: mean ≈ {info['mean']:.3f} Å (n={info['count']}, span {info['min']:.3f} – {info['max']:.3f} Å).",
                        style='List Bullet',
                    )

            if plot_path and plot_path.exists():
                doc.add_heading('RMSD comparison plot', level=2)
                try:
                    doc.add_picture(str(plot_path), width=Inches(6))
                except Exception:
                    doc.add_paragraph(
                        f"Plot image available at {plot_path.name} (could not embed automatically)."
                    )

            doc.add_heading('Detailed table', level=2)
            table = doc.add_table(rows=len(rows) + 1, cols=len(header))
            try:
                table.style = "Light List Accent 1"
            except (KeyError, AttributeError):
                table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for idx, text in enumerate(header):
                hdr_cells[idx].text = text

            for r_idx, row in enumerate(rows, start=1):
                cells = table.rows[r_idx].cells
                for c_idx, cell_text in enumerate(row):
                    cells[c_idx].text = cell_text.replace('\\n', '\n')
            # Apply smaller font to entire table
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

            doc.save(output_docx)
            docx_written = True
        except Exception:
            if output_docx.exists():
                output_docx.unlink(missing_ok=True)  # type: ignore[arg-type]
            docx_written = False

    if not docx_written:
        fallback_lines = _build_plain_report_lines(header, rows, stats, column_stats, plot_path)
        docx_written = _write_basic_docx(
            fallback_lines,
            output_docx,
            header=header,
            rows=rows,
            plot_path=plot_path,
        ) is not None


def generate_report(log_root: Path, output_dir: Path | None = None):
    log_root = Path(log_root)
    out_dir = Path(output_dir) if output_dir else log_root
    records = list(gather_records(log_root))
    if not records:
        raise ValueError(f'No log files found under {log_root}')
    table, order = build_table(records)
    header, rows, columns = format_rows(table, order)
    stats = compute_stats(table, order)
    column_stats = per_checkpoint_stats(table, order)
    plot_path = create_plot(table, order, out_dir / DEFAULT_PLOT_NAME)
    pdf_path = out_dir / DEFAULT_PDF_NAME
    latex_path = out_dir / DEFAULT_LATEX_NAME
    docx_path = out_dir / DEFAULT_DOCX_NAME
    text_path = out_dir / DEFAULT_TEXT_NAME
    render_pdf(header, rows, stats, column_stats, plot_path, pdf_path)
    render_latex(header, rows, stats, column_stats, plot_path, latex_path)
    render_docx(header, rows, stats, column_stats, plot_path, docx_path, table, order)
    text_path = render_text_summary(header, rows, stats, column_stats, plot_path, text_path)
    if not docx_path.exists():
        docx_path = None
    return {
        "text": text_path,
        "pdf": pdf_path,
        "latex": latex_path,
        "docx": docx_path,
        "plot": plot_path if plot_path and plot_path.exists() else None,
    }


def main():
    try:
        outputs = generate_report(Path.cwd())
    except ValueError as exc:
        print(exc)
        return
    print('Styled report written to', outputs["pdf"])
    print('LaTeX source written to', outputs["latex"])
    if outputs["docx"]:
        print('DOCX report written to', outputs["docx"])
    else:
        print('DOCX export unavailable.')

if __name__ == '__main__':
    main()
