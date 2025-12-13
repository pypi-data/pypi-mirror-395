from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx import Document, shared

Default_indent = shared.Inches(0.3)


def get_styles(doc: Document):
    """
    Not strictly necessary like ReportLab, but we configure consistent fonts here.
    """
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc.styles


def add_name_header(doc: Document, name: str):
    p = doc.add_heading(level=0)
    run = p.add_run(name)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_info_bar(doc: Document, content: list[str]):
    info = " | ".join(content)
    p = doc.add_paragraph(info)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Spacer (empty line)
    doc.add_paragraph("")


def add_section_header(doc: Document, title: str):
    p = doc.add_heading(level=1)
    run = p.add_run(title)
    run.bold = True


def add_subheader(doc: Document, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    

def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Default_indent



def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.left_indent = Default_indent


def add_education_line(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Default_indent


def add_two_column_table(doc: Document, left_items: list[str], right_items: list[str]):
    """
    Creates a 2-column table for skills, with indentation aligned to other content.
    """
    table = doc.add_table(rows=len(left_items), cols=2)

    for i, (left, right) in enumerate(zip(left_items, right_items)):
        left_cell = table.cell(i, 0)
        right_cell = table.cell(i, 1)

        # Set cell text
        left_cell.text = f"• {left}" if left else ""
        right_cell.text = f"• {right}" if right else ""

        # Apply indent to EACH cell's paragraph (cells default to 1 paragraph)
        for paragraph in left_cell.paragraphs:
            paragraph.paragraph_format.left_indent = Default_indent

        for paragraph in right_cell.paragraphs:
            paragraph.paragraph_format.left_indent = Default_indent
