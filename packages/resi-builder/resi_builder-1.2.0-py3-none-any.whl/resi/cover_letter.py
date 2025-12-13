from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from .open_ai_writer import cover_letter_generator
from .utils import pdf_utils, word_utils
import json
from typing import Union
import copy
import os
from docx import Document, shared

def build_cover_letter_preview(
        job_desc: str,
        user_history: Union[str, dict],
        hiring_manager: str = 'Hiring Manager',
        additional_prompts: str = '',
    ) -> dict:
    """
    Build the cover letter data preview dictionary.

    :param job_desc: Job Description
    :param user_history: Either a dictionary of the user's resume work history,
                         or a path to a JSON file containing that dictionary.
    :param hiring_manager: (Optional) Name of the Hiring manager
    :param additional_prompts: (Optional) Additional prompts for the LLM
    :return: Cover Letter preview dictionary
    """

    # Normalize input: if user_history is a str, load JSON file
    if isinstance(user_history, str):
        with open(user_history, "r") as f:
            user_history = json.load(f)

    user_history_copy = copy.deepcopy(user_history)

    # Delete the contact info to avoid passing personal data to the LLM except for name
    del user_history_copy['contact_info']
    user_history_copy['contact_info'] = {'name': user_history['contact_info']['name']}

    # Step 1: Generate initial cover letter text
    body = cover_letter_generator(
        job_desc,
        user_history_copy,
        additional_prompts
    )

    # Step 2: Write it to a temp text file for approval
    paragraphs = body.strip().split("\n\n")

    # Step 3: return the file for preview
    body = {
        'intro': f"Dear {hiring_manager},",
        'paragraphs': {k:v for k,v in enumerate(paragraphs)},
    }

    return body


def build_cover_letter_pdf(
        cover_letter_data: dict,
        user_history: Union[str, dict],
        file_name: str = 'cover_letter.pdf'
    ) -> None:
    """
    Build the cover letter as a pdf file.

    :param cover_letter_data: Dictionary containing cover letter data
    :param user_history: Either a dictionary of the user's resume work history.
                         or a path to a JSON file containing that dictionary.
    :param file_name: (Optional) file name of the output. This can be a path to the output
    :return: PDF cover letter file
    """

    # Normalize input: if user_history is a str, load JSON file
    if isinstance(user_history, str):
        with open(user_history, "r") as f:
            user_history = json.load(f)

    # Check if there is a file name provided and normalize to pdf
    base, ext = os.path.splitext(file_name)
    if ext.lower() != '.pdf':
        file_name = f"{base}.pdf"

    # Remove any spaces at the end
    paragraphs = [x.strip() for x in cover_letter_data['paragraphs'].values()]

    # Build the PDF
    doc = SimpleDocTemplate(file_name, pagesize=LETTER, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=24)
    Story = []

    styles = pdf_utils.get_styles()
    # Add name title
    pdf_utils.add_name_header(Story, styles, user_history['contact_info']['name'])
    # Add info bar
    pdf_utils.add_info_bar(Story, styles, [x for x in user_history['contact_info'].values()])

    # Greeting
    Story.append(Paragraph(cover_letter_data['intro'], styles['CustomBodyText']))
    Story.append(Spacer(1, 12))  # Paragraph break

    # Paragraph section
    for para in paragraphs:
        Story.append(Paragraph(para, styles['CustomBodyText']))
        Story.append(Spacer(1, 12))  # Paragraph break

    # Sign off
    Story.append(Paragraph(f"Sincerely<br/>{user_history['contact_info']['name']}", styles['CustomBodyText']))

    doc.build(Story)
    print(f"Cover letter generated: {file_name}")

    
def build_cover_letter_word(
        cover_letter_data: dict,
        user_history: Union[str, dict],
        file_name: str = 'cover_letter.docx'
) -> None:
    """
    Build the cover letter as an MS Word file.

    :param cover_letter_data: Dictionary containing cover letter data
    :param user_history: Either a dictionary of the user's resume work history.
                         or a path to a JSON file containing that dictionary.
    :param file_name: (Optional) file name of the output. This can be a path to the output
    :return: Word cover letter file
    """

    # Normalize input: if user_history is a str, load JSON file
    if isinstance(user_history, str):
        with open(user_history, "r") as f:
            user_history = json.load(f)

    # Check if there is a file name provided and normalize to pdf
    base, ext = os.path.splitext(file_name)
    if ext.lower() != '.docx':
        file_name = f"{base}.docx"

    # Remove any spaces at the end
    paragraphs = [x.strip() for x in cover_letter_data['paragraphs'].values()]

    # Build the Word doc
    doc = Document()

    # Adjust the margins
    section = doc.sections[0]
    section.left_margin = shared.Inches(0.5)
    section.right_margin = shared.Inches(0.5)
    section.top_margin = shared.Inches(0.5)
    section.bottom_margin = shared.Inches(0.5)

    # Name Header
    word_utils.add_name_header(doc, user_history['contact_info']['name'])

    # Info Bar
    contact_values = [x for x in user_history['contact_info'].values()]
    word_utils.add_info_bar(doc, contact_values)

    # Greeting
    word_utils.add_paragraph(doc, cover_letter_data['intro'])

    # Paragraph section
    for para in paragraphs:
        word_utils.add_paragraph(doc, para)

    # Sign off
    word_utils.add_paragraph(doc, f"Sincerely\n{user_history['contact_info']['name']}")

    # Save file
    doc.save(file_name)
    print(f"Cover Letter generated: {file_name}")
    