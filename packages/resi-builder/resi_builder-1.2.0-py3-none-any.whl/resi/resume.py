from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib import colors
from .open_ai_writer import generate_job_bullets, generate_profile
from .utils import pdf_utils, word_utils
import textwrap
from typing import Union
import json
import copy
import os
from docx import Document, shared

def build_resume_preview(
        job_desc: str,
        user_history: Union[str, dict],
        additional_prompts: str = ""
    ) -> dict:
    """
    Build resume data preview dictionary.

    :param job_desc: Job description
    :param user_history: Either a dictionary of the user's resume work history,
                         or a path to a JSON file containing that dictionary.
    :param additional_prompts: (Optional) Additional prompts for the LLM
    :return: Resume preview dictionary
    """

    # Normalize input: if user_history is a str, load JSON file
    if isinstance(user_history, str):
        with open(user_history, "r") as f:
            user_history = json.load(f)

    user_history_copy = copy.deepcopy(user_history)

    # Delete the contact info to avoid passing personal data to the LLM except for name
    del user_history_copy['contact_info']
    user_history_copy['contact_info'] = {'name': user_history['contact_info']['name']}

    # Step 1: Generate initial resume content

    # Check if Profile is missing, if so, generate a profile
    if 'profile' not in user_history:
        user_profile = generate_profile(
            job_desc,
            user_history_copy
        )
    else:
        user_profile = user_history_copy['profile']
    

    wrapped_profile = textwrap.fill(user_profile.strip(), width=80)

    # bullet points
    bullets = generate_job_bullets(
        job_desc,
        user_history_copy,
        additional_prompts
    )

    # skills
    skills = user_history_copy['skills']

    # Step 2: Build preview dictionary
    body = {
        'profile': wrapped_profile,
        'bullets': bullets,
        'skills': skills,
        'education': user_history_copy['education']
    }

    return body

def build_resume_pdf(
        resume_data: dict,
        user_history: Union[str, dict],
        file_name: str = 'resume.pdf'
    ) -> None:
    """
    Build the resume as a pdf file

    :param resume_data: Dictionary containing resume data
    :param user_history: Either a dictionary of the user's resume work history,
                         or a path to a JSON file containing that dictionary.
    :param file_name: (Optional) file name of the output. This can be a path to the output
    :return: PDF resume file
    """

    # Normalize input: if user_history is a str, load JSON file
    if isinstance(user_history, str):
        with open(user_history, "r") as f:
            user_history = json.load(f)

    # Check if there is a file name provided and normalize to pdf
    base, ext = os.path.splitext(file_name)
    if ext.lower() != '.pdf':
        file_name = f"{base}.pdf"

    # Build the PDF
    doc = SimpleDocTemplate(file_name, pagesize=LETTER, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=24)
    Story = []

    styles = pdf_utils.get_styles()

    # Add name title
    pdf_utils.add_name_header(Story, styles, user_history['contact_info']['name'])

    pdf_utils.add_info_bar(Story, styles, [x for x in user_history['contact_info'].values()])

    # Profile
    pdf_utils.add_section(Story, "Profile", styles, content=resume_data['profile'])

    # Experience
    pdf_utils.add_section(Story, "Experience", styles)
    for exp in resume_data['bullets']:
        pdf_utils.add_section(
            Story,
            f"{exp['role'].upper()} | {exp['company'].upper()} | {exp['dates'].upper()}",
            styles,
            bullets=exp['experience']
        )

    # Education
    pdf_utils.add_section(Story, "Education", styles)
    for education in user_history['education']:
            pdf_utils.add_education_section(
                story=Story,
                title=f"{education['degree'].upper()} IN {education['field_of_study'].upper()} | {education['school'].upper()}, {education['location'].upper()}",
                styles=styles,
                spacer_height=5
            )

    # Skills

    # Make sure number of skills is even (pad if needed)
    if len(resume_data['skills']) % 2 != 0:
        resume_data['skills'].append("")

    half = len(resume_data['skills']) // 2
    data = list(zip(
        [f"• {skill}" for skill in resume_data['skills'][:half]],
        [f"• {skill}" if skill else '' for skill in resume_data['skills'][half:]] # ensure that the last value is not displayed if empty
    ))

    table = Table(data, colWidths=[250, 250])  # Adjust widths as needed

    table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
    ]))

    pdf_utils.add_section(story=Story, title="Skills & Abilities", styles=styles)
    Story.append(table)
    Story.append(Spacer(1, 2))

    # Activities
    pdf_utils.add_section(story=Story, title="Activities and Interests", styles=styles, content=f"{user_history['activities_and_interests']}")

    # Build PDF
    doc.build(Story)
    print(f"Resume generated: {file_name}")


def build_resume_word(
        resume_data: dict,
        user_history: Union[str, dict],
        file_name: str = "resume.docx"
) -> None:
    """
    Build the resume as an MS Word file

    :param resume_data: Dictionary containing resume data
    :param user_history: Either a dictionary of the user's resume work history,
                         or a path to a JSON file containing that dictionary.
    :param file_name: (Optional) file name of the output. Can include a path.
    :return: DOCX resume file
    """

    # Normalize input: if user_history is a str, load JSON file
    if isinstance(user_history, str):
        with open(user_history, "r") as f:
            user_history = json.load(f)

    # Ensure .docx extension
    base, ext = os.path.splitext(file_name)
    if ext.lower() != '.docx':
        file_name = f"{base}.docx"

    # Create document
    doc = Document()

    # Adjust the margins
    section = doc.sections[0]
    section.left_margin = shared.Inches(0.5)
    section.right_margin = shared.Inches(0.5)
    section.top_margin = shared.Inches(0.5)
    section.bottom_margin = shared.Inches(0.5)

    # Name Header
    word_utils.add_name_header(doc, user_history['contact_info']['name'])

    # Info Bar
    contact_values = [x for x in user_history['contact_info'].values()]
    word_utils.add_info_bar(doc, contact_values)

    # Profile
    word_utils.add_section_header(doc, "Profile")
    word_utils.add_paragraph(doc, resume_data['profile'])

    # Experience section
    word_utils.add_section_header(doc, "Experience")
    for exp in resume_data['bullets']:
        title = f"{exp['role'].upper()} | {exp['company'].upper()} | {exp['dates'].upper()}"
        word_utils.add_subheader(doc, title)

        for bullet in exp['experience']:
            word_utils.add_bullet(doc, bullet)

    # Education section
    word_utils.add_section_header(doc, "Education")
    for edu in user_history['education']:
        title = f"{edu['degree'].upper()} IN {edu['field_of_study'].upper()} | {edu['school'].upper()}, {edu['location'].upper()}"
        word_utils.add_education_line(doc, title)

    # Skills
    word_utils.add_section_header(doc, "Skills & Abilities")

    # Make skills even count
    skills = resume_data['skills']
    if len(skills) % 2 != 0:
        skills.append("")

    half = len(skills) // 2
    left_col = skills[:half]
    right_col = skills[half:]

    word_utils.add_two_column_table(doc, left_col, right_col)

    # Activities
    word_utils.add_section_header(doc, "Activities and Interests")
    word_utils.add_paragraph(doc, user_history["activities_and_interests"])

    # Save file
    doc.save(file_name)
    print(f"Resume generated: {file_name}")
